from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
XLSX_PATH = OUT / "Clarity_Business_Model_and_Configurable_Rules_Workbook.xlsx"
DOCX_PATH = OUT / "Clarity_Business_Model_and_Rules_Playbook.docx"
PDF_PATH = OUT / "Clarity_Business_Model_and_Rules_Playbook.pdf"

PURPLE = "4F46A3"
LAV = "EEEAFE"
TEAL = "42A89C"
MINT = "E9F8F5"
AMBER = "F6B84B"
AMBER_LIGHT = "FFF4D8"
RED = "DC626C"
RED_LIGHT = "FDEBEC"
DARK = "1F2937"
MUTED = "64748B"
LINE = "D9DEE8"
BLUE_LIGHT = "E8F1FC"


def rule(rule_id, domain, subdomain, question, default, source, links, owner, editable, priority, acceptance, notes=""):
    return {
        "Rule ID": rule_id,
        "Domain": domain,
        "Subdomain": subdomain,
        "Discovery question / rule to review": question,
        "Suggested default or example": default,
        "Data source": source,
        "Must update / link to": links,
        "Configuration owner": owner,
        "Chief editable?": editable,
        "MVP priority": priority,
        "Acceptance check": acceptance,
        "Hospital answer / decision": "",
        "IT build notes": "",
        "Sales / implementation notes": notes,
    }


RULES = [
    rule("PGM-001", "Program Setup", "Academic structure", "How many scheduling blocks exist per academic year?", "13 blocks; some programs may use 12 monthly or 26 two-week blocks.", "Chief/admin setup", "All block selectors, request deadlines, master schedule, monthly builders, analytics", "Chief + IT", "Chief can edit before launch", "MVP", "Changing block count updates every annual and monthly workflow.", "This is one of the first discovery questions."),
    rule("PGM-002", "Program Setup", "Dates", "What are the exact start and end dates for each block?", "Block 1 starts July 1; later blocks follow department calendar.", "Chief/admin setup", "Resident request deadlines, service schedules, holidays, PTO, attendance", "Chief", "Chief editable", "MVP", "Every block has date range and no gaps/overlap unless intentionally allowed."),
    rule("PGM-003", "Program Setup", "Deadline policy", "How many days before a block must residents submit requests?", "90 days before block start.", "Institution rule", "Resident request forms, reminders, chief approvals", "Chief", "Chief editable", "MVP", "Late requests are flagged but chiefs can override."),
    rule("PGM-004", "Program Setup", "Classes", "Which trainee groups need schedules?", "PGY-1, PGY-2, PGY-3, Med-Peds, outside rotators.", "Program configuration", "Profiles, master schedules, service coverage, eligibility", "IT + Chief", "IT config plus chief edit", "MVP", "Every resident belongs to one class/year/track at a time."),
    rule("PGM-005", "Program Setup", "Tracks", "Does the program use tracks with different curriculum requirements?", "Primary care, hospitalist, subspecialty, critical care, categorical, Med-Peds.", "Resident intake + chief rules", "Master builder, elective eligibility, service minimums", "Chief", "Chief editable", "MVP+", "Track-specific requirements appear in master builder and resident profile."),
    rule("PGM-006", "Program Setup", "Institutions", "Which institutions send residents or outside rotators?", "Main institution plus outside rotator institutions.", "Institution profile", "Protected time, didactics, clinics, eligibility, call rules", "Chief/admin", "Chief editable", "MVP", "Every resident has an institution profile or an exception flag."),
    rule("PGM-007", "Program Setup", "User roles", "Who can view, edit, approve, publish, and export schedules?", "Chief admin, chief editor, coordinator, resident, read-only leadership.", "Security design", "All portals, audit log, request decisions", "IT + Leadership", "IT config", "MVP", "Residents cannot see private master schedules of others."),
    rule("PGM-008", "Program Setup", "Import mode", "Can the department upload an already completed master schedule?", "Yes: upload Excel/PDF, map tabs and columns, review unmatched names/rotations.", "Existing Excel/PDF", "Profiles, master schedule, service coverage, monthly builder", "IT + Chief", "Chief runs import; IT builds parser", "MVP+", "Imported assignments populate profiles and builder source data."),

    rule("RES-001", "Resident Profile", "Fixed facts", "What fixed resident facts must be collected once?", "Name, email, PGY, institution, track, outside rotator status, fellowship interest.", "Resident intake / admin upload", "Eligibility, request forms, attendance, call switch, schedule builder", "Resident + Chief", "Resident proposes; chief approves", "MVP", "Profile cannot be used for scheduling until fixed facts are complete."),
    rule("RES-002", "Resident Profile", "Private visibility", "What resident information should be private from other residents?", "Individual master schedule, requests, PTO, leave, sick days, special circumstances.", "Security design", "Resident portal and chief portal", "IT + Leadership", "IT config", "MVP", "Resident sees only their own private data; chiefs see all."),
    rule("RES-003", "Resident Profile", "Profile as source", "Should resident profile data feed schedules automatically?", "Yes. Clinic, didactics, approved PTO, institution, PGY, master block, call pool eligibility.", "Profiles + institution rules", "Master builder, monthly builder, call switches, attendance", "IT", "Automated", "MVP", "Profile changes recalculate affected schedule warnings."),
    rule("RES-004", "Resident Profile", "Profile review tabs", "Which tabs should chiefs see inside each resident profile?", "Overview, clinic plan, master schedule, didactics, PTO/leave/exams, requests/switches, eligibility.", "Product design", "Chief profile screen", "IT + Chief", "Chief editable labels optional", "MVP", "Chief can review all relevant scheduling inputs from one profile."),
    rule("RES-005", "Resident Profile", "Resident view tabs", "Which tabs should residents see?", "My overview, my schedule, my master schedule, published schedules, requests, call switches, PTO/sick leave.", "Product design", "Resident portal", "IT + Leadership", "IT config", "MVP", "Resident portal mirrors approved chief data without exposing others' private details."),

    rule("INT-001", "Resident Intake", "Annual timing", "When do residents submit annual preferences?", "Before chiefs build the annual master schedule.", "Resident annual form", "Chief master builder", "Chief", "Chief editable", "MVP", "Every annual submission has timestamp and status."),
    rule("INT-002", "Resident Intake", "Master ranking", "How many master schedule templates must residents rank?", "All available templates for their PGY year; example 20 options.", "Resident annual form", "Master builder rankings step", "Resident + Chief", "Resident submits; chief reviews", "MVP", "Duplicate ranking numbers are blocked or flagged."),
    rule("INT-003", "Resident Intake", "Template uniqueness", "Can two residents request the same top master template?", "Yes, but conflicts are flagged for chief review.", "Resident annual form", "Master builder conflict view", "IT + Chief", "Automated flag; chief decides", "MVP", "System highlights competing top choices and decision factors."),
    rule("INT-004", "Resident Intake", "Submission ordering", "Should first-come-first-serve affect priority?", "Yes as a signal, not an absolute decision.", "Timestamp", "Master builder sorting and priority notes", "Chief", "Chief override", "MVP", "Chief can sort by submission time and override with documented reason."),
    rule("INT-005", "Resident Intake", "Fellowship timing", "Do fellowship applicants receive priority for early electives?", "Often yes, especially Blocks 3-5 or interview season.", "Resident annual form", "Master builder and elective assignment", "Chief", "Chief decides", "MVP", "Fellowship flag appears beside rankings and requests."),
    rule("INT-006", "Resident Intake", "Life events", "How are major life events handled?", "Maternity/paternity, wedding, visa travel, board review, family need.", "Resident comments + chief review", "Requests center, master builder, PTO/LOA", "Chief", "Chief approves", "MVP", "Major life event requests are visible but controlled by privacy policy."),
    rule("INT-007", "Resident Intake", "Holiday requests", "Do residents rank holidays they prefer off?", "Christmas, New Year, Thanksgiving, Labor Day, July 4, MLK, Memorial Day.", "Resident annual form", "Holiday workbook", "Resident + Chief", "Resident submits; chief approves", "MVP+", "Holiday request ranks feed holiday staffing and conflict view."),
    rule("INT-008", "Resident Intake", "Elective ranking", "How many elective choices are ranked?", "Example: top 12 electives; PGY-3 may have advanced electives.", "Resident annual form", "Elective assignment and master builder", "Resident + Chief", "Resident submits; chief approves", "MVP+", "Elective choices have rank, eligibility, availability, and capacity."),
    rule("INT-009", "Resident Intake", "Two-week electives", "Can residents request split two-week electives?", "Yes if department allows; may affect vacation eligibility.", "Resident elective form", "Master/elective builder", "Chief", "Chief editable", "MVP+", "Split block requests are visible and do not silently break vacation rules."),
    rule("INT-010", "Resident Intake", "Comment parsing", "Should free-text comments be structured for chiefs?", "Yes: show comment plus tags like travel, fellowship, maternity, board review.", "Resident annual form", "Chief review panel", "IT + Chief", "Chief can retag", "MVP+", "Comments are searchable and linked to resident profile."),

    rule("MAS-001", "Master Schedule", "Core rotations", "Which rotations are mandatory for each PGY class?", "PGY-specific list: floor teams, ED, NICU, PICU, cardiology, etc.", "Chief curriculum rules", "Master builder, service coverage, profile", "Chief", "Chief editable", "MVP", "Core rotations can be added, renamed, dragged, deleted, and mapped to services."),
    rule("MAS-002", "Master Schedule", "Elective rotations", "Which rotations are elective/vacation/call-pool eligible?", "White boxes or elective/outpatient/easy rotations.", "Chief curriculum rules", "Call pool, vacation capacity, master builder", "Chief", "Chief editable", "MVP", "Elective rotations are marked for vacation/call eligibility separately."),
    rule("MAS-003", "Master Schedule", "PGY differences", "How do requirements differ by PGY?", "PGY-1 no PICU in some programs; PGY-2/3 ICU; PGY-2 night blocks; PGY-3 advanced electives.", "Chief curriculum rules", "Master builder and eligibility checks", "Chief", "Chief editable", "MVP", "Each PGY has its own core/elective/minimum rule set."),
    rule("MAS-004", "Master Schedule", "Number of templates", "How many master schedules must be built for each class?", "Match number of residents in that class.", "Chief setup + resident roster", "Master builder rankings", "Chief", "Chief editable", "MVP", "System flags mismatch between resident count and template count."),
    rule("MAS-005", "Master Schedule", "Service minimums", "For each block, what minimum and maximum residents are needed per rotation/service?", "Example: Purple 2-4 interns/block; Cardiology 2-3/block.", "Chief rules", "Block coverage, service coverage, alerts", "Chief", "Chief editable", "MVP", "Counts per block and service show OK/short/over."),
    rule("MAS-006", "Master Schedule", "Vacation capacity", "How many vacation/elective slots are allowed per block?", "Set min/max by PGY, service, and block.", "Chief rules", "Vacation requests, master builder", "Chief", "Chief editable", "MVP", "Vacation requests cannot silently exceed block capacity."),
    rule("MAS-007", "Master Schedule", "Elective distribution", "Should electives be distributed to reduce burnout?", "At least one elective/easier block early and later when possible.", "Chief rule", "Master builder suggestions", "Chief", "Chief editable", "MVP+", "System flags residents with long inpatient-heavy stretches."),
    rule("MAS-008", "Master Schedule", "Night-to-next-block safety", "Should residents ending a block on nights avoid inpatient service next block?", "Yes when possible to protect post-call transition.", "Chief rule", "Master builder conflict inspector", "Chief", "Chief override", "MVP", "Unsafe night-to-inpatient transition is flagged."),
    rule("MAS-009", "Master Schedule", "Manual edits", "Can chiefs override generated master assignments?", "Yes. Chief remains final decision maker.", "Chief action", "Master grid, audit trail, profiles", "Chief", "Chief editable", "MVP", "Manual changes update service coverage and resident profiles."),
    rule("MAS-010", "Master Schedule", "Finalize source", "When is master schedule considered official source data?", "After chief validates requirements, approves requests, and saves final master.", "Chief action", "Service coverage, monthly builders, resident portals", "Chief", "Chief publishes", "MVP", "Downstream builders use only finalized or explicitly selected draft source."),
    rule("MAS-011", "Master Schedule", "Excel-like view", "Do chiefs need an Excel-like view of the annual master?", "Yes as optional transition view.", "Design preference", "Master schedule UI", "IT", "User toggle", "MVP+", "Digital and Excel-like views show same source data."),
    rule("MAS-012", "Master Schedule", "Import existing master", "How should an existing Excel master be reviewed after upload?", "Map PGY tabs, names, blocks, rotations, notes; show unmatched items.", "Excel/PDF upload", "Profiles, service coverage, call pools", "IT + Chief", "Chief approves import", "MVP+", "No imported data becomes official until chief approves mapping."),

    rule("SCV-001", "Service Coverage", "Purpose", "Should there be a service coverage workbook separate from monthly schedules?", "Yes. It shows who covers each service each block before day-by-day schedule building.", "Final master schedule", "Monthly service builder, overview, analytics", "IT", "Automated with chief edits", "MVP", "Service coverage names are pulled from finalized master, not random."),
    rule("SCV-002", "Service Coverage", "Name drilldown", "Can users click service counts to see names?", "Yes. Counts stay clean; details expand on click.", "Service coverage data", "Resident profiles, service detail panel", "IT", "Automated", "MVP", "Clicking a service shows assigned residents and profile links."),
    rule("SCV-003", "Service Coverage", "Year at glance", "Should counts be visible across all blocks?", "Yes. Compact count grid with status color and optional Excel-like view.", "Final master schedule", "Chief overview and service coverage", "IT", "User toggle", "MVP", "Chief can see block-service staffing across the year without crowding."),
    rule("SCV-004", "Service Coverage", "Outside rotators", "How are outside rotators counted?", "Separate flag by institution, PGY equivalent, service eligibility, dates.", "Outside rotator profile", "Service coverage, institution rules, monthly builder", "Chief", "Chief editable", "MVP", "Outside rotator load is tracked by block and service."),
    rule("SCV-005", "Service Coverage", "Short/over alerts", "When should alerts appear?", "Below min = short; above max = over; outside rotator overload flagged.", "Service min/max rules", "Overview, master builder, service coverage", "IT", "Automated", "MVP", "Alerts show service, block, assigned count, target, and names."),

    rule("MON-001", "Monthly Scheduler", "Service selection", "When building a monthly schedule, how are residents pulled?", "Select block + service; system pulls residents assigned to that service in master schedule.", "Final master schedule + service coverage", "Monthly service schedule", "IT", "Automated with chief override", "MVP", "PICU Block 2 pulls only residents assigned to PICU Block 2 plus eligible supplements."),
    rule("MON-002", "Monthly Scheduler", "Service-specific rules", "Does each service have its own rules?", "Yes: required residents, PGY mix, locations/units, shifts, calls, nights, weekends.", "Institution/service rules", "Monthly generator and validation", "Chief", "Chief editable", "MVP", "Each service tab can be configured independently."),
    rule("MON-003", "Monthly Scheduler", "Coverage locations", "Can services have units/teams/locations?", "Yes: PICU Unit 1/2, floor colored teams, nursery, ICU zones.", "Service rule setup", "Schedule builder and analytics", "Chief", "Chief editable", "MVP", "Locations can be added, renamed, colored, and given minimums."),
    rule("MON-004", "Monthly Scheduler", "Role templates", "Can services define custom roles?", "Day, night, short call, long call, procedure call, clinic call, jeopardy, crossover.", "Service rule setup", "Shift cards, schedule builder, fairness stats", "Chief", "Chief editable", "MVP", "Roles have time, color, eligibility, count target, and fairness category."),
    rule("MON-005", "Monthly Scheduler", "Shift templates", "What timing fields are required for a shift?", "Start, end, duration, label, color, service, day/night/protected/off category.", "Service rules", "Monthly builder, hours analytics", "Chief", "Chief editable", "MVP", "Hours calculate from shift templates and manual overrides."),
    rule("MON-006", "Monthly Scheduler", "Protected time import", "Should clinics/didactics/exams auto-place before shifts?", "Yes. Pull from resident profile and institution rules.", "Profiles + institution rules", "Monthly builder conflict protection", "IT", "Automated; chief can edit", "MVP", "Protected time appears before assignments and blocks unsafe calls."),
    rule("MON-007", "Monthly Scheduler", "Fairness targets", "What fairness metrics are balanced?", "Weekend days, golden weekends, nights, calls, clinic calls, hours, days off.", "Schedule generator", "Insights and analytics", "IT + Chief", "Chief config", "MVP", "Chief sees imbalance before publishing."),
    rule("MON-008", "Monthly Scheduler", "Manual edits", "Can chiefs drag/drop and override generated schedules?", "Yes. Generated schedule is draft, not final.", "Chief action", "Monthly schedule, audit trail, resident view after publish", "Chief", "Chief editable", "MVP", "Manual edit recalculates coverage and conflict warnings."),
    rule("MON-009", "Monthly Scheduler", "Publishing", "When do residents see monthly schedules?", "Only after chiefs publish block/service schedule.", "Chief action", "Resident portal, published schedules", "Chief", "Chief publishes", "MVP", "Drafts are hidden from residents unless shared intentionally."),

    rule("NCL-001", "Night/Call/Clinic", "Night eligibility", "Which residents can do nights on each service?", "Service-specific PGY and institution eligibility.", "Service rules + profiles", "Monthly builder", "Chief", "Chief editable", "MVP", "Ineligible residents are blocked or warned."),
    rule("NCL-002", "Night/Call/Clinic", "Consecutive nights", "What is the maximum consecutive night stretch?", "Example: up to 5 nights, service-dependent.", "Service rules", "Monthly builder, fairness stats", "Chief", "Chief editable", "MVP", "Generator avoids exceeding max; chief override flagged."),
    rule("NCL-003", "Night/Call/Clinic", "Post-call", "What post-call/rest rules apply?", "Post-call day off after night/call when required.", "Service rules", "Monthly builder, resident schedule", "Chief", "Chief editable", "MVP", "Post-call rest appears automatically when rule requires."),
    rule("NCL-004", "Night/Call/Clinic", "Clinic same day", "Can residents be assigned call on clinic day?", "Usually no.", "Institution/clinic profile", "Call switch and monthly builder", "Chief", "Chief editable", "MVP", "Call on clinic day is flagged or blocked."),
    rule("NCL-005", "Night/Call/Clinic", "Clinic next day", "Can residents take call the day before clinic?", "Usually no because post-call clinic is unsafe.", "Institution/clinic profile", "Call switch and monthly builder", "Chief", "Chief editable", "MVP", "Call before clinic is flagged or blocked."),
    rule("NCL-006", "Night/Call/Clinic", "Back-to-back calls", "Can residents have back-to-back calls?", "No unless chief overrides.", "Call rules", "Call switch and monthly builder", "Chief", "Chief override", "MVP", "Back-to-back call request shows reason and warning."),
    rule("NCL-007", "Night/Call/Clinic", "Golden weekends", "How is golden weekend defined?", "Saturday and Sunday both off.", "Program rule", "Fairness stats and resident portal", "Chief", "Chief editable", "MVP", "System counts golden weekends per resident per block/year."),
    rule("NCL-008", "Night/Call/Clinic", "Weekend fairness", "What weekend targets apply?", "Each resident should have fair Saturday, Sunday, weekend worked, and golden weekend balance.", "Program/service rules", "Monthly builder analytics", "Chief", "Chief editable", "MVP", "Fairness panel compares residents on same service/block."),

    rule("CSW-001", "Call Switch", "Offer creation", "Can residents offer a call for switch?", "Yes, by selecting call date and optional alternate dates/notes.", "Resident portal", "Call switch marketplace, chief approval", "Resident", "Resident submits", "MVP+", "Offered call appears to eligible residents."),
    rule("CSW-002", "Call Switch", "Eligibility check", "What checks run before another resident accepts?", "No clinic same day, no clinic next day, no back-to-back calls, no duplicate assignment, PGY/service eligibility.", "Profiles + schedules", "Resident switch view and chief approval", "IT", "Automated", "MVP+", "Every switch shows eligibility reasons before acceptance."),
    rule("CSW-003", "Call Switch", "Receiver acceptance", "Can eligible residents accept offered calls?", "Yes, but final schedule change requires chief approval.", "Resident portal", "Chief request center", "Resident + Chief", "Resident accepts; chief approves", "MVP+", "Switch status moves offered -> accepted -> chief review -> approved/declined."),
    rule("CSW-004", "Call Switch", "Chief review", "What should chiefs see?", "Offering resident, receiving resident, date, call type, eligibility checks, clinic conflicts, back-to-back status.", "Call switch workflow", "Chief approvals and schedules", "Chief", "Chief final decision", "MVP+", "Chief can approve, decline, or request changes."),
    rule("CSW-005", "Call Switch", "After approval", "What updates after switch approval?", "Both resident schedules, service schedule, call counts, fairness stats, audit trail.", "Chief action", "Resident portal, monthly schedule, analytics", "IT", "Automated", "MVP+", "No manual PDF searching is needed after approval."),

    rule("REQ-001", "Requests", "Request types", "What request types can residents submit?", "Weekend off, day off, vacation/PTO, holiday off, clinic request, elective request, special need, call switch.", "Resident portal", "Resident profile, chief approvals, schedule builders", "Resident + Chief", "Resident submits; chief approves", "MVP", "Each request has type, date/block, reason, priority, status."),
    rule("REQ-002", "Requests", "Approval status", "What request states exist?", "Submitted, under review, approved, declined, needs info, withdrawn, applied to schedule.", "Workflow", "All request screens and builder inputs", "IT + Chief", "Chief action", "MVP", "Only approved/applied requests become generator constraints."),
    rule("REQ-003", "Requests", "Conflict priority", "How are competing requests handled?", "First come first serve, fellowship needs, life events, maternity, chief discretion.", "Program policy", "Approval center and master builder", "Chief", "Chief decides", "MVP", "System suggests priority but does not force final decision."),
    rule("REQ-004", "Requests", "Pending reminders", "Should pending requests appear during schedule building?", "Yes, as reminders; approved requests are constraints.", "Request center", "Master and monthly builders", "IT", "Automated", "MVP", "Chief can see whether request is approved/pending/declined while building."),

    rule("PTO-001", "PTO/Leave", "PTO tracking", "What PTO fields are tracked?", "Date/block, type, approved status, rotation, inpatient/elective, clinic impact, service coverage impact.", "Resident request + chief action", "Resident profile, monthly builder, analytics", "Chief", "Chief approves", "MVP+", "Resident and chief see running totals and affected services."),
    rule("PTO-002", "PTO/Leave", "Sick leave", "How are sick days recorded?", "Date, service, call impact, clinic impact, coverage needed, excused/unexcused if applicable.", "Resident/chief entry", "Profile, service schedule, analytics", "Chief + Resident", "Chief finalizes", "MVP+", "Sick leave updates profile and coverage alerts."),
    rule("PTO-003", "PTO/Leave", "LOA/maternity", "How are longer leaves handled?", "Date range, blocks affected, coverage impact, privacy level, replacement needs.", "Chief profile", "Master, service coverage, monthly builders", "Chief", "Chief restricted edit", "MVP+", "LOA appears as protected/unavailable without exposing private details unnecessarily."),
    rule("PTO-004", "PTO/Leave", "Exams/conferences", "How are one-time protected events entered?", "ITE, board exam, conference, retreat, simulation, travel buffer.", "Chief or resident request", "Monthly builder and profile", "Chief", "Chief editable", "MVP", "Events block unsafe assignments and show in resident schedule."),

    rule("DID-001", "Didactics/Attendance", "Didactic rules", "Which didactics require residents to leave service?", "Only didactics that affect coverage should be protected.", "Institution rules", "Monthly builder and attendance", "Chief", "Chief editable", "MVP", "Morning reports that do not affect coverage can be optional/untracked."),
    rule("DID-002", "Didactics/Attendance", "Institution patterns", "Do institutions have recurring didactic patterns?", "Example: Friday 12:30-17:00, virtual allowed for some rotations.", "Institution profile", "Resident profile, monthly builder, attendance", "Chief", "Chief editable", "MVP", "Outside rotator didactics auto-apply based on institution."),
    rule("DID-003", "Didactics/Attendance", "Attendance check-in", "How is attendance captured?", "Tablet-friendly roster with checkmarks for each lecture.", "Attendance tab", "Resident/chief profiles and analytics", "Chief/admin", "Chief creates session", "MVP+", "Checking a resident updates attendance count automatically."),
    rule("DID-004", "Didactics/Attendance", "Lecture setup", "What fields define a lecture?", "Block, date, time, lecture type, required/optional, in-person/virtual, eligible residents.", "Attendance tab", "Attendance analytics", "Chief/admin", "Chief editable", "MVP+", "Session can be reused or copied."),
    rule("DID-005", "Didactics/Attendance", "Missed lecture stats", "What attendance stats are shown?", "Attended/missed by week, block, month, year, rotation, institution, excused/unexcused.", "Attendance records", "Resident profile, chief analytics", "IT", "Automated", "MVP+", "Chief can see missed count and context for each resident."),

    rule("HOL-001", "Holidays", "Holiday list", "Which holidays require special staffing?", "July 4, Labor Day, Thanksgiving, Christmas Eve/Day, New Year Eve/Day, MLK, Memorial Day.", "Program setup", "Holiday workbook and resident requests", "Chief", "Chief editable", "MVP+", "Holiday list can be customized by institution."),
    rule("HOL-002", "Holidays", "Holiday request ranking", "Do residents rank preferred holidays off?", "Yes, with comments and first-come signal.", "Resident annual form", "Holiday workbook", "Resident + Chief", "Resident submits; chief approves", "MVP+", "Holiday requests appear beside staffing needs."),
    rule("HOL-003", "Holidays", "Holiday service coverage", "Which services need holiday coverage?", "PICU, NICU, floor senior, floor intern, gold senior, cardiology, heme/onc, ED, jeopardy.", "Holiday rules", "Holiday workbook", "Chief", "Chief editable", "MVP+", "Each holiday has service-specific minimum counts."),
    rule("HOL-004", "Holidays", "Holiday break", "Does program have one- or two-week holiday break schedule?", "If yes, define break windows and service coverage minimums.", "Program rules", "Holiday breaks workbook", "Chief", "Chief editable", "MVP+", "Holiday break schedule links to master and service coverage."),

    rule("INS-001", "Institution Rules", "Clinic patterns", "How are clinic days assigned?", "Resident-specific by block, or institution-level pattern inherited by residents.", "Institution profile + resident profile", "Monthly builder and call switch", "Chief/admin", "Chief editable", "MVP", "Clinic patterns protect time and block unsafe calls."),
    rule("INS-002", "Institution Rules", "Didactics patterns", "How are didactics assigned?", "Institution-wide, resident-specific, block-specific, virtual/in-person rules.", "Institution profile", "Monthly builder and attendance", "Chief/admin", "Chief editable", "MVP", "Protected didactics appear automatically in schedule builder."),
    rule("INS-003", "Institution Rules", "Outside rotator eligibility", "Which services/roles can outside rotators cover?", "By institution, PGY equivalent, rotation, block, and service.", "Institution profile", "Service coverage and monthly builder", "Chief", "Chief editable", "MVP", "Ineligible outside rotators are not auto-assigned."),
    rule("INS-004", "Institution Rules", "PICU/NICU off-service rules", "Do some institutions have special PICU/NICU didactic or off-service rules?", "Example: off-service didactics vary by institution.", "Institution profile", "PICU/NICU schedule builder", "Chief", "Chief editable", "MVP+", "Institution-specific exceptions appear when that resident is pulled."),

    rule("ANL-001", "Analytics/Alerts", "Coverage alerts", "Which coverage alerts are required?", "Service below minimum, over maximum, missing PGY mix, missing night coverage, missing weekend coverage.", "Rules + schedule data", "Overview, builders, service coverage", "IT", "Automated", "MVP", "Alerts are clickable and explain cause/action."),
    rule("ANL-002", "Analytics/Alerts", "Fairness analytics", "Which fairness summaries are required?", "Weekends worked, golden weekends, nights, calls, hours, clinic calls, jeopardy calls.", "Schedule data", "Resident profile, chief analytics", "IT", "Automated", "MVP", "Resident totals can be compared to peer group."),
    rule("ANL-003", "Analytics/Alerts", "Work hours", "How should work hours be calculated?", "From shift templates plus manual overrides; weekly/block totals.", "Shift templates and schedules", "Resident/chief analytics", "IT", "Automated", "MVP", "Hours recalculate when chiefs edit shifts."),
    rule("ANL-004", "Analytics/Alerts", "Readiness", "What readiness status should each schedule show?", "Not started, needs inputs, generated draft, needs review, ready, published.", "Workflow status", "Overview and schedules tab", "IT", "Automated with chief control", "MVP", "Overview shows each block and service status."),

    rule("DES-001", "Design/UX", "Portals", "Should chief and resident views be separate?", "Yes. Same data backbone, different permission and workflow views.", "Product design", "All UI", "IT", "IT config", "MVP", "Chief and resident portals have clear toggle/roles."),
    rule("DES-002", "Design/UX", "Progressive disclosure", "How do we avoid crowded screens?", "Show counts first, names/details on click, drawers/modals for deeper data.", "Product design", "Master, service coverage, call pools, holidays", "IT", "Design standard", "MVP", "Large Excel-like data can be explored without overwhelming default view."),
    rule("DES-003", "Design/UX", "Excel-like view", "Where should Excel-like views be offered?", "Master, service coverage, call pools, holidays, holiday breaks, monthly schedule final view.", "User preference", "View toggles", "IT", "User toggle", "MVP+", "Digital and Excel-like views reflect the same data."),
    rule("DES-004", "Design/UX", "Clickable names", "Should resident names/initials link to profiles?", "Yes, with a back button to return to prior context.", "Product design", "All roster/count views", "IT", "Design standard", "MVP", "Back navigation preserves where user came from."),
    rule("DES-005", "Design/UX", "Design customization", "Can departments customize colors/labels?", "Yes for service colors, role colors, shift card labels, tab names.", "Department preferences", "Schedule UI and exports", "Chief/admin", "Chief editable", "MVP+", "Design customization does not break analytics categories."),
    rule("DES-006", "Design/UX", "Import review", "How should uploaded Excel/PDF data be reviewed?", "Stepwise mapping, preview, unmatched names, confidence score, approve import.", "Upload workflow", "Master/service/call/holiday modules", "IT + Chief", "Chief approves", "MVP+", "No upload silently overwrites official data."),
]


BUSINESS_ROWS = [
    ["Business layer", "What the company offers", "Hospital-facing value", "IT build implication"],
    ["Discovery framework", "A structured interview and rule checklist for every program.", "Captures hidden scheduling logic that usually lives in chiefs' heads and Excel comments.", "Rules engine must be configurable, not hardcoded."],
    ["Data model", "Residents, institutions, rotations, blocks, services, shifts, requests, attendance, PTO, call switches.", "One source of truth replaces scattered forms, PDFs, and spreadsheets.", "Relational backend with audit trail and permissions."],
    ["Scheduling engine", "Builds master schedules and monthly service schedules from rules and approved inputs.", "Saves chiefs from rebuilding from scratch and exposes gaps early.", "Constraint solver/draft generator plus manual chief editing."],
    ["Workflow platform", "Approvals, call switches, PTO, attendance, publishing, notifications.", "Reduces manual PDF searching and three-person coordination for simple changes.", "State machines for requests and schedule publishing."],
    ["Transition design", "Digital views plus Excel-like views, import existing Excel/PDF, export schedules.", "Meets chiefs where they are instead of forcing immediate behavior change.", "Reusable view layer and upload/mapping tools."],
    ["Analytics layer", "Coverage, fairness, hours, attendance, call balance, vacation capacity.", "Leadership sees risk and workload distribution before issues become conflicts.", "Metrics generated from schedule facts, not manual summaries."],
]


LINKAGE_ROWS = [
    ["Source event", "Feeds into", "Modules affected", "Why linkage matters"],
    ["Resident submits annual ranking", "Chief master builder", "Resident profile, master schedule, service coverage", "Preferences become decision support for the annual source schedule."],
    ["Chief approves vacation/request", "Master/monthly builder constraints", "Resident profile, schedule builder, analytics", "Only approved requests should influence generated schedules."],
    ["Chief finalizes master schedule", "Official source of rotations by block", "Service coverage, call pools, resident portal, monthly builder", "Monthly schedules must pull real assigned residents, not random samples."],
    ["Institution clinic/didactic rule changes", "Protected time generator", "Resident profiles, monthly builder, call switch checks, attendance", "Outside rotator rules apply automatically without entering each resident manually."],
    ["Service rule changes", "Monthly schedule builder", "Coverage targets, shift templates, alerts, analytics", "Each service can have different rules, roles, and staffing needs."],
    ["Call switch accepted and approved", "Published schedule update", "Both residents' schedules, service tab, call counts, chief audit trail", "Prevents manual PDF review and missed clinic/back-to-back conflicts."],
    ["Attendance check-in", "Didactics record", "Resident profile, chief analytics, attendance dashboard", "Lecture attendance updates instantly for both resident and chief view."],
    ["PTO/sick leave entered", "Availability and coverage", "Resident profile, monthly builder, service coverage, analytics", "Chiefs see service impact and resident totals."],
]


def style_header(row, fill=PURPLE, font_color="FFFFFF"):
    for cell in row:
        cell.fill = PatternFill("solid", fgColor=fill)
        cell.font = Font(color=font_color, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def style_sheet(ws, freeze="A2"):
    ws.freeze_panes = freeze
    ws.sheet_view.showGridLines = False
    thin = Side(style="thin", color=LINE)
    for row in ws.iter_rows():
        for cell in row:
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if cell.row > 1:
                cell.font = Font(color=DARK, size=10)


def set_widths(ws, widths):
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width


def add_rules_sheet(wb, title, rows):
    ws = wb.create_sheet(title)
    headers = list(rows[0].keys())
    ws.append(headers)
    style_header(ws[1])
    for item in rows:
        ws.append([item[h] for h in headers])
    set_widths(ws, [13, 18, 22, 48, 34, 24, 42, 20, 18, 14, 40, 34, 34, 34])
    style_sheet(ws)
    ws.auto_filter.ref = ws.dimensions
    for row in ws.iter_rows(min_row=2):
        priority = row[9].value
        if priority == "MVP":
            for cell in row:
                cell.fill = PatternFill("solid", fgColor="FFFFFF")
        elif priority == "MVP+":
            for cell in row:
                cell.fill = PatternFill("solid", fgColor="F7FBFF")
    return ws


def create_workbook():
    OUT.mkdir(exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "00 Business Model"
    for row in BUSINESS_ROWS:
        ws.append(row)
    style_header(ws[1])
    set_widths(ws, [24, 42, 48, 42])
    style_sheet(ws)

    ws = wb.create_sheet("01 Linkage Map")
    for row in LINKAGE_ROWS:
        ws.append(row)
    style_header(ws[1], TEAL)
    set_widths(ws, [32, 32, 44, 54])
    style_sheet(ws)

    add_rules_sheet(wb, "02 Full Rule Register", RULES)

    domains = [
        ("03 Program Setup", "Program Setup"),
        ("04 Resident Intake", "Resident Intake"),
        ("05 Master Rules", "Master Schedule"),
        ("06 Service Coverage", "Service Coverage"),
        ("07 Monthly Scheduler", "Monthly Scheduler"),
        ("08 Nights Calls Clinic", "Night/Call/Clinic"),
        ("09 Call Switch Rules", "Call Switch"),
        ("10 Requests PTO Leave", ("Requests", "PTO/Leave")),
        ("11 Attendance Didactics", "Didactics/Attendance"),
        ("12 Holidays Breaks", "Holidays"),
        ("13 Institution Rules", "Institution Rules"),
        ("14 Analytics Design", ("Analytics/Alerts", "Design/UX")),
    ]
    for sheet_title, domain in domains:
        if isinstance(domain, tuple):
            rows = [r for r in RULES if r["Domain"] in domain]
        else:
            rows = [r for r in RULES if r["Domain"] == domain]
        add_rules_sheet(wb, sheet_title, rows)

    ws = wb.create_sheet("15 Discovery Interview")
    headers = ["Meeting section", "Who should answer", "Questions to ask", "Evidence/document to request", "Decision / answer", "Follow-up owner", "Status"]
    ws.append(headers)
    style_header(ws[1], PURPLE)
    interview_rows = [
        ["Program structure", "Program director/chiefs", "How many blocks, classes, tracks, institutions, outside rotators, and services?", "Current master schedule workbook and service coverage sheets", "", "", ""],
        ["Master schedule", "Chiefs", "What are core vs elective rotations by PGY? Which blocks are vacation/call eligible?", "PGY1/PGY2/PGY3/Med-Peds tabs", "", "", ""],
        ["Service coverage", "Chiefs/GME", "What minimum and maximum coverage is required per service and block?", "Services tab and minimum staffing documents", "", "", ""],
        ["Monthly scheduling", "Chiefs", "What shifts, roles, nights, weekends, calls, and fairness rules apply per service?", "Final monthly schedules and service instructions", "", "", ""],
        ["Clinic/didactics", "Chiefs/coordinators", "What protected time applies by resident, institution, and rotation?", "Clinic templates, didactic calendar, outside rotator rules", "", "", ""],
        ["Requests", "Chiefs/residents", "What request types exist and who approves them before schedule generation?", "Google Forms and request logs", "", "", ""],
        ["Call switches", "Chiefs/residents", "What makes a switch eligible or unsafe?", "Current switch policy and sample switch requests", "", "", ""],
        ["PTO/leave", "Chiefs/coordinators", "How are PTO, sick leave, LOA, conferences, and exams tracked?", "PTO tracker, sick leave logs, leave policy", "", "", ""],
        ["Attendance", "Chiefs/education admin", "Which lectures are tracked and what counts as required/optional/virtual?", "Didactic schedule and attendance roster", "", "", ""],
        ["Design transition", "Chiefs", "Which workflows need digital view, Excel-like view, import, export, or both?", "Screenshots of current Excel/PDF outputs", "", "", ""],
    ]
    for row in interview_rows:
        ws.append(row)
    set_widths(ws, [24, 24, 56, 44, 34, 22, 16])
    style_sheet(ws)

    ws = wb.create_sheet("16 IT Build Tasks")
    headers = ["Epic", "Build task", "Depends on", "Rules involved", "MVP phase", "Acceptance criteria", "Owner notes"]
    ws.append(headers)
    style_header(ws[1], TEAL)
    tasks = [
        ["Data foundation", "Create resident, institution, block, rotation, service, shift, request, attendance, and PTO data models.", "Program setup rules", "PGM, RES, INS", "Phase 1", "Can import/create fake and real resident profiles with linked schedule data.", ""],
        ["Master builder", "Build stepwise PGY master workflow with editable core rotations, resident rankings, block coverage, and finalize grid.", "Resident intake + PGY curriculum", "INT, MAS", "Phase 2", "Chief can build or import master and see counts/gaps.", ""],
        ["Service coverage", "Generate service coverage workbook from finalized master with clickable names and Excel-like toggle.", "Final master", "SCV", "Phase 2", "Selecting block/service shows assigned residents and profile links.", ""],
        ["Monthly scheduler", "Build block/service schedule generator pulling residents, protected time, shift templates, and requests.", "Service coverage + profiles", "MON, NCL, REQ", "Phase 3", "Chief can generate, edit, validate, and publish one service schedule.", ""],
        ["Resident portal", "Build personal schedule, private master, requests, call switches, PTO/sick leave, published schedules.", "Published schedules + profile", "RES, REQ, CSW, PTO", "Phase 3", "Resident sees only their own private data and shared published service rosters.", ""],
        ["Call switches", "Build offer/accept/chief approval workflow with eligibility checks.", "Published calls + profiles", "CSW, NCL, INS", "Phase 4", "Switch updates both residents after chief approval.", ""],
        ["Attendance", "Build tablet attendance session check-in and resident/chief stats.", "Resident roster + didactic rules", "DID", "Phase 5", "Checking a name updates profile and attendance dashboard.", ""],
        ["Import/export", "Build Excel/PDF upload mapping plus Excel-like exports.", "Data model and mapping rules", "PGM, MAS, SCV, DES", "Phase 4", "Uploaded master can be mapped, reviewed, approved, and applied.", ""],
    ]
    for row in tasks:
        ws.append(row)
    set_widths(ws, [22, 54, 30, 22, 16, 52, 30])
    style_sheet(ws)

    ws = wb.create_sheet("17 Hospital Decisions")
    headers = ["Decision area", "Decision needed from hospital", "Options", "Selected decision", "Who signs off", "Due date", "Risk if unanswered"]
    ws.append(headers)
    style_header(ws[1], PURPLE)
    decisions = [
        ["Block model", "Select number/date model for academic year.", "13 blocks / 12 monthly / 26 two-week / custom", "", "", "", "Cannot build calendar, requests, or schedules."],
        ["Core rotations", "Approve core rotations by PGY/track.", "Use current curriculum / customize / import", "", "", "", "Master builder cannot validate requirements."],
        ["Service minimums", "Approve min/max residents by service and block.", "Single range / block-specific / seasonal", "", "", "", "Coverage alerts will be inaccurate."],
        ["Clinic protection", "Approve call restrictions around clinic.", "No same day / no next day / AM-PM logic / custom", "", "", "", "Unsafe call switches may be allowed."],
        ["Night rules", "Approve max nights and post-call rules.", "5 nights / custom / service-specific", "", "", "", "Night schedule may violate program norms."],
        ["Request priority", "Approve priority policy.", "First come / fellowship / life event / chief discretion", "", "", "", "Conflicting requests cannot be resolved consistently."],
        ["Attendance", "Approve which lectures are tracked.", "All didactics / required only / selected sessions", "", "", "", "Attendance stats may not reflect education policy."],
        ["Design mode", "Choose default view style.", "Digital first / Excel-like first / both", "", "", "", "Adoption risk if workflow feels unfamiliar."],
    ]
    for row in decisions:
        ws.append(row)
    set_widths(ws, [24, 44, 36, 28, 22, 16, 42])
    style_sheet(ws)

    for sheet in wb.worksheets:
        sheet.row_dimensions[1].height = 34
    wb.save(XLSX_PATH)


def add_doc_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        hdr[i]._tc.get_or_add_tcPr().append(parse_shading(PURPLE))
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def parse_shading(hex_color):
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    return parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')


def create_docx():
    OUT.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for name, size, color in [("Heading 1", 16, PURPLE), ("Heading 2", 12, DARK)]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Clarity Scheduling Business Model and Configurable Rules")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(PURPLE)
    sub = doc.add_paragraph(
        "A practical operating model for sales, business analysts, IT builders, chiefs, coordinators, and hospital leadership."
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("What This Document Is", level=1)
    doc.add_paragraph(
        "This is the structure your company can use when meeting any residency program or hospital department. "
        "It separates what the product builds from the rules that must be collected, reviewed, edited, approved, and linked across the platform."
    )
    add_doc_table(
        doc,
        BUSINESS_ROWS[0],
        BUSINESS_ROWS[1:],
        [1.4, 2.6, 3.2, 2.5],
    )

    doc.add_heading("The Core Business Model", level=1)
    rows = [
        ["1. Discover", "Sales/business analyst interviews chiefs and coordinators using the rule workbook.", "Completed hospital decisions and rule register."],
        ["2. Configure", "IT configures blocks, PGY classes, rotations, services, shifts, institutions, and imports.", "Department-specific environment ready for testing."],
        ["3. Build master", "Chiefs build or import annual master schedule using resident submissions and service minimums.", "Finalized source schedule for each resident."],
        ["4. Build monthly schedules", "Chiefs generate and edit service schedules from master, profiles, protected time, and approved requests.", "Published schedules visible to residents."],
        ["5. Operate", "Residents request, switch calls, submit PTO, view schedules; chiefs approve and track analytics.", "Live scheduling operations and reporting."],
    ]
    add_doc_table(doc, ["Company stage", "What happens", "Output"], rows, [1.4, 4.5, 3.4])

    doc.add_heading("Critical Linkage Model", level=1)
    doc.add_paragraph(
        "These are the places where the product must behave as one connected system. These linkages are the main difference between a static design mockup and a real scheduling MVP."
    )
    add_doc_table(doc, LINKAGE_ROWS[0], LINKAGE_ROWS[1:], [2.1, 2.1, 2.8, 3.0])

    doc.add_heading("Rule Categories the Hospital Must Review and Edit", level=1)
    category_summary = []
    for domain in sorted(set(r["Domain"] for r in RULES)):
        subset = [r for r in RULES if r["Domain"] == domain]
        category_summary.append([domain, len(subset), ", ".join(sorted(set(r["MVP priority"] for r in subset))), subset[0]["Discovery question / rule to review"]])
    add_doc_table(doc, ["Rule category", "# rules", "Priority mix", "Example question"], category_summary, [1.8, 0.7, 1.2, 6.0])

    doc.add_heading("High-Risk Rules That Must Be Explicitly Confirmed", level=1)
    high_risk_ids = ["MAS-005", "MAS-008", "MON-001", "MON-002", "NCL-004", "NCL-005", "NCL-006", "CSW-002", "REQ-003", "DID-002", "PTO-003", "DES-003"]
    rows = [[r["Rule ID"], r["Domain"], r["Discovery question / rule to review"], r["Acceptance check"]] for r in RULES if r["Rule ID"] in high_risk_ids]
    add_doc_table(doc, ["Rule ID", "Domain", "Rule to confirm", "Acceptance check"], rows, [0.8, 1.5, 4.2, 3.1])

    doc.add_heading("How to Use the Editable Workbook", level=1)
    workbook_steps = [
        ["Start with 00 Business Model", "Use this to explain the product as a company offering, not just a screen."],
        ["Then 01 Linkage Map", "Show the hospital how resident submissions, chief approvals, master schedule, service coverage, and monthly schedules connect."],
        ["Use 02 Full Rule Register", "This is the master checklist. Every row has a decision field, IT notes field, and implementation notes field."],
        ["Use focused tabs 03-14", "Run department interviews by category so the conversation does not become overwhelming."],
        ["Use 15 Discovery Interview", "This is the sales/business analyst meeting script."],
        ["Use 16 IT Build Tasks", "Translate the rules into implementation epics."],
        ["Use 17 Hospital Decisions", "Capture items requiring leadership sign-off before build or launch."],
    ]
    add_doc_table(doc, ["Workbook area", "How the team should use it"], workbook_steps, [2.2, 7.4])

    doc.add_heading("Non-Negotiable Product Principle", level=1)
    doc.add_paragraph(
        "The master schedule is the source of truth for who belongs to each block and service. "
        "The monthly builder, service coverage workbook, call pool, holiday coverage, resident portal, chief analytics, and request workflows must pull from that source unless chiefs explicitly approve an override."
    )

    doc.add_heading("Deliverables Created", level=1)
    doc.add_paragraph(f"Editable workbook: {XLSX_PATH.name}")
    doc.add_paragraph(f"Business rules playbook: {DOCX_PATH.name}")
    doc.save(DOCX_PATH)


def hx(value):
    return colors.HexColor(value if value.startswith("#") else f"#{value}")


def p(text, style):
    safe = str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(safe, style)


def create_pdf():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("TitleX", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=21, leading=25, textColor=hx(PURPLE)))
    styles.add(ParagraphStyle("SubX", parent=styles["Normal"], fontName="Helvetica", fontSize=10, leading=14, textColor=hx(MUTED), spaceAfter=10))
    styles.add(ParagraphStyle("H1X", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=hx(PURPLE), spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle("BodyX", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.4, leading=9.1, textColor=hx(DARK)))
    styles.add(ParagraphStyle("HeadCell", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=6.2, leading=7.4, textColor=colors.white, alignment=1))
    styles.add(ParagraphStyle("Cell", parent=styles["BodyText"], fontName="Helvetica", fontSize=5.8, leading=7.0, textColor=hx(DARK)))

    def table(headers, rows, widths, fill=PURPLE):
        data = [[p(h, styles["HeadCell"]) for h in headers]]
        for row in rows:
            data.append([p(v, styles["Cell"]) for v in row])
        t = Table(data, colWidths=[w * inch for w in widths], repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), hx(fill)),
            ("GRID", (0, 0), (-1, -1), 0.25, hx(LINE)),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, hx("FAFBFE")]),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return t

    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=letter, leftMargin=0.45*inch, rightMargin=0.45*inch, topMargin=0.45*inch, bottomMargin=0.45*inch)
    story = [
        Paragraph("Clarity Scheduling Business Model and Configurable Rules", styles["TitleX"]),
        Paragraph("Operational rule structure for sales discovery, hospital review, and IT implementation.", styles["SubX"]),
        Paragraph("Business Model", styles["H1X"]),
        table(BUSINESS_ROWS[0], BUSINESS_ROWS[1:], [1.15, 1.85, 2.15, 1.8]),
        PageBreak(),
        Paragraph("Critical Linkage Model", styles["H1X"]),
        table(LINKAGE_ROWS[0], LINKAGE_ROWS[1:], [1.35, 1.35, 1.8, 2.45], TEAL),
        PageBreak(),
        Paragraph("Hospital Rule Categories", styles["H1X"]),
    ]
    category_summary = []
    for domain in sorted(set(r["Domain"] for r in RULES)):
        subset = [r for r in RULES if r["Domain"] == domain]
        category_summary.append([domain, len(subset), ", ".join(sorted(set(r["MVP priority"] for r in subset))), subset[0]["Discovery question / rule to review"]])
    story.append(table(["Rule category", "#", "Priority", "Example question"], category_summary, [1.55, 0.35, 0.8, 4.25]))
    story.append(PageBreak())
    for domain in [
        "Program Setup", "Resident Intake", "Master Schedule", "Service Coverage",
        "Monthly Scheduler", "Night/Call/Clinic", "Call Switch", "Requests",
        "PTO/Leave", "Didactics/Attendance", "Holidays", "Institution Rules",
        "Analytics/Alerts", "Design/UX"
    ]:
        story.append(Paragraph(domain, styles["H1X"]))
        rows = [[r["Rule ID"], r["Discovery question / rule to review"], r["Must update / link to"], r["Chief editable?"], r["Acceptance check"]] for r in RULES if r["Domain"] == domain]
        story.append(table(["ID", "Rule / question", "Linked modules", "Editable", "Acceptance"], rows, [0.55, 2.25, 1.75, 0.9, 1.5], PURPLE))
        story.append(PageBreak())
    story.append(Paragraph("Use the editable workbook for live hospital discovery", styles["H1X"]))
    story.append(Paragraph("The PDF is a readable playbook. The Excel workbook is the actual review/edit tool with fields for hospital answer, IT notes, and implementation status.", styles["BodyX"]))
    doc.build(story)


if __name__ == "__main__":
    create_workbook()
    create_docx()
    create_pdf()
    with ZipFile(DOCX_PATH) as z:
        assert "word/document.xml" in z.namelist()
    print(XLSX_PATH)
    print(DOCX_PATH)
    print(PDF_PATH)
