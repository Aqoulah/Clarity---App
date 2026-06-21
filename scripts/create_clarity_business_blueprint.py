from __future__ import annotations

from datetime import date
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
ASSETS = OUT / "brochure-v2-assets-anon"
DOCX_PATH = OUT / "Clarity_Scheduling_Business_Blueprint_and_Discovery_Guide.docx"
PDF_PATH = OUT / "Clarity_Scheduling_Business_Blueprint_and_Discovery_Guide.pdf"

PURPLE = "4F4698"
DARK = "23263A"
MUTED = "70788E"
LINE = "D9DDE8"
TEAL = "299A8F"
SOFT = "F7F8FC"
WARN = "B7791F"
RED = "D9534F"


def rgb(value: str) -> RGBColor:
    value = value.replace("#", "")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, color=DARK, size=8.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def configure_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.58)
    sec.right_margin = Inches(0.58)
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"].font.color.rgb = rgb(DARK)
    for style_name, size, color in [
        ("Title", 23, DARK),
        ("Subtitle", 11, MUTED),
        ("Heading 1", 15, PURPLE),
        ("Heading 2", 11.5, DARK),
        ("Heading 3", 10, PURPLE),
    ]:
        s = styles[style_name]
        s.font.name = "Aptos Display" if "Heading" in style_name or style_name == "Title" else "Aptos"
        s.font.size = Pt(size)
        s.font.color.rgb = rgb(color)
        s.font.bold = style_name != "Subtitle"
    return doc


def add_table(doc: Document, headers, rows, widths=None, font_size=7.4, header_fill=PURPLE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True, "FFFFFF", font_size)
        shade(hdr[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, DARK, font_size)
            if len(table.rows) % 2 == 0:
                shade(cells[i], "FAFBFE")
    if widths:
        for idx, width in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = Inches(width)
    doc.add_paragraph()
    return table


def callout(doc: Document, title: str, body: str, kind="info"):
    colors_by_kind = {
        "info": ("EDEAFB", PURPLE),
        "success": ("EAF8F6", TEAL),
        "warn": ("FFF4DD", WARN),
        "danger": ("FDECEC", RED),
    }
    fill, border = colors_by_kind[kind]
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    for edge in ["top", "left", "bottom", "right"]:
        border_el = OxmlElement(f"w:{edge}")
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), "8")
        border_el.set(qn("w:color"), border)
        tc_pr.append(border_el)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(9.2)
    r.font.color.rgb = rgb(border)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8.2)
        run.font.color.rgb = rgb(DARK)
    doc.add_paragraph()


def add_image(doc: Document, filename: str, caption: str, width=6.8):
    path = ASSETS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(7.4)
        run.font.color.rgb = rgb(MUTED)


def bullet(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1.5)
        p.add_run(item)


def numbered(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


CORE_WORKFLOW = [
    ["1", "Annual program setup", "Define academic year, block count, dates, PGY classes, tracks, institutions, holidays, rotations, services, roles, minimums."],
    ["2", "Resident intake", "Residents submit fixed facts and annual preferences: master ranking, vacation, holidays, electives, fellowship timing, life events."],
    ["3", "Chief review and approval", "Chiefs review requests, resolve conflicts, approve or decline requests, and mark what can be used by the generator."],
    ["4", "Master schedule creation", "System proposes assignments by PGY and block using rankings plus service minimums; chiefs edit and finalize."],
    ["5", "Service coverage workbook", "Final master becomes service coverage lists by block and service, including outside rotators and supplements."],
    ["6", "Call pool and holidays", "Elective/call-eligible residents flow into call pools; approved holiday requests guide holiday and break coverage."],
    ["7", "Monthly service schedule builder", "For each block and service, the system pulls residents from master/service coverage, protects clinics/didactics/PTO, and drafts editable schedules."],
    ["8", "Publish and operate", "Residents see personal schedules; chiefs monitor stats, switches, PTO, attendance, alerts, and changes."],
]


MODULES = [
    ["Resident portal", "Resident-facing input and visibility", "Annual preferences, monthly requests, call switches, PTO, attendance, personal schedule", "Feeds resident profile, chief approvals, master builder, monthly scheduler"],
    ["Chief portal", "Command center for building and approving", "All resident profiles, rules, requests, rosters, schedules", "Feeds every builder and published output"],
    ["Resident profile", "Single source of resident-specific facts", "PGY, institution, track, clinic pattern, didactics, PTO, master assignments, calls, attendance", "Used by master, service builder, call switch checks, stats"],
    ["Institution rules", "Reusable patterns for home and outside institutions", "Clinic/didactic schedules, attendance requirements, virtual rules, eligibility", "Auto-protects time when residents are scheduled"],
    ["Master schedule builder", "Annual source of truth", "Resident rankings, vacation/elective requests, core rules, service minimums", "Outputs rotation assignment per resident per block"],
    ["Service coverage workbook", "Bridge from master to monthly schedules", "Finalized master, outside rotators, service mapping", "Outputs who covers each service each block"],
    ["Schedule builder", "Monthly schedule generator/editor", "Service roster, shift templates, protected time, approved requests, coverage rules", "Outputs editable and publishable service schedules"],
    ["Call pool workbook", "Call-eligible resident pools", "Elective/call-eligible master blocks, no-call rules, LOA/PTO", "Outputs clinic call, jeopardy, cross-cover, no-call lists"],
    ["Holiday workbook", "Holiday coverage planning", "Holiday calendar, resident holiday rankings, service minimums", "Outputs holiday staffing and fairness stats"],
    ["Attendance", "Didactic check-in and tracking", "Lecture sessions, requirements, checkmarks, excused absences", "Updates chief and resident attendance stats"],
    ["Requests and approvals", "Workflow engine for decisions", "Vacation, day off, call switch, PTO, sick leave, special needs", "Approved items become scheduling constraints"],
    ["Analytics and alerts", "Decision support", "Assignments, hours, calls, clinics, weekends, attendance, conflicts", "Shows fairness, gaps, risks, and action items"],
]


LINKAGES = [
    ["Annual master ranking", "Resident portal", "Chief master builder", "Preferences shown by submission order, template rank, fellowship/life-event flags; chiefs can approve or override."],
    ["Vacation and holiday requests", "Resident portal", "Master builder, holiday workbook, resident profile", "Only approved requests become protected constraints; pending requests remain warnings."],
    ["Fellowship/maternity/life events", "Resident portal/profile", "Master builder and chief review", "Used as decision support, not automatic final authority."],
    ["Clinic pattern", "Institution profile or resident profile", "Monthly schedule builder, call switch checker", "Protects clinic time and prevents calls on clinic day or before clinic when rule applies."],
    ["Didactics rules", "Institution profile or resident profile", "Monthly builder and attendance", "Defines required attendance, virtual/in-person mode, and protected time."],
    ["Final master assignment", "Chief master builder", "Service coverage, schedule builder, resident view", "Primary source for who belongs on each service in each block."],
    ["Service coverage roster", "Service coverage workbook", "Monthly schedule builder", "Pulls names into service tabs like PICU, NICU, Purple, Orange, Gold, etc."],
    ["Approved PTO/sick/leave", "Resident/Chief request workflow", "Resident profile, schedule builder, analytics", "Blocks time, updates leave history, identifies service impact."],
    ["Call switch offer", "Resident portal", "Other residents, chief approval, schedule builder", "Eligibility checked automatically; final switch requires chief approval."],
    ["Attendance checkmark", "Attendance tab", "Resident profile and chief analytics", "Updates attended/missed counts by week, block, year, and requirement type."],
]


DISCOVERY = [
    ["Program structure", "How many blocks per year? What are block start/end dates? Are there 13 four-week blocks, 26 two-week blocks, monthly blocks, or mixed blocks? Do holiday breaks split blocks?"],
    ["Training classes", "Which PGY classes exist? PGY-1, PGY-2, PGY-3, fellows, Med-Peds, off-cycle residents, outside rotators? How many residents per class?"],
    ["Tracks", "Do tracks change required rotations? Examples: primary care, critical care, hospitalist, subspecialty, research, chief, Med-Peds."],
    ["Core rotations", "Which rotations are mandatory per PGY? How many times must each resident complete each? Which PGY levels are eligible?"],
    ["Electives", "Which rotations are elective? Which are vacation-eligible or call-eligible? Which are advanced electives? Which PGY levels can request them?"],
    ["Master templates", "Does the program use lettered templates A-T or individualized assignment? How many templates are available? Do residents rank templates?"],
    ["Capacity rules", "For each block, what is the minimum and maximum number of residents needed per service? Are these fixed or census-dependent?"],
    ["Requests priority", "What is the priority logic: first come first serve, fellowship timing, maternity/major life event, seniority, historical fairness, chief discretion?"],
    ["Protected time", "What clinic, didactic, conference, ITE, retreat, exam, and administrative days must be protected? Are they institution-specific?"],
    ["Night rules", "Who can do nights? Max consecutive nights? Post-call rules? No inpatient after night block? Weekend night rules? Split night blocks?"],
    ["Call rules", "What types exist: clinic call, J1/J2/J3 jeopardy, crossover, procedure call, long call, short call, 24-hour call? Who is eligible?"],
    ["Weekend fairness", "How many Saturdays, Sundays, worked weekends, and golden weekends should each resident receive per block and year?"],
    ["Holidays", "Which holidays matter? How do residents rank holiday requests? What minimum staffing is required per holiday and service?"],
    ["Attendance", "Which conferences require check-in? Are sessions required for all residents or only certain rotations? Is virtual attendance allowed?"],
    ["PTO and leave", "How are PTO, sick, maternity, conference, interview, and LOA tracked? What approval steps are needed? What reports does leadership need?"],
    ["Publishing", "What is public to all residents? What is private? Can residents see only their master schedule but all published service schedules?"],
    ["Audit and compliance", "What changes require notes? Who can override warnings? What reports are needed for GME, leadership, or ACGME-style monitoring?"],
]


RULES_MASTER = [
    ["PGY-specific curriculum", "Core rotations differ by year. Example concept: PGY-1 has many core floor/outpatient blocks; PGY-2 may add ICU/night float; PGY-3 may have advanced electives and leadership/board review."],
    ["Core vs elective classification", "Every rotation must be tagged as core, elective, vacation-eligible, call-eligible, no-call, ICU, floor, ED, outpatient, or custom category."],
    ["Template ranking", "Residents rank master schedule templates. The system must enforce unique ranks and show chiefs conflicts when multiple residents prefer the same template."],
    ["Preference priority", "Resident preferences guide the draft, but chiefs need final control. Suggested ranking signals: submission timestamp, fellowship application, major life event, seniority, and fairness."],
    ["Service minimums", "Each block must meet service minimums such as Purple/Orange, NICU, PICU, ED, Cardiology, Newborn, Heme/Onc, floor senior, gold senior, etc."],
    ["Vacation capacity", "Vacation-eligible blocks must have maximum allowed vacation slots; the system should flag blocks with too many vacation requests."],
    ["Unsafe transitions", "If a resident ends a block on nights, ICU nights, NICU nights, floor nights, or night float, the next block should avoid inpatient start if the program requires post-call recovery."],
    ["Outside rotator planning", "If a service is below minimum, the system should flag need for outside rotators or master schedule adjustment."],
    ["Editable final grid", "Chiefs must be able to drag/drop, change rotations, lock cells, validate, and save the finalized master as the source of truth."],
]


RULES_MONTHLY = [
    ["Roster pull", "When chiefs choose Block X and Service Y, residents assigned to that service in the finalized master should auto-load into the schedule builder."],
    ["Supplement pull", "Call pool or outside rotator residents can be added as supplements with visible source labels."],
    ["Protected time pull", "Clinics, didactics, exams, retreats, PTO, sick leave, maternity leave, conferences, and approved special requests should appear before shifts are assigned."],
    ["Service-specific rules", "Each service has its own staffing minimums, roles, units, shift types, and PGY eligibility. PICU may have Unit 1/Unit 2; NICU may have procedure call; floors may have colored teams."],
    ["Shift templates", "Day, night, long call, short call, procedure call, post-call, protected, off, crossover, and custom shifts should be configurable by service."],
    ["Fairness engine", "Track nights, weekends, Saturdays, Sundays, golden weekends, calls, clinic calls, hours, and days off across the block and year."],
    ["Chief editability", "The system drafts; chiefs edit. Drag/drop, copy, custom time, resident swap, row reorder, undo/redo, and notes are required."],
    ["Publish behavior", "Published service schedules become visible to residents; personal schedules combine service assignments, calls, clinics, didactics, PTO, and attendance obligations."],
]


RULES_CALL_SWITCH = [
    ["Offer", "Resident selects an assigned call and offers it for switch, optionally listing acceptable alternative dates."],
    ["Marketplace", "Other residents see available calls with eligibility indicators before accepting."],
    ["Automatic checks", "No clinic on call day, no clinic next day if post-call rule applies, no back-to-back calls, no duplicate assignment, PGY/service eligibility, duty hour risk."],
    ["Acceptance", "Receiving resident requests to take or swap. The system shows why they are eligible or not eligible."],
    ["Chief approval", "Chiefs see both residents, call details, clinic/didactic conflicts, back-to-back status, and final recommendation; chiefs approve or decline."],
    ["Propagation", "Approved switch updates the service schedule, resident personal schedules, call totals, profile histories, and audit log."],
]


DESIGN_REQUIREMENTS = [
    ["Dual experience", "Separate Chief portal and Resident portal with different permissions and visibility."],
    ["Stepwise builders", "Master builder and monthly service builder should walk users through setup, people, protected time, requests, rules, generate, edit, publish."],
    ["Digital and Excel-like views", "Every major schedule should support modern card/table view plus Excel-like grid view for transition comfort."],
    ["Upload-first option", "Programs with existing Excel/PDF masters can upload files; the system parses names, PGY, blocks, rotations, service coverage, call pools, holidays, and maps them to profiles."],
    ["Clickable details", "Counts should be clickable. Clicking a service, block, holiday, call pool, or resident should show names, sources, conflicts, and links to profiles."],
    ["Not crowded", "Use progressive disclosure: show counts first, then open details on click. Keep side panels for insights instead of forcing all data into one grid."],
    ["Visual alerts", "Use clear badges: Ready, Review, Gap, Over, Pending, Approved, Declined, Private, Published."],
    ["Autosave and audit", "Draft autosave, save draft, publish, export, undo/redo, and change notes should be obvious."],
    ["Accessibility", "Colors should help but not be the only signal. Use labels and status text in addition to color."],
]


def create_docx():
    OUT.mkdir(exist_ok=True)
    doc = configure_doc()
    doc.add_paragraph("Clarity Scheduling System", style="Title")
    doc.add_paragraph("Business blueprint, discovery guide, and implementation rulebook", style="Subtitle")
    doc.add_paragraph(f"Prepared: {date.today().strftime('%B %d, %Y')}")
    callout(
        doc,
        "Purpose of this document",
        "This is not a simple feature list. It is a company-facing guide for sales, business analysts, IT, chiefs, coordinators, and hospital leadership. It explains what schedules the platform builds, what rules must be collected, what data links across portals, and how to run a discovery conversation with a clinical department.",
        "success",
    )
    add_table(
        doc,
        ["Audience", "How they should use this document"],
        [
            ["Sales / discovery team", "Use the checklists to interview a department even without medical scheduling background."],
            ["IT / product team", "Use the module map, data linkages, and rule tables to design the MVP and data model."],
            ["Chief residents", "Use the rule categories to explain current Excel workflows and required customizations."],
            ["Program leadership", "Use the scope sections to decide what is core, what is add-on, and what reports matter."],
        ],
        [1.7, 4.9],
        8,
    )

    doc.add_heading("1. Product Thesis", level=1)
    doc.add_paragraph(
        "The core product is a clinical scheduling operating system. Its main job is to build and maintain the annual master schedule and monthly service schedules. Everything else, including attendance, call switches, PTO tracking, request approvals, analytics, and communication, exists because those items feed into or are affected by the schedules."
    )
    callout(
        doc,
        "Core idea to repeat in every hospital meeting",
        "We are not replacing Excel with a prettier spreadsheet. We are turning hidden scheduling logic into structured rules, linked resident profiles, editable draft schedules, and visible decision support.",
        "info",
    )
    add_table(doc, ["Layer", "What it means"], [
        ["Core scheduling engine", "Master schedules, service coverage, call pools, holiday coverage, monthly service schedules."],
        ["Linked resident profile", "The source of truth for each resident's PGY, institution, clinics, didactics, requests, leave, calls, and attendance."],
        ["Workflow tools", "Request approval, call switches, PTO/sick leave, attendance, publishing, audit trail."],
        ["Decision support", "Coverage gaps, over-capacity, fairness, hours, weekends, calls, missed didactics, and conflict warnings."],
    ], [1.65, 4.85], 8)

    doc.add_heading("2. End-to-End Scheduling Workflow", level=1)
    add_table(doc, ["Step", "Workflow", "What must be linked"], CORE_WORKFLOW, [0.45, 1.55, 4.5], 7.3)

    doc.add_heading("3. Core Modules and Ownership", level=1)
    add_table(doc, ["Module", "Purpose", "Primary data", "Feeds / depends on"], MODULES, [1.28, 1.45, 2.05, 1.75], 6.3)

    doc.add_heading("4. Critical Linkage Map", level=1)
    doc.add_paragraph(
        "These are the tasks that cannot be built as isolated screens. They must share data. This is the section the IT team should use to avoid creating disconnected prototypes."
    )
    add_table(doc, ["Data / action", "Starts in", "Must update", "Business rule"], LINKAGES, [1.45, 1.35, 1.65, 2.1], 6.5, TEAL)

    doc.add_heading("5. What Is Core vs Additional", level=1)
    add_table(doc, ["Category", "Features", "Why it matters"], [
        ["Core MVP", "Resident profiles, annual intake, master builder, service coverage, monthly service builder, publish resident view", "Without these, the product does not solve the main scheduling problem."],
        ["High-value add-ons", "Call switch workflow, PTO/sick leave tracking, requests approval, attendance, analytics", "These reduce daily chief/resident burden and make the system sticky."],
        ["Advanced add-ons", "Optimization engine, import parser, integrations, messaging, audit exports, duty-hour compliance reports", "These increase enterprise value after the MVP proves the workflow."],
    ], [1.35, 3.1, 2.05], 7.4)

    doc.add_heading("6. Discovery Checklist for Any Department", level=1)
    doc.add_paragraph("Use this list in the first discovery meeting. The goal is to understand how the department thinks, not only what their spreadsheet looks like.")
    add_table(doc, ["Category", "Questions to ask"], DISCOVERY, [1.25, 5.25], 6.7, PURPLE)

    doc.add_heading("7. Master Schedule Rules Checklist", level=1)
    doc.add_paragraph("The annual master is the source of truth. The product must support building it from resident requests or importing an already completed Excel/PDF master.")
    add_table(doc, ["Rule area", "Detailed rule / question"], RULES_MASTER, [1.35, 5.15], 7.0)
    callout(
        doc,
        "Important master schedule principle",
        "Resident rankings are decision support, not automatic truth. The system should suggest based on rules and priorities, but chiefs must retain final approval and override power.",
        "warn",
    )

    doc.add_heading("8. Annual Resident Intake Forms", level=1)
    add_table(doc, ["Form section", "Fields to collect", "Where it links"], [
        ["Fixed resident facts", "Name, email, PGY, institution, track, fellowship interest, outside rotator status", "Resident profile, eligibility checks, institution rules"],
        ["Master schedule ranking", "Rank templates or preferences; enforce unique ranks; timestamp submission", "Chief master builder, ranking review, conflict detection"],
        ["Vacation requests", "Preferred blocks/dates, priority rank, reason/comment, flexibility", "Master builder, PTO workflow, resident profile"],
        ["Holiday requests", "Rank holidays to be off or willing to work; comments and priority", "Holiday workbook, fairness analytics"],
        ["Elective requests", "Top 12 electives, advanced electives for eligible PGY levels, split-block requests", "Master builder, elective assignment, resident profile"],
        ["Life events and special needs", "Maternity/paternity, wedding, visa/travel, interview season, board review, conference", "Chief review, decision support, protected time"],
    ], [1.45, 2.85, 2.2], 6.8)

    doc.add_heading("9. Master Builder Design Requirements", level=1)
    add_table(doc, ["Step", "Screen goal", "Required functions"], [
        ["Choose PGY", "Select PGY class and define curriculum", "Number of blocks, number of residents/templates, core rotations, min/max per block, elective/vacation-eligible blocks, drag/add/delete rotations."],
        ["Resident rankings", "Review resident submitted preferences", "Sort by submission time, alphabetical, template rank, priority flags; open resident profile from row."],
        ["Request review", "See special requests without cluttering the builder", "Show only relevant flags; link to full request approval center/profile."],
        ["Block coverage", "See service count by block before final grid", "Clickable block and service counts, names, target min/max, shortage/overage alerts."],
        ["Finalize master", "Editable annual grid", "Drag/drop cells, lock selected, validate, show conflicts, save final source, export Excel-like view."],
        ["Import existing master", "Skip builder when master already exists", "Upload Excel/PDF, parse tabs, map residents/rotations/blocks, review unmatched names, approve import."],
    ], [1.25, 1.85, 3.4], 6.8, TEAL)
    add_image(doc, "chief-master.png", "Example: editable master schedule concept with conflict inspector and capacity counts.", 6.6)

    doc.add_heading("10. Service Coverage Workbook", level=1)
    doc.add_paragraph(
        "After the master is finalized, the service coverage workbook translates rotations into rosters. Chiefs are used to Excel tabs that list names under every service and block. The platform should show counts first and names on click so it stays readable."
    )
    add_table(doc, ["Required behavior", "Details"], [
        ["Auto-pull from master", "A resident assigned to PICU in Block 2 should appear in PICU service coverage for Block 2. A resident assigned Purple/Floor should appear in the correct floor service group if mapping rules say so."],
        ["Allow supplements", "Outside rotators, call pool residents, and manual supplements can be added with labels showing source."],
        ["Names without crowding", "Default view shows counts and status; clicking a service opens names, PGY, institution, source, and profile links."],
        ["Excel-like option", "Provide a grid view that resembles the old Services workbook for transition comfort."],
        ["Statistics", "Per block: assigned count, min/max, shortage/overage, outside rotator burden, service readiness."],
    ], [1.7, 4.8], 7.2)

    doc.add_heading("11. Monthly Service Schedule Builder Rules", level=1)
    add_table(doc, ["Rule area", "Detailed rule / question"], RULES_MONTHLY, [1.35, 5.15], 7.0)
    add_table(doc, ["Service example", "Possible custom configuration"], [
        ["PICU", "Unit 1, Unit 2, day coverage, night coverage, long/short call, procedure call, PGY-2/3 eligibility, outside rotators."],
        ["NICU", "Day team, night team, procedure call, long/short call, senior/intern split, no newborn overlap if program requires."],
        ["Purple/Orange floor", "Minimum interns and seniors, colored team assignment, weekend coverage, night stretch, post-call day."],
        ["Gold/Newborn", "Day-only or limited nights, senior support, newborn not eligible for some call types depending on program."],
        ["Jeopardy / Call pool", "J1/J2/J3, clinic call, crossover, no-call lists, call-eligible blocks only."],
    ], [1.2, 5.3], 7.0)
    add_image(doc, "chief-schedule.png", "Example: monthly schedule builder with service tabs, protected time, coverage alerts, and editable draft assignments.", 6.7)

    doc.add_heading("12. Night Shift, Call, and Clinic Protection Rules", level=1)
    add_table(doc, ["Rule family", "Specific questions for discovery"], [
        ["Night eligibility", "Which PGY levels and services can do nights? Are outside rotators eligible? Are interns eligible?"],
        ["Night stretch", "Maximum consecutive nights? Common example discussed: up to 5 nights in a stretch."],
        ["Post-call", "Is post-call day required? What happens if the next block starts immediately after nights?"],
        ["Unsafe next block", "Should a resident ending with night shift avoid starting inpatient block next?"],
        ["Clinic protection", "Can calls be assigned on clinic day? Can calls be assigned the day before clinic? Does this differ by clinic AM/PM?"],
        ["Back-to-back calls", "Are residents allowed to have calls on consecutive days? Which call types count?"],
        ["Call categories", "Clinic call, jeopardy 1/2/3, crossover, 24-hour call, procedure call, long call, short call, custom."],
        ["Fairness counters", "Calls per block/year, weekend calls, night calls, procedure calls, clinic calls, golden weekends, hours."],
    ], [1.35, 5.15], 7.0, PURPLE)

    doc.add_heading("13. Call Switch Workflow Rules", level=1)
    add_table(doc, ["Stage", "Rules and system behavior"], RULES_CALL_SWITCH, [1.15, 5.35], 7.2, TEAL)
    callout(doc, "Why this is high value", "This replaces the current three-person manual process: the offering resident, receiving resident, and chief all checking PDFs, clinics, back-to-back calls, and eligibility by hand.", "success")

    doc.add_heading("14. Attendance and Didactics", level=1)
    add_table(doc, ["Feature", "Details"], [
        ["Attendance tab", "Chiefs/coordinators create a session with block, date, time, lecture type, requirement, and allowed mode."],
        ["Resident check-in", "All resident names appear with checkmarks, designed for tablet/iPad sign-in during lecture."],
        ["Profile linkage", "Clicking attendance updates resident profile, resident view, chief profile, and analytics."],
        ["Lecture types", "Morning report, grand rounds, radiology rounds, simulation, board review, institutional didactics, custom."],
        ["Requirement rules", "Required, optional/service need, excused, virtual allowed, in-person only, not required on certain rotations."],
        ["Statistics", "Attended/missed by week, month, block, year, rotation, institution, and excused vs unexcused."],
    ], [1.4, 5.1], 7.0)

    doc.add_heading("15. PTO, Sick Leave, LOA, Conferences, and Exams", level=1)
    add_table(doc, ["Leave type", "What to collect and track"], [
        ["PTO/vacation", "Requested dates/blocks, priority, approval status, service impact, vacation capacity, resident yearly total."],
        ["Sick leave", "Date, service, clinic impact, inpatient/elective status, coverage needed, running total."],
        ["Maternity/paternity/LOA", "Date range, blocks affected, service coverage impact, privacy rules, chief-only notes."],
        ["Conference/exam/ITE", "Protected day or partial day, travel buffer, whether service coverage is needed, resident profile update."],
        ["Switch/request history", "Every request should show submitted, approved/declined, who decided, when, and why."],
    ], [1.35, 5.15], 7.2)

    doc.add_heading("16. Resident Portal Requirements", level=1)
    add_table(doc, ["Resident area", "What residents can do / see", "Restrictions"], [
        ["My overview", "Summary of current block, upcoming calls, protected time, PTO, attendance status, alerts.", "Only personal data plus published shared schedules."],
        ["My schedule", "Monthly personal schedule with work hours, calls, clinics, didactics, off days, post-call, golden weekends.", "Cannot edit published assignments directly."],
        ["My master schedule", "Private annual master assignment and submitted annual preferences.", "Other residents cannot view it."],
        ["Published schedules", "View shared service schedules published by chiefs.", "Only after publish."],
        ["Requests", "Submit days off, weekends, vacation, holiday preferences, fellowship/elective requests, special needs.", "Deadlines and approval rules apply."],
        ["Call switches", "Offer assigned call, see available calls, check eligibility before accepting.", "Final approval remains chief-owned."],
        ["PTO/sick leave", "Submit/track PTO and sick days and see status.", "Chiefs approve; emergency sick process may differ."],
        ["Attendance", "See personal attendance and missed didactics.", "Cannot edit attendance unless program allows correction requests."],
    ], [1.45, 3.1, 1.95], 6.7)
    add_image(doc, "resident-schedule-main.png", "Example: resident view showing personal schedule, calls, protected activities, and monthly summary.", 6.8)

    doc.add_heading("17. Chief Portal Requirements", level=1)
    add_table(doc, ["Chief area", "Core functions"], [
        ["Overview", "Block readiness, service readiness, conflicts, coverage score, pending approvals, requests, weekend/call/hour stats."],
        ["Master schedule", "Build/import/finalize master schedules by PGY and track; validate service minimums and resident requests."],
        ["Service coverage", "View names and counts by service/block; choose digital or Excel-like grid."],
        ["Schedule builder", "Build monthly schedules by block/service; auto-pull residents, protected time, and approved requests."],
        ["Residents/profiles", "Review and edit individual profiles, clinic plans, didactics, PTO, requests, master schedule, eligibility."],
        ["Rules/institutions", "Configure program rules, service rules, institution rules, shift templates, holidays, blocks."],
        ["Attendance", "Create didactic sessions and check attendance with live profile updates."],
        ["Analytics", "Fairness, workload, hours, calls, weekends, golden weekends, missed didactics, PTO, coverage gaps."],
    ], [1.45, 5.05], 7.2)
    add_image(doc, "chief-overview-main.png", "Example: chief overview designed to show readiness, pending risks, and schedule status.", 6.8)

    doc.add_heading("18. Requests and Approval Center", level=1)
    add_table(doc, ["Request type", "Required status flow", "What changes after approval"], [
        ["Annual master requests", "Submitted -> chief review -> approved/declined/needs info", "Feeds master schedule decision support."],
        ["Vacation/weekend/day off", "Submitted -> conflict check -> chief decision", "Approved requests protect dates in scheduler."],
        ["Call switch", "Offered -> accepted/requested -> eligibility checked -> chief decision", "Updates both residents and published schedule."],
        ["PTO/sick leave", "Submitted/reported -> chief/coordinator review -> approved/recorded", "Updates profile, coverage alerts, leave totals."],
        ["Special circumstance", "Submitted -> private chief review -> applied/not applied", "May affect priority, protected time, or master placement."],
    ], [1.45, 2.3, 2.75], 6.8)
    add_image(doc, "chief-approvals.png", "Example: approval center connects resident requests, switch checks, and chief decisions.", 6.8)

    doc.add_heading("19. Statistics and Alerts Catalog", level=1)
    add_table(doc, ["Statistic / alert", "Where shown", "Why it matters"], [
        ["Residents per service per block", "Master, service coverage, overview", "Prevents short or overstaffed services."],
        ["Outside rotator burden", "Service coverage, overview", "Prevents overcalling rotators in one block and underusing another."],
        ["Weekend days and golden weekends", "Schedule builder, resident stats", "Keeps fairness visible."],
        ["Night shift totals and stretches", "Schedule builder, analytics", "Avoids uneven night burden and unsafe sequences."],
        ["Call totals by type", "Resident profile, analytics", "Tracks J1/J2/J3, clinic call, crossover, procedure call."],
        ["Clinic conflicts", "Schedule builder, call switch checker", "Avoids call on clinic day or day before when prohibited."],
        ["Didactic missed/attended", "Attendance, profile, analytics", "Supports education tracking and interventions."],
        ["PTO/sick leave by service type", "Resident profile, analytics", "Shows if absences affect ICU/floor/ED/elective differently."],
        ["Pending approvals", "Overview, resident profile", "Prevents unapproved requests from silently affecting schedules."],
        ["Unsafe transitions", "Master and monthly builders", "Flags nights followed by inpatient or other prohibited transitions."],
    ], [1.55, 1.65, 3.3], 6.8)

    doc.add_heading("20. Design and User Experience Standards", level=1)
    add_table(doc, ["Design requirement", "Implementation meaning"], DESIGN_REQUIREMENTS, [1.55, 4.95], 7.1, PURPLE)
    callout(doc, "Design principle", "Show the summary first, then reveal details on click. Chiefs need Excel-level detail, but not all at once.", "info")

    doc.add_heading("21. Import and Migration Requirements", level=1)
    add_table(doc, ["Upload type", "What the system should extract", "Human review needed"], [
        ["Existing master Excel", "Sheets by PGY, resident names, block dates, rotations, colors/categories, notes, vacation columns.", "Unmatched names, ambiguous rotations, split blocks, comments, hidden notes."],
        ["Service coverage Excel", "Service rows, block columns, resident names, outside rotators, target counts.", "Service mapping, red/struck names, partial blocks."],
        ["Call pool Excel", "Eligible call pools, no-call lists, floor cross-cover, LOA/unavailable residents.", "Eligibility rules and pool definitions."],
        ["Holiday Excel", "Holiday dates, services, resident names, worked/not worked status.", "Holiday requirements and fairness expectations."],
        ["PDF final schedules", "Published schedule dates, service tabs, residents, work hours, calls, nights, protected time.", "PDF parsing is harder; user review should be expected."],
    ], [1.45, 3.15, 1.9], 6.8)

    doc.add_heading("22. Configuration Split: IT Builds vs Chiefs Customize", level=1)
    add_table(doc, ["Owned by IT/product team", "Owned by chiefs/program team"], [
        ["Database, portals, authentication, permissions, import tools, report templates, audit logs.", "Actual program rules, resident roster, PGY curriculum, service minimums, holiday rules."],
        ["Reusable configuration screens for rotations, services, shift templates, institutions.", "Which rotations/services exist and how they are named for their department."],
        ["Eligibility engine and warning framework.", "Whether a warning is hard stop, soft warning, or informational."],
        ["Digital and Excel-like views.", "Which format they prefer for each workflow and published output."],
        ["Default templates for common specialties.", "Final customization for Pediatrics, IM, ED, Neurology, Surgery, etc."],
    ], [3.25, 3.25], 7.0)

    doc.add_heading("23. MVP Phasing Recommendation", level=1)
    add_table(doc, ["Phase", "Build", "Definition of done"], [
        ["Phase 1: Data foundation", "Resident profiles, institution rules, rotations/services, block calendar, upload/import skeleton.", "Can represent real residents and real master assignments."],
        ["Phase 2: Master + service coverage", "Master builder/import, service coverage workbook, clickable names, counts, Excel-like view.", "Chiefs can validate annual staffing by block and service."],
        ["Phase 3: Monthly scheduler", "Service schedule builder, protected time, shift templates, editable grid/cards, publish resident view.", "One full block can be generated, edited, and published."],
        ["Phase 4: Requests and switches", "Resident request forms, approval center, call switch eligibility checks, PTO/sick leave.", "Approved decisions update schedules/profiles."],
        ["Phase 5: Attendance and analytics", "Didactic attendance check-in, missed counts, dashboards, exports.", "Leadership can track education/workload/fairness."],
    ], [1.35, 2.65, 2.5], 6.9, TEAL)

    doc.add_heading("24. Sales Discovery Meeting Script", level=1)
    numbered(doc, [
        "Start by asking for the department's current Excel/PDF outputs: master schedules, service coverage, call pools, holidays, holiday breaks, and one published monthly block.",
        "Ask them to explain how a resident request moves from form submission to final approval.",
        "Identify their core pain: master building, monthly service scheduling, call switches, vacation/holiday fairness, attendance, or reporting.",
        "Walk through the rule checklist and mark each rule as fixed, configurable, or unknown.",
        "Ask what they want to keep familiar from Excel and where they want a cleaner digital workflow.",
        "End by agreeing on the first MVP test: usually one PGY class, one block, and two or three services.",
    ])

    doc.add_heading("25. Department Configuration Worksheet", level=1)
    add_table(doc, ["Question", "Answer field for discovery notes"], [
        ["Program name and specialty", ""],
        ["Number of blocks and date model", ""],
        ["PGY classes / tracks included", ""],
        ["Core rotations by PGY", ""],
        ["Elective/vacation/call-eligible rotations", ""],
        ["Services and locations to schedule", ""],
        ["Minimum/maximum staffing per service", ""],
        ["Shift templates and call types", ""],
        ["Clinic and didactic protection rules", ""],
        ["Night/post-call rules", ""],
        ["Request priority rules", ""],
        ["Holiday and break coverage rules", ""],
        ["Attendance requirements", ""],
        ["Reports/statistics leadership wants", ""],
        ["Preferred design: digital, Excel-like, or both", ""],
    ], [2.25, 4.25], 7.0)

    doc.add_heading("26. Final Company Positioning", level=1)
    callout(
        doc,
        "Positioning statement",
        "Clarity is a scheduling workflow platform for residency and clinical training programs. It helps departments convert hidden Excel logic into configurable rules, resident-linked data, automated draft schedules, chief-controlled editing, resident transparency, and operational analytics.",
        "success",
    )
    doc.add_paragraph(
        "The system should be sold as both a scheduling product and a discovery framework. The value is not only the software screen; it is the structured way the team helps each hospital uncover, document, and improve its scheduling process."
    )

    doc.save(DOCX_PATH)


def hx(value: str):
    return colors.HexColor(value if value.startswith("#") else f"#{value}")


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("TitleX", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24, leading=28, textColor=hx(DARK), spaceAfter=10))
    styles.add(ParagraphStyle("SubX", parent=styles["Normal"], fontName="Helvetica", fontSize=11, leading=15, textColor=hx(MUTED), spaceAfter=10))
    styles.add(ParagraphStyle("H1X", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, leading=17, textColor=hx(PURPLE), spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle("BodyX", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.1, leading=10.2, textColor=hx(DARK), spaceAfter=4))
    styles.add(ParagraphStyle("CellX", parent=styles["BodyText"], fontName="Helvetica", fontSize=6.3, leading=7.5, textColor=hx(DARK)))
    styles.add(ParagraphStyle("CellBoldX", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=6.4, leading=7.6, textColor=colors.white, alignment=TA_CENTER))
    return styles


def pp(text, style):
    safe = str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(safe, style)


def pdf_table(headers, rows, widths, styles, header_fill=PURPLE):
    data = [[pp(h, styles["CellBoldX"]) for h in headers]]
    for row in rows:
        data.append([pp(v, styles["CellX"]) for v in row])
    table = Table(data, colWidths=[w * inch for w in widths], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), hx(header_fill)),
        ("GRID", (0, 0), (-1, -1), 0.25, hx(LINE)),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, hx("FAFBFE")]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


def pdf_callout(title, body, styles, fill="EAF8F6", border=TEAL):
    text = f"<b>{title}</b><br/>{body}".replace("&", "&amp;")
    table = Table([[Paragraph(text, styles["BodyX"])]], colWidths=[7.05 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), hx(fill)),
        ("BOX", (0, 0), (-1, -1), 0.5, hx(border)),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    return table


def pdf_image(filename, caption, max_width=7.05):
    path = ASSETS / filename
    if not path.exists():
        return []
    img = Image(str(path))
    img._restrictSize(max_width * inch, 3.2 * inch)
    return [img, Spacer(1, 4), Paragraph(caption, pdf_styles()["BodyX"]), Spacer(1, 10)]


def create_pdf():
    OUT.mkdir(exist_ok=True)
    styles = pdf_styles()
    story = []
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.45 * inch,
    )
    story.append(Paragraph("Clarity Scheduling System", styles["TitleX"]))
    story.append(Paragraph("Business blueprint, discovery guide, and implementation rulebook", styles["SubX"]))
    story.append(pdf_callout("Purpose", "A detailed guide for sales, IT, chiefs, coordinators, and hospital leadership to understand what the platform builds, what rules must be collected, and what screens must be linked.", styles))
    story.append(Spacer(1, 8))
    story.append(pdf_table(["Step", "Workflow", "What must be linked"], CORE_WORKFLOW, [0.38, 1.45, 5.2], styles))
    story.append(PageBreak())

    story.append(Paragraph("Visual Examples for Sales and IT Conversations", styles["H1X"]))
    story.append(Paragraph(
        "These screenshots are included to make the discussion concrete. They show the design direction: stepwise workflows, chief and resident portals, linked profiles, schedule workspaces, request approvals, and readiness views.",
        styles["BodyX"],
    ))
    for filename, caption in [
        ("chief-master.png", "Chief master schedule view: capacity, conflicts, editable annual source data, and profile-linked residents."),
        ("chief-schedule.png", "Monthly schedule builder: service tabs, resident rows, shift cards, coverage alerts, and editable drafts."),
        ("resident-schedule-main.png", "Resident portal: personal schedule, protected time, calls, master schedule access, requests, and PTO areas."),
        ("chief-overview-main.png", "Chief overview: block readiness, pending risks, coverage, fairness, and attention center."),
        ("chief-approvals.png", "Approval center: resident requests, call switch checks, and chief decision workflow."),
    ]:
        story.extend(pdf_image(filename, caption, 7.05))
    story.append(PageBreak())

    sections = [
        ("Product Thesis", [["Layer", "What it means"], [["Core scheduling engine", "Master schedules, service coverage, call pools, holiday coverage, monthly service schedules."], ["Workflow tools", "Request approval, call switches, PTO/sick leave, attendance, publishing, audit trail."], ["Decision support", "Coverage gaps, over-capacity, fairness, hours, weekends, calls, missed didactics, conflict warnings."]], [1.5, 5.5], PURPLE]),
        ("Core Modules and Ownership", [["Module", "Purpose", "Feeds / depends on"], [[m[0], m[1], m[3]] for m in MODULES], [1.35, 2.0, 3.65], PURPLE]),
        ("Critical Linkage Map", [["Data / action", "Starts in", "Must update", "Business rule"], LINKAGES, [1.35, 1.2, 1.55, 2.9], TEAL]),
        ("Discovery Checklist for Any Department", [["Category", "Questions to ask"], DISCOVERY, [1.25, 5.75], PURPLE]),
        ("Master Schedule Rules Checklist", [["Rule area", "Detailed rule / question"], RULES_MASTER, [1.35, 5.65], PURPLE]),
        ("Annual Resident Intake Forms", [["Form section", "Fields to collect", "Where it links"], [["Fixed resident facts", "Name, email, PGY, institution, track, fellowship interest, outside rotator status", "Resident profile, eligibility checks, institution rules"], ["Master schedule ranking", "Rank templates/preferences; enforce unique ranks; timestamp submission", "Chief master builder"], ["Vacation/holiday/elective requests", "Preferred blocks, holiday rank, top electives, comments", "Master builder, holiday workbook, PTO workflow"], ["Life events and special needs", "Maternity/paternity, wedding, visa/travel, interviews, board review", "Chief review and protected time"]], [1.45, 3.1, 2.45], TEAL]),
        ("Master Builder Design Requirements", [["Step", "Screen goal", "Required functions"], [["Choose PGY", "Select class/curriculum", "Blocks, residents/templates, rotations, min/max, eligible blocks, drag/add/delete rotations"], ["Resident rankings", "Review preferences", "Sort by submission, alphabet, template; open resident profile"], ["Block coverage", "See counts before final grid", "Clickable block/service counts, names, min/max, gaps"], ["Finalize master", "Editable annual grid", "Drag/drop, lock, validate, save source, export Excel-like view"], ["Import existing master", "Skip builder", "Upload Excel/PDF, parse, map, review, approve import"]], [1.2, 1.8, 4.0], TEAL]),
        ("Monthly Service Schedule Builder Rules", [["Rule area", "Detailed rule / question"], RULES_MONTHLY, [1.35, 5.65], PURPLE]),
        ("Night Shift, Call, and Clinic Protection Rules", [["Rule family", "Specific questions for discovery"], [["Night eligibility", "Which PGY levels/services can do nights? Outside rotators? Interns?"], ["Night stretch", "Maximum consecutive nights, post-call rules, unsafe next block rules."], ["Clinic protection", "Can calls occur on clinic day or the day before? Does AM/PM matter?"], ["Call categories", "Clinic call, J1/J2/J3, crossover, procedure call, long/short call, 24-hour call."], ["Fairness counters", "Calls, weekends, nights, golden weekends, hours, days off."]], [1.35, 5.65], PURPLE]),
        ("Call Switch Workflow Rules", [["Stage", "Rules and system behavior"], RULES_CALL_SWITCH, [1.2, 5.8], TEAL]),
        ("Attendance and Didactics", [["Feature", "Details"], [["Attendance tab", "Create lecture session by block/date/time/type/requirement/mode."], ["Resident check-in", "Tablet-friendly list of residents with checkmarks."], ["Profile linkage", "Updates resident/chief profile and analytics."], ["Statistics", "Attended/missed by week, block, year, rotation, institution, excused/unexcused."]], [1.45, 5.55], PURPLE]),
        ("PTO, Sick Leave, LOA, Conferences, Exams", [["Leave type", "What to collect and track"], [["PTO/vacation", "Dates/blocks, priority, approval, service impact, vacation capacity, yearly total."], ["Sick leave", "Date, service, clinic impact, inpatient/elective status, coverage needed, running total."], ["Maternity/LOA", "Date range, blocks affected, privacy rules, coverage impact."], ["Conference/exam/ITE", "Protected day/partial day, travel buffer, coverage need, resident profile update."]], [1.4, 5.6], PURPLE]),
        ("Resident and Chief Portal Requirements", [["Portal", "Required areas"], [["Resident", "My overview, my schedule, private master, published schedules, requests, call switches, PTO/sick leave, attendance."], ["Chief", "Overview, master schedule, service coverage, monthly builder, residents/profiles, rules/institutions, attendance, analytics."]], [1.2, 5.8], TEAL]),
        ("Statistics and Alerts Catalog", [["Statistic / alert", "Where shown", "Why it matters"], [["Residents per service per block", "Master, service coverage, overview", "Prevents short or overstaffed services."], ["Outside rotator burden", "Service coverage", "Prevents overcalling one block."], ["Weekend/golden weekend fairness", "Schedule builder", "Keeps fairness visible."], ["Night/call totals", "Profile and analytics", "Avoids uneven workload."], ["Clinic conflicts", "Call switch/scheduler", "Prevents unsafe assignments."], ["Attendance missed/attended", "Attendance/profile", "Supports education tracking."]], [1.8, 1.55, 3.65], PURPLE]),
        ("Design and User Experience Standards", [["Design requirement", "Implementation meaning"], DESIGN_REQUIREMENTS, [1.65, 5.35], PURPLE]),
        ("Import and Migration Requirements", [["Upload type", "What the system should extract", "Human review needed"], [["Existing master Excel", "PGY tabs, names, blocks, rotations, colors, notes, vacation columns", "Unmatched names, ambiguous rotations, split blocks"], ["Service/call/holiday Excel", "Service rows, block columns, names, pools, holiday staffing", "Mappings, red/struck names, partial blocks"], ["PDF final schedules", "Dates, services, residents, work hours, calls, nights, protected time", "PDF parsing needs review"]], [1.45, 3.35, 2.2], TEAL]),
        ("MVP Phasing Recommendation", [["Phase", "Build", "Definition of done"], [["Phase 1", "Data foundation", "Can represent real residents and master assignments."], ["Phase 2", "Master + service coverage", "Chiefs validate annual staffing."], ["Phase 3", "Monthly scheduler", "One block can be generated, edited, published."], ["Phase 4", "Requests and switches", "Approved decisions update schedules/profiles."], ["Phase 5", "Attendance and analytics", "Leadership can track education/workload/fairness."]], [1.0, 2.5, 3.5], TEAL]),
    ]
    for title, (headers, rows, widths, fill) in sections:
        story.append(Paragraph(title, styles["H1X"]))
        story.append(pdf_table(headers, rows, widths, styles, fill))
        story.append(PageBreak())

    story.append(Paragraph("Department Configuration Worksheet", styles["H1X"]))
    story.append(pdf_table(["Question", "Answer field for discovery notes"], [["Program name and specialty", ""], ["Number of blocks and date model", ""], ["PGY classes / tracks", ""], ["Core rotations by PGY", ""], ["Elective/vacation/call-eligible rotations", ""], ["Services and locations", ""], ["Minimum/maximum staffing", ""], ["Shift templates and call types", ""], ["Clinic and didactic protection rules", ""], ["Night/post-call rules", ""], ["Request priority rules", ""], ["Holiday and break rules", ""], ["Attendance requirements", ""], ["Reports/statistics needed", ""], ["Preferred design", "Digital, Excel-like, or both"]], [2.4, 4.6], styles))
    doc.build(story)


if __name__ == "__main__":
    create_docx()
    create_pdf()
    with ZipFile(DOCX_PATH) as z:
        assert "word/document.xml" in z.namelist()
    print(DOCX_PATH)
    print(PDF_PATH)
