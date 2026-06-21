from pathlib import Path
from datetime import date
from textwrap import shorten

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)

DOCX_PATH = OUT / "Clarity_Scheduling_System_Implementation_Playbook.docx"
PDF_PATH = OUT / "Clarity_Scheduling_System_Implementation_Playbook.pdf"
ASSETS = OUT / "brochure-v2-assets-anon"

PURPLE = "4F4698"
PURPLE_DARK = "24243E"
TEAL = "299A8F"
AMBER = "D1912E"
RED = "C65353"
INK = "22263A"
MUTED = "697087"
PALE = "F4F1FF"
PALE_TEAL = "EAF8F6"
PALE_AMBER = "FFF6E4"
PALE_RED = "FFF0F0"
LINE = "D9DDE8"
WHITE = "FFFFFF"
LIGHT = "F7F8FC"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def style_run(run, size=None, color=None, bold=None):
    run.font.name = "Aptos"
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold


def configure_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    style_specs = [
        ("Title", 26, PURPLE_DARK, 0, 8),
        ("Subtitle", 12.5, MUTED, 0, 14),
        ("Heading 1", 16, PURPLE, 14, 7),
        ("Heading 2", 12.8, PURPLE_DARK, 10, 5),
        ("Heading 3", 10.8, PURPLE, 7, 3),
    ]
    for name, size, color, before, after in style_specs:
        st = styles[name]
        st.font.name = "Aptos Display" if name in ("Title", "Heading 1", "Heading 2") else "Aptos"
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.font.bold = name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = name.startswith("Heading")

    header = section.header.paragraphs[0]
    header.text = "CLARITY SCHEDULING SYSTEM  |  IMPLEMENTATION PLAYBOOK"
    style_run(header.runs[0], 7.8, MUTED, True)

    footer = section.footer.paragraphs[0]
    footer.add_run("Implementation playbook  |  Prepared for residency scheduling discovery     ")
    style_run(footer.runs[0], 7.8, MUTED, False)
    add_page_number(footer)
    return doc


def add_docx_callout(doc, title, body, tone="purple"):
    fills = {"purple": PALE, "teal": PALE_TEAL, "amber": PALE_AMBER, "red": PALE_RED}
    accents = {"purple": PURPLE, "teal": TEAL, "amber": AMBER, "red": RED}
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.68)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fills[tone])
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    style_run(r, 10.2, accents[tone], True)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        style_run(run, 9.2, INK, False)


def add_docx_table(doc, headers, rows, widths=None, font_size=8.0, header_fill=PURPLE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Inches(width)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(headers) > 3 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(header))
        style_run(r, font_size, WHITE, True)
    repeat_header(table.rows[0])

    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        no_split(table.rows[-1])
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_margins(cell)
            if r_idx % 2:
                set_cell_shading(cell, "FAFBFE")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if isinstance(value, (int, float)) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            style_run(r, font_size, INK, False)
    return table


def add_docx_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            r = p.add_run(item[0] + ": ")
            style_run(r, 9.3, INK, True)
            r = p.add_run(item[1])
            style_run(r, 9.3, INK, False)
        else:
            r = p.add_run(item)
            style_run(r, 9.3, INK, False)


def add_docx_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        style_run(r, 9.3, INK, False)


def add_docx_image(doc, image_path, caption):
    image_path = Path(image_path)
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(6.35))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(1)
    cap.paragraph_format.space_after = Pt(7)
    r = cap.add_run(caption)
    style_run(r, 8.0, MUTED, False)


def flow_rows():
    return [
        ["1", "Annual setup", "Blocks, PGY classes, rotations, service minimums, holidays, institutions", "Chiefs + IT"],
        ["2", "Resident intake", "Master ranking, vacation requests, holidays, fellowship timing, life events, elective choices", "Residents"],
        ["3", "Master schedule", "Assign each resident to a rotation per block while meeting service minimums", "Chiefs"],
        ["4", "Service coverage", "Convert master rotations into service rosters by block, service, role, and location", "Chiefs"],
        ["5", "Call pool", "Pull elective or call-eligible residents into clinic call, jeopardy, cross-cover, no-call, and LOA pools", "Chiefs"],
        ["6", "Holiday coverage", "Assign holiday staffing using service needs and approved holiday requests", "Chiefs"],
        ["7", "Monthly schedules", "Build editable service schedules with shifts, calls, clinics, didactics, leave, and coverage checks", "Chiefs"],
        ["8", "Publish and track", "Resident views, switches, PTO, attendance, statistics, and audit history", "Chiefs + Residents"],
    ]


DATA_TABLE = [
    ["Program structure", "Fixed yearly configuration", "Chiefs / program leadership", "Academic year, number of blocks, block dates, PGY classes, departments, institutions"],
    ["Resident profile", "Mostly fixed, updated as needed", "Resident + chiefs", "Name, PGY, institution, track, clinic pattern, didactic rules, eligibility flags"],
    ["Resident requests", "Variable, submission-based", "Resident", "Master preference rank, vacation/PTO, holiday rank, elective rank, fellowship timing, comments"],
    ["Rotation rules", "Fixed per PGY and department", "Chiefs", "Core rotations, required counts, eligible blocks, vacation-eligible blocks, unsafe transitions"],
    ["Service rules", "Configurable by service", "Chiefs with IT support", "Service names, roles, staffing min/max, locations, shift types, night rules, call types"],
    ["Monthly events", "Variable by block", "Chiefs / coordinator", "Holidays, ITE, retreats, conferences, exam days, one-time clinic changes"],
    ["Generated schedules", "System-created, chief-editable", "System + chiefs", "Master, service rosters, call pools, holidays, monthly service schedules"],
    ["Operational records", "Live transactional data", "Residents + chiefs", "Call switches, PTO approvals, sick leave, attendance, publish status, notes"],
]


MODULE_TABLE = [
    ["Master Schedule Builder", "Creates the annual source of truth", "PGY templates, resident rankings, vacation requests, rotation minimums", "Final master grid, service coverage counts, resident master view"],
    ["Service Coverage Workbook", "Shows who covers every service each block", "Final master schedule, service mapping rules", "Service roster by block, gap/overage alerts, Excel-like view"],
    ["Call Pool Workbook", "Builds eligible call and no-call pools", "Elective/call-eligible master blocks, leave, no-call rules", "Clinic call, jeopardy, cross-cover, no-call, LOA lists"],
    ["Holiday Workbook", "Plans coverage around holidays", "Holiday calendar, resident holiday requests, service minimums", "Holiday staffing grid, gap alerts, resident holiday assignments"],
    ["Monthly Schedule Builder", "Builds editable service schedules", "Service roster, shift templates, clinics, didactics, PTO, requests", "Published service schedules, resident personal schedules, statistics"],
    ["Resident Portal", "Collects requests and shows personal schedule", "Published schedules, individual profile, request forms", "Requests, call switches, PTO/sick leave, personal statistics"],
    ["Attendance", "Tracks didactic presence", "Lecture sessions, resident roster, attendance requirements", "Resident attendance history, missed-session counts, chief reports"],
    ["Rules and Institutions", "Stores reusable program logic", "Institution profiles, service rules, shift templates, safety rules", "Rules automatically applied to future schedules"],
]


IT_CHIEF_TABLE = [
    ["Build once by IT", "Authentication, database, import engine, audit logs, permission model, scheduling rule engine, export tools, notification framework"],
    ["Configure with chiefs", "PGY classes, block calendar, service catalog, rotation-to-service mapping, shift templates, min/max staffing, holiday rules, call switch rules"],
    ["Chiefs manage yearly", "Resident roster, outside rotators, approved requests, master schedule decisions, service coverage corrections, holiday staffing"],
    ["Residents submit", "Annual preferences, vacation/PTO, holiday requests, elective rankings, call switch offers/acceptances, sick leave, attendance check-ins"],
    ["System calculates", "Coverage gaps, overages, fairness stats, conflict flags, clinic/didactic conflicts, back-to-back call risk, attendance totals"],
]


FEATURE_TABLE = [
    ["Core MVP", "Resident profiles, annual intake, master schedule, service coverage, monthly schedule builder, publishable resident view"],
    ["Operational MVP", "Call pool, PTO/sick leave, call switches, approval queue, attendance check-in, export to Excel/PDF"],
    ["Advanced MVP", "Rule scoring, optimization suggestions, import from legacy Excel/PDF, notifications, audit trail, analytics dashboard"],
    ["Future expansion", "Multi-program support, role-based administration, mobile check-in, integrations with Outlook/Google calendar, institutional reporting"],
]


INTERVIEW_CHECKLIST = [
    ["Program map", "How many PGY classes, residents, institutions, tracks, outside rotators, and block systems exist?"],
    ["Annual master", "Which rotations are core, elective, vacation eligible, call eligible, or restricted by PGY year?"],
    ["Coverage targets", "What minimum and maximum number of residents must cover each service per block? Are ranges different by PGY?"],
    ["Service logic", "For each service, what roles exist: day, night, short call, long call, procedure call, unit 1, unit 2, senior, intern?"],
    ["Protected time", "Which clinics, didactics, exams, retreats, conferences, and institution-specific rules must block scheduling?"],
    ["Requests", "What can residents request, what is the deadline, what is first-come-first-serve, and what requires chief discretion?"],
    ["Fairness", "What should be balanced: nights, weekends, golden weekends, holidays, call types, clinic calls, hours, or difficult rotations?"],
    ["Safety", "Which sequences are unsafe: post-night to inpatient, back-to-back calls, call before clinic, max consecutive nights, duty hours?"],
    ["Outputs", "What published PDFs/Excels are expected today, and what digital/Excel-like views must replace them?"],
    ["Audit", "Who can approve, edit, override, publish, view resident-private data, and see other residents' schedules?"],
]


def create_docx():
    doc = configure_docx()
    add_docx_callout(
        doc,
        "Purpose of this playbook",
        "This document translates the residency scheduling process into a structured system blueprint. It is designed for IT, chiefs, coordinators, and program leadership to gather the right information, build the right modules, and avoid missing the hidden clinical rules that make scheduling difficult.",
        "teal",
    )
    p = doc.add_paragraph("Clarity Scheduling System")
    p.style = "Title"
    doc.add_paragraph("Implementation playbook, data-gathering guide, and product structure for a residency scheduling MVP", style="Subtitle")
    doc.add_paragraph(f"Prepared: {date.today().strftime('%B %d, %Y')}")
    add_docx_table(
        doc,
        ["Who should use this", "How to use it"],
        [
            ["IT / product team", "Use this as the discovery checklist, data model starter, and implementation map."],
            ["Chief residents", "Use this to explain scheduling rules, exceptions, fairness goals, and monthly workflows."],
            ["Program leadership", "Use this to define scope, approve priorities, and understand operational value."],
        ],
        [2.0, 4.5],
        8.4,
        TEAL,
    )

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "The scheduling system should not be treated as one calendar. It is a layered clinical operations system. "
        "The annual master schedule becomes the source of truth; service coverage converts that master into staffing rosters; call pools and holiday schedules manage special coverage; monthly schedules become editable service-specific outputs; and resident profiles keep every clinic, didactic, PTO, call switch, and attendance record connected."
    )
    add_docx_callout(
        doc,
        "The central design principle",
        "Every schedule should be generated from structured resident profiles and configurable department rules, then remain editable by chiefs. The system should suggest and warn, but chiefs retain final control.",
        "purple",
    )
    add_docx_table(doc, ["Step", "Workflow", "Main decision", "Owner"], flow_rows(), [0.45, 1.35, 3.65, 1.1], 7.4)

    doc.add_heading("2. Roles and Permissions", level=1)
    add_docx_table(
        doc,
        ["Role", "Needs", "Should be able to"],
        [
            ["Resident", "Submit requests and understand personal obligations", "View personal master, published service schedules, calls, PTO, switches, attendance, and submit requests"],
            ["Chief resident", "Build and adjust all schedules", "Configure rules, approve requests, edit schedules, publish outputs, override warnings, export reports"],
            ["Program director / leadership", "Oversight and risk visibility", "View dashboards, coverage trends, resident workload summaries, and policy reports"],
            ["Coordinator / admin", "Operational support", "Import data, maintain resident roster, add events, help track attendance and leave"],
            ["IT / product admin", "Configuration and maintenance", "Manage templates, integrations, backups, access, and department-level customization"],
        ],
        [1.15, 2.25, 3.1],
        7.8,
    )

    doc.add_heading("3. Data Taxonomy: Fixed, Variable, and Generated", level=1)
    doc.add_paragraph(
        "A major source of scheduling pain is that fixed rules, resident preferences, and generated assignments are often mixed together inside Excel. The MVP should separate them clearly."
    )
    add_docx_table(doc, ["Data type", "Nature", "Provided by", "Examples"], DATA_TABLE, [1.4, 1.3, 1.35, 2.45], 7.0)

    doc.add_heading("4. Recommended Product Modules", level=1)
    add_docx_table(doc, ["Module", "Purpose", "Inputs", "Outputs"], MODULE_TABLE, [1.45, 1.55, 1.9, 1.6], 6.7)
    add_docx_image(doc, ASSETS / "chief-overview-main.png", "Example chief overview: readiness, conflicts, and schedule status should be visible without opening multiple files.")

    doc.add_heading("5. Department Discovery Process", level=1)
    doc.add_paragraph("Use this sequence when meeting with any department. It keeps the conversation organized and reveals rules that teams often forget to mention.")
    add_docx_numbered(doc, [
        "Map the annual structure: academic year, number of blocks, block dates, PGY classes, tracks, institutions, and rotator types.",
        "Define resident profile fields: PGY, institution, track, clinic pattern, didactics, electives, PTO balances, and eligibility flags.",
        "Build the service catalog: every team, unit, rotation, call pool, and special responsibility that must appear in a schedule.",
        "Collect service rules: minimum coverage, maximum coverage, allowed PGY levels, shift types, night rules, weekend rules, and backup logic.",
        "Collect protected time: clinics, didactics, exams, retreats, holidays, conferences, maternity leave, and institution-specific absences.",
        "Collect resident request rules: annual rankings, vacation requests, holiday requests, elective ranking, switch requests, and approval workflow.",
        "Define outputs: what chiefs publish today, what residents need to see, and which Excel-like views must remain available.",
        "Define statistics: duty hours, weekend days, golden weekends, nights, clinic calls, jeopardy, attendance, sick days, PTO, and coverage gaps.",
    ])
    add_docx_table(doc, ["Interview area", "Questions to ask"], INTERVIEW_CHECKLIST, [1.65, 4.85], 7.5, TEAL)

    doc.add_heading("6. Annual Master Schedule Builder", level=1)
    doc.add_paragraph(
        "The annual master schedule is the foundation. It assigns each resident to one rotation per block and feeds every downstream schedule. "
        "For programs with PGY-1, PGY-2, PGY-3, Med-Peds, or other tracks, each class may require a different curriculum template."
    )
    add_docx_table(
        doc,
        ["Configuration area", "Chiefs provide", "IT builds as options"],
        [
            ["PGY setup", "PGY classes, number of residents, number of master schedules needed", "PGY selector, resident count matching, missing-profile alerts"],
            ["Rotation catalog", "Core rotations, electives, vacation-eligible rotations, call-eligible rotations", "Editable rotation cards with add, rename, drag, delete, color, and service mapping"],
            ["Coverage targets", "Min/max residents per service per block", "Coverage counters, shortage/overage warnings, block detail drawer"],
            ["Resident preference intake", "Master ranking, vacation blocks, fellowship timing, elective ranking, special notes", "Annual form, timestamps, sortable resident ranking queue, profile links"],
            ["Finalize grid", "Manual overrides and final chief decisions", "Drag/drop editable grid, lock/finalize button, audit trail, source data for all modules"],
        ],
        [1.45, 2.55, 2.5],
        7.4,
    )
    add_docx_image(doc, ASSETS / "chief-master.png", "Example master schedule view: the final grid should remain editable, but it should also show coverage warnings.")

    doc.add_heading("7. Service Coverage Workbook", level=1)
    doc.add_paragraph(
        "After the master schedule is finalized, the system should automatically produce a service coverage workbook. "
        "This replaces the Excel tab where chiefs list who is on each service for every block."
    )
    add_docx_table(
        doc,
        ["Requirement", "System behavior"],
        [
            ["Pull names from master", "If a resident is assigned to PICU in Block 2, they appear in the PICU Block 2 service roster."],
            ["Support service aliases", "A master rotation such as Floor can feed Purple, Orange, Gold, or floor senior service depending on rules."],
            ["Show counts first", "Keep the main view clean with counts, gaps, and status; reveal names when a service or block is clicked."],
            ["Allow Excel-like view", "Chiefs can switch to a familiar spreadsheet-style roster while transitioning away from Excel."],
            ["Link to profiles", "Clicking a resident opens their profile and then returns the chief to the same service context."],
        ],
        [1.65, 4.85],
        7.8,
    )
    add_docx_image(doc, ASSETS / "service-builder-main.png", "Example service/rules builder: each service needs its own roles, shifts, minimums, and location structure.")

    doc.add_heading("8. Call Pool, Holiday, and Holiday Break Workbooks", level=1)
    add_docx_table(
        doc,
        ["Workbook", "Who gets pulled in", "Chief review questions"],
        [
            ["Call pool", "Residents in elective, call-eligible, or vacation-eligible blocks", "Who can do clinic call, jeopardy, cross-cover, no-call, or LOA?"],
            ["Holiday coverage", "Residents available during holiday blocks, filtered by service needs and approved holiday requests", "Are PICU, NICU, floors, cardiology, heme/onc, ED, call, and jeopardy covered?"],
            ["Holiday breaks", "Residents assigned to Christmas/New Year break staffing windows", "Are minimum services covered for both break weeks? Who is off and who is working?"],
        ],
        [1.25, 2.65, 2.6],
        7.5,
        AMBER,
    )
    add_docx_callout(
        doc,
        "Keep names available without crowding",
        "Use card views and counts for the default screen. When chiefs click a call pool, holiday, or service cell, show the names, source rotation, eligibility reason, and profile links in a detail drawer or section below.",
        "amber",
    )

    doc.add_heading("9. Monthly Service Schedule Builder", level=1)
    doc.add_paragraph(
        "Monthly scheduling starts after the master and service coverage layers are ready. Each block has multiple service schedules such as NICU, PICU, Orange, Purple, Gold, Newborn, Heme/Onc, Jeopardy, and Night Senior. Each service has different rules."
    )
    add_docx_table(
        doc,
        ["Builder step", "What should happen"],
        [
            ["Block setup", "Select block, service, dates, holidays, service locations, and whether to use existing master/service coverage data."],
            ["People", "Pull residents from the master schedule for that service and block; allow chiefs to add outside rotators or call-pool residents."],
            ["Protected time", "Auto-pull clinic, didactics, ITE, conferences, PTO, leave, and institution rules; allow edits and one-time overrides."],
            ["Requests", "Show approved, pending, and denied requests as reminders. Do not let unapproved requests drive the schedule automatically."],
            ["Rules", "Apply service-specific shift templates, min coverage, PGY eligibility, nights, weekends, long/short call, procedure call, and fairness targets."],
            ["Review and generate", "Create an editable draft, flag gaps/conflicts, allow drag/drop and manual changes, then publish."],
        ],
        [1.4, 5.1],
        7.5,
    )
    add_docx_image(doc, ASSETS / "chief-schedule.png", "Example monthly schedule editor: generated templates should stay editable and show coverage/fairness checks.")

    doc.add_heading("10. Resident Portal", level=1)
    add_docx_table(
        doc,
        ["Resident feature", "Purpose", "Linked chief-side data"],
        [
            ["My schedule", "Personal month view with work hours, calls, clinics, didactics, PTO, and weekends", "Published service schedules and resident profile"],
            ["My master schedule", "Private annual schedule visible only to resident and chiefs", "Final master schedule"],
            ["Published schedules", "Shared service rosters for the month", "Published service schedules"],
            ["Requests and approvals", "Day off, weekend, holiday, elective, fellowship, or special requests", "Chief approval queue"],
            ["Call switches", "Offer a call, accept a switch, and receive eligibility checks", "Call schedule, clinic rules, back-to-back call rules"],
            ["PTO and sick leave", "Submit and track leave balances", "Leave tracker and monthly schedule builder"],
            ["Attendance", "See didactic sessions attended/missed", "Attendance module and resident profile"],
        ],
        [1.45, 2.3, 2.75],
        7.1,
        TEAL,
    )
    add_docx_image(doc, ASSETS / "resident-schedule-main.png", "Example resident personal schedule: residents see only their own private schedule plus shared published rosters.")

    doc.add_heading("11. Requests, Approvals, and Communication", level=1)
    doc.add_paragraph(
        "Requests should be advisory until approved. The system can show conflicts and suggestions, but the final decision should remain with chiefs."
    )
    add_docx_table(
        doc,
        ["Request type", "Resident provides", "System checks", "Chief action"],
        [
            ["Annual master preference", "Ranked master templates, fellowship timing, track, comments", "Duplicate ranks, timing conflicts, coverage impact", "Approve inputs, adjust assignment"],
            ["Vacation/PTO", "Dates or blocks, priority, reason, backup flexibility", "Capacity, rotation type, holiday conflict", "Approve, deny, or request change"],
            ["Holiday request", "Ranked holidays desired off", "Coverage impact, prior year equity if tracked", "Assign holiday coverage"],
            ["Call switch", "Call offered, desired date/alternate, receiving resident", "Clinic same day/next day, back-to-back calls, duplicate assignment", "Approve switch or deny"],
            ["Sick leave", "Date, affected service, reason category", "Coverage gap and duty impact", "Record, arrange coverage"],
        ],
        [1.15, 1.4, 1.95, 1.55],
        6.6,
    )
    add_docx_image(doc, ASSETS / "chief-approvals.png", "Example approvals view: chiefs need request context, system checks, and final approve/decline controls.")

    doc.add_heading("12. Attendance and Didactics", level=1)
    add_docx_table(
        doc,
        ["Feature", "How it should work"],
        [
            ["Attendance session setup", "Chief selects block, date, time, lecture type, requirement, and virtual/in-person rule."],
            ["Resident check-in", "Residents tap/check their name from a roster during didactics."],
            ["Profile link", "Attendance updates resident profile, chief view, and resident view automatically."],
            ["Requirement logic", "Some rotations require attendance, some allow virtual attendance, and some are optional due to service needs."],
            ["Reports", "Show missed per week, month, block, year, and reason if protected by service obligation."],
        ],
        [1.35, 5.15],
        7.6,
    )

    doc.add_heading("13. Statistics and Decision Support", level=1)
    add_docx_bullets(doc, [
        ("Coverage", "Residents scheduled per service, per block, per role, with short/over/covered status."),
        ("Fairness", "Weekend days, golden weekends, night shifts, call counts, holiday distribution, clinic calls, jeopardy calls."),
        ("Duty burden", "Expected hours, nights in a row, post-call recovery, inpatient after night block, unsafe transition alerts."),
        ("Resident history", "PTO days, sick days, didactic attendance, missed sessions, call switches, approvals, comments."),
        ("Operational readiness", "Which blocks and services are ready, pending, missing outside rotators, or not yet published."),
    ])

    doc.add_heading("14. What IT Should Build vs What Chiefs Should Customize", level=1)
    add_docx_table(doc, ["Ownership", "Examples"], IT_CHIEF_TABLE, [1.7, 4.8], 7.8)
    add_docx_callout(
        doc,
        "Provider-perspective advantage",
        "The product should not force chiefs to become software configurators. IT should build reusable options for common clinical scheduling patterns, while chiefs provide the medical judgment: coverage minimums, fairness priorities, and final approval.",
        "teal",
    )

    doc.add_heading("15. Design and Usability Standards", level=1)
    add_docx_bullets(doc, [
        ("Default to clean digital views", "Use cards, counts, status labels, and drill-down details so chiefs are not overwhelmed."),
        ("Preserve Excel-like views", "Offer spreadsheet-style views for master, service coverage, call pools, holidays, and monthly schedules."),
        ("Make every count clickable", "A number should open names, source rules, missing coverage, and profile links."),
        ("Support step-wise workflows", "Complex scheduling should be broken into setup, people, protected time, rules, review, and publish."),
        ("Keep source links visible", "Every assignment should show whether it came from master schedule, resident profile, chief edit, or imported file."),
        ("Make warnings advisory", "The system should explain risk and suggest fixes, but not block chiefs unless a hard rule is configured."),
    ])
    add_docx_image(doc, ASSETS / "institution-rules.png", "Example institution rules: reusable patterns prevent repeated manual entry for outside rotators and clinic/didactic rules.")

    doc.add_heading("16. MVP Build Phases", level=1)
    add_docx_table(doc, ["Phase", "Scope"], FEATURE_TABLE, [1.4, 5.1], 7.8, PURPLE)
    add_docx_table(
        doc,
        ["MVP acceptance test", "Pass condition"],
        [
            ["Import existing master Excel", "Residents, PGY, blocks, rotations, service coverage, and profiles are created correctly."],
            ["Build service schedule", "Selecting a block/service pulls correct residents from master and flags gaps."],
            ["Apply protected time", "Clinic/didactic/PTO rules appear automatically before shifts are assigned."],
            ["Publish schedule", "Resident view updates with personal schedule and shared roster."],
            ["Submit call switch", "Eligibility checks show clinic/back-to-back conflicts and chiefs can approve or decline."],
            ["Take attendance", "Checking a resident updates attendance totals in chief and resident views."],
            ["Export", "Chiefs can produce a PDF/Excel-like output similar enough to their existing process."],
        ],
        [2.2, 4.3],
        7.6,
        TEAL,
    )

    doc.add_heading("17. Implementation Data Model Starter", level=1)
    add_docx_table(
        doc,
        ["Object", "Key fields"],
        [
            ["Resident", "id, name, pgy, institution, track, email, clinic_patterns, didactic_rules, leave_balance, eligibility_flags"],
            ["Block", "id, name, start_date, end_date, year, holidays, request_deadline"],
            ["Rotation", "id, name, pgy_allowed, core_required, vacation_eligible, call_eligible, service_mapping, color"],
            ["Service", "id, name, category, locations, roles, shift_templates, min/max staffing, fairness_rules"],
            ["MasterAssignment", "resident_id, block_id, rotation_id, source, locked, notes, conflicts"],
            ["ServiceCoverage", "block_id, service_id, resident_ids, target_min, target_max, status, gaps"],
            ["Request", "resident_id, type, dates/block, priority, submitted_at, status, chief_decision, notes"],
            ["ScheduleAssignment", "resident_id, service_id, block_id, date, role, shift_template, source, status"],
            ["AttendanceSession", "date, block_id, lecture_type, requirement, allowed_mode, attendees, excused_absences"],
        ],
        [1.45, 5.05],
        7.0,
    )

    doc.add_heading("18. Team Working Checklist", level=1)
    add_docx_bullets(doc, [
        "Before discovery: collect existing Excel/PDF examples, resident intake forms, service rules, and published schedules.",
        "During discovery: separate facts from preferences, hard rules from soft rules, and chief decisions from resident requests.",
        "After discovery: convert every rule into either a data field, configurable option, warning, approval step, or report.",
        "Before MVP testing: create fake resident accounts, imported master schedule data, institution patterns, service rules, and monthly requests.",
        "During MVP testing: test one PGY class, one block, one inpatient service, one call pool, one holiday, one call switch, and one attendance session end to end.",
        "Before pilot launch: confirm privacy, permissions, export format, chief override workflow, and rollback plan.",
    ])

    doc.add_heading("19. Suggested Meeting Agenda for Any New Department", level=1)
    add_docx_numbered(doc, [
        "Show the current pain: master schedule, service coverage, call pool, holidays, monthly schedule, requests, and attendance are fragmented.",
        "Ask them to list every schedule they produce and who uses it.",
        "Ask them to identify what is fixed yearly, what changes monthly, and what residents submit.",
        "Walk through one block from master schedule to final published service schedule.",
        "Identify the top five rules that cause errors or chief workload.",
        "Agree on which rules should be automated, which should only warn, and which should stay manual.",
        "Choose MVP scope: one class, one service, one block, one request type, one export format.",
    ])

    doc.add_heading("20. Final Recommendation", level=1)
    add_docx_callout(
        doc,
        "Build the platform around source-of-truth layers",
        "Start with resident profiles and master schedule data. Then build service coverage, call pools, holidays, monthly schedules, requests, and attendance as connected layers. This keeps the system explainable, flexible across specialties, and familiar enough for chiefs who are used to Excel.",
        "purple",
    )

    doc.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("CoverTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=25, leading=29, textColor=colors.HexColor("#24243E"), spaceAfter=10))
    styles.add(ParagraphStyle("CoverSub", parent=styles["Normal"], fontName="Helvetica", fontSize=12, leading=16, textColor=colors.HexColor("#697087"), spaceAfter=14))
    styles.add(ParagraphStyle("H1x", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=15, leading=18, textColor=colors.HexColor("#4F4698"), spaceBefore=10, spaceAfter=7))
    styles.add(ParagraphStyle("H2x", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11.5, leading=14, textColor=colors.HexColor("#24243E"), spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle("Bodyx", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.7, leading=11.3, textColor=colors.HexColor("#22263A"), spaceAfter=5))
    styles.add(ParagraphStyle("Smallx", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.2, leading=9, textColor=colors.HexColor("#697087"), spaceAfter=3))
    styles.add(ParagraphStyle("Cellx", parent=styles["BodyText"], fontName="Helvetica", fontSize=6.7, leading=8.0, textColor=colors.HexColor("#22263A")))
    styles.add(ParagraphStyle("CellBoldx", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=6.8, leading=8.1, textColor=colors.white, alignment=TA_CENTER))
    styles.add(ParagraphStyle("CallTitlex", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=8.7, leading=10, textColor=colors.HexColor("#4F4698"), spaceAfter=3))
    return styles


def hx(value):
    return colors.HexColor(value if str(value).startswith("#") else f"#{value}")


def para(text, style):
    return Paragraph(str(text).replace("&", "&amp;"), style)


def pdf_table(headers, rows, widths, styles, header_fill="#4F4698"):
    data = [[para(h, styles["CellBoldx"]) for h in headers]]
    for row in rows:
        data.append([para(value, styles["Cellx"]) for value in row])
    table = Table(data, colWidths=[w * inch for w in widths], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), hx(header_fill)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, hx(LINE)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, hx("FAFBFE")]),
    ]))
    return table


def pdf_callout(title, body, styles, fill="#EAF8F6", border="#299A8F"):
    table = Table([[para(f"<b>{title}</b><br/>{body}", styles["Bodyx"])]], colWidths=[6.55 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), hx(fill)),
        ("BOX", (0, 0), (-1, -1), 0.6, hx(border)),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    return table


def pdf_image(path, caption, styles):
    path = Path(path)
    if not path.exists():
        return []
    img = Image(str(path))
    max_w = 6.45 * inch
    max_h = 3.45 * inch
    ratio = min(max_w / img.imageWidth, max_h / img.imageHeight)
    img.drawWidth = img.imageWidth * ratio
    img.drawHeight = img.imageHeight * ratio
    return [img, para(caption, styles["Smallx"])]


def create_pdf():
    styles = pdf_styles()
    doc = BaseDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.62 * inch,
        bottomMargin=0.55 * inch,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")

    def footer(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(hx(MUTED))
        canvas.drawString(0.65 * inch, 0.35 * inch, "Clarity Scheduling System Implementation Playbook")
        canvas.drawRightString(7.85 * inch, 0.35 * inch, f"Page {doc_obj.page}")
        canvas.restoreState()

    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=footer)])
    story = []
    story.append(para("Clarity Scheduling System", styles["CoverTitle"]))
    story.append(para("Implementation playbook, data-gathering guide, and product structure for a residency scheduling MVP", styles["CoverSub"]))
    story.append(pdf_callout("Purpose", "A structured guide for IT, chiefs, coordinators, and leadership to gather the correct scheduling data and build a flexible system across departments.", styles))
    story.append(Spacer(1, 10))
    story.append(pdf_table(["Who should use this", "How to use it"], [
        ["IT / product team", "Use as discovery checklist, data model starter, and implementation map."],
        ["Chief residents", "Use to explain scheduling rules, exceptions, fairness goals, and monthly workflows."],
        ["Program leadership", "Use to approve scope and understand operational value."],
    ], [2.0, 4.5], styles, "#299A8F"))
    story.append(PageBreak())

    sections = [
        ("1. Executive Summary", [
            para("The scheduling system should not be treated as one calendar. It is a layered clinical operations system: resident profiles, annual master schedule, service coverage, call pools, holidays, monthly schedules, resident portal, attendance, and analytics.", styles["Bodyx"]),
            pdf_callout("Central design principle", "Every schedule should be generated from structured resident profiles and configurable department rules, then remain editable by chiefs.", styles, "#F4F1FF", "#4F4698"),
            pdf_table(["Step", "Workflow", "Main decision", "Owner"], flow_rows(), [0.38, 1.2, 3.55, 1.05], styles),
        ]),
        ("2. Data Taxonomy", [
            para("Separate fixed rules, resident-submitted preferences, and generated assignments. This prevents the Excel problem where everything is mixed together and hard to audit.", styles["Bodyx"]),
            pdf_table(["Data type", "Nature", "Provided by", "Examples"], DATA_TABLE, [1.25, 1.15, 1.2, 2.8], styles),
        ]),
        ("3. Product Modules", [
            pdf_table(["Module", "Purpose", "Inputs", "Outputs"], MODULE_TABLE, [1.35, 1.55, 1.85, 1.65], styles),
        ] + pdf_image(ASSETS / "chief-overview-main.png", "Chief overview example: readiness, conflicts, and schedule status should be visible without opening multiple files.", styles)),
        ("4. Department Discovery Process", [
            para("Use this sequence with any new department. It keeps discovery organized and reveals rules that teams often forget to mention.", styles["Bodyx"]),
            pdf_table(["Interview area", "Questions to ask"], INTERVIEW_CHECKLIST, [1.55, 4.95], styles, "#299A8F"),
        ]),
        ("5. Master Schedule Builder", [
            para("The annual master schedule is the source of truth. It assigns each resident to one rotation per block and feeds every downstream schedule.", styles["Bodyx"]),
            pdf_table(["Configuration area", "Chiefs provide", "IT builds as options"], [
                ["PGY setup", "PGY classes, resident count, number of master schedules needed", "PGY selector, count matching, missing-profile alerts"],
                ["Rotation catalog", "Core, elective, vacation-eligible, call-eligible rotations", "Editable cards with add, rename, drag, delete, color, service mapping"],
                ["Coverage targets", "Min/max residents per service per block", "Coverage counters, shortage/overage warnings"],
                ["Resident preference intake", "Master ranking, vacation, fellowship timing, elective ranking", "Annual form, timestamps, sortable ranking queue"],
                ["Finalize grid", "Overrides and final chief decisions", "Editable grid, lock/finalize, audit trail"],
            ], [1.45, 2.55, 2.5], styles),
        ] + pdf_image(ASSETS / "chief-master.png", "Master schedule example: final grid remains editable while showing coverage warnings.", styles)),
        ("6. Service Coverage Workbook", [
            pdf_table(["Requirement", "System behavior"], [
                ["Pull names from master", "If a resident is assigned to PICU in Block 2, they appear in the PICU Block 2 service roster."],
                ["Support service aliases", "A master rotation such as Floor can feed Purple, Orange, Gold, or floor senior depending on rules."],
                ["Show counts first", "Keep the main view clean with counts and status; reveal names when clicked."],
                ["Allow Excel-like view", "Chiefs can switch to a familiar spreadsheet-style roster."],
                ["Link to profiles", "Clicking a resident opens profile and returns chief to the same context."],
            ], [1.65, 4.85], styles),
        ] + pdf_image(ASSETS / "service-builder-main.png", "Service builder example: each service needs its own roles, shifts, minimums, and locations.", styles)),
        ("7. Call Pools, Holidays, and Breaks", [
            pdf_table(["Workbook", "Who gets pulled in", "Chief review questions"], [
                ["Call pool", "Elective, call-eligible, or vacation-eligible residents", "Who can do clinic call, jeopardy, cross-cover, no-call, or LOA?"],
                ["Holiday coverage", "Available residents filtered by service needs and requests", "Are all services covered for each holiday?"],
                ["Holiday breaks", "Residents assigned to Christmas/New Year break windows", "Are minimum services covered for both break weeks?"],
            ], [1.25, 2.6, 2.65], styles, "#D1912E"),
            pdf_callout("Keep names available without crowding", "Use card views and counts by default. When chiefs click a cell or card, show names, source rotation, eligibility reason, and profile links.", styles, "#FFF6E4", "#D1912E"),
        ]),
        ("8. Monthly Service Schedule Builder", [
            pdf_table(["Builder step", "What should happen"], [
                ["Block setup", "Select block, service, dates, holidays, locations, and source data."],
                ["People", "Pull residents from master; allow outside rotators and call-pool additions."],
                ["Protected time", "Auto-pull clinic, didactics, ITE, conferences, PTO, leave, and institution rules."],
                ["Requests", "Show approved, pending, and denied requests as reminders."],
                ["Rules", "Apply shift templates, min coverage, eligibility, nights, weekends, and fairness targets."],
                ["Review/generate", "Create an editable draft, flag conflicts, allow drag/drop, then publish."],
            ], [1.35, 5.15], styles),
        ] + pdf_image(ASSETS / "chief-schedule.png", "Monthly schedule editor example: generated templates stay editable and show live checks.", styles)),
        ("9. Resident Portal and Approvals", [
            pdf_table(["Resident feature", "Purpose", "Linked chief-side data"], [
                ["My schedule", "Personal work hours, calls, clinics, didactics, PTO, weekends", "Published schedules and resident profile"],
                ["My master schedule", "Private annual schedule", "Final master schedule"],
                ["Requests", "Days off, weekends, holidays, electives, fellowship notes", "Chief approval queue"],
                ["Call switches", "Offer or accept call switch with eligibility checks", "Call schedule, clinic rules, back-to-back rules"],
                ["PTO/sick leave", "Submit and track leave", "Leave tracker and schedule builder"],
                ["Attendance", "See sessions attended/missed", "Attendance module"],
            ], [1.35, 2.25, 2.9], styles),
        ] + pdf_image(ASSETS / "resident-schedule-main.png", "Resident view example: private personal schedule plus shared published rosters.", styles)),
        ("10. Requests, Attendance, and Statistics", [
            pdf_table(["Area", "Required capability"], [
                ["Requests", "Approval queue with resident context, system checks, and final chief decision."],
                ["Call switches", "Check clinic same day/next day, back-to-back calls, and duplicate assignments."],
                ["Attendance", "Chief creates session; residents check in; profiles update automatically."],
                ["Statistics", "Coverage, fairness, hours, weekends, golden weekends, calls, leave, attendance."],
            ], [1.5, 5.0], styles, "#299A8F"),
        ] + pdf_image(ASSETS / "chief-approvals.png", "Approvals example: chiefs need request context, system checks, and approve/decline controls.", styles)),
        ("11. Ownership and MVP Phases", [
            pdf_table(["Ownership", "Examples"], IT_CHIEF_TABLE, [1.65, 4.85], styles),
            pdf_table(["Phase", "Scope"], FEATURE_TABLE, [1.35, 5.15], styles, "#4F4698"),
        ]),
        ("12. Data Model Starter and Acceptance Tests", [
            pdf_table(["Object", "Key fields"], [
                ["Resident", "id, name, pgy, institution, track, clinic_patterns, didactic_rules, leave_balance, eligibility_flags"],
                ["Block", "id, name, start_date, end_date, year, holidays, request_deadline"],
                ["Rotation", "id, name, pgy_allowed, core_required, vacation_eligible, call_eligible, service_mapping, color"],
                ["Service", "id, name, category, locations, roles, shift_templates, min/max staffing, fairness_rules"],
                ["MasterAssignment", "resident_id, block_id, rotation_id, source, locked, notes, conflicts"],
                ["Request", "resident_id, type, dates/block, priority, submitted_at, status, decision"],
                ["ScheduleAssignment", "resident_id, service_id, block_id, date, role, shift_template, source, status"],
                ["AttendanceSession", "date, block_id, lecture_type, requirement, allowed_mode, attendees, excused_absences"],
            ], [1.45, 5.05], styles),
            pdf_callout("Minimum end-to-end MVP test", "Import a master schedule, generate one service roster, build one monthly service schedule, publish to one resident, submit one call switch, record one attendance session, and export one Excel-like/PDF output.", styles, "#EAF8F6", "#299A8F"),
        ]),
        ("13. Final Recommendation", [
            pdf_callout("Build around source-of-truth layers", "Start with resident profiles and master schedule data. Then build service coverage, call pools, holidays, monthly schedules, requests, and attendance as connected layers. This keeps the system explainable, flexible across specialties, and familiar enough for chiefs moving from Excel.", styles, "#F4F1FF", "#4F4698"),
        ]),
    ]

    for title, parts in sections:
        story.append(para(title, styles["H1x"]))
        for part in parts:
            story.append(part)
            story.append(Spacer(1, 6))

    doc.build(story)


if __name__ == "__main__":
    create_docx()
    create_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
