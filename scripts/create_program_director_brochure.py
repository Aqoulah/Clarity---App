from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "outputs" / "brochure-assets"
OUTPUT = ROOT / "outputs" / "Clarity_Schedule_Program_Director_Brochure.docx"

PURPLE = "50459B"
PURPLE_DARK = "26243A"
PURPLE_LIGHT = "EFECFB"
TEAL = "269889"
TEAL_LIGHT = "E7F6F3"
AMBER = "D79A32"
AMBER_LIGHT = "FFF4DF"
RED = "D75A64"
RED_LIGHT = "FDECEF"
INK = "222536"
MUTED = "676D80"
LINE = "E3E4EA"
WHITE = "FFFFFF"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size=8):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def set_font(run, size=10.5, color=INK, bold=False, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def style_paragraph(paragraph, before=0, after=6, line=1.12, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def add_text(doc, text, size=10.5, color=INK, bold=False, italic=False,
             align=None, before=0, after=6, line=1.12, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    style_paragraph(p, before, after, line, keep)
    set_font(p.add_run(text), size, color, bold, italic)
    return p


def add_heading(doc, text, level=1):
    specs = {
        1: (20, PURPLE_DARK, 8, 8),
        2: (14, PURPLE, 8, 5),
        3: (11, PURPLE_DARK, 5, 3),
    }
    size, color, before, after = specs[level]
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    style_paragraph(p, before, after, 1.05, True)
    set_font(p.add_run(text), size, color, True)
    return p


def add_kicker(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = add_text(doc, text.upper(), 8.5, TEAL, True, align=align, after=4, keep=True)
    for run in p.runs:
        run.font.all_caps = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, after=4, line=1.12)
    set_font(p.add_run(text), 9.5, INK)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])
    set_font(run, 8, MUTED)


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, after=0)
    set_font(p.add_run("CLARITY SCHEDULE  |  PROGRAM CONCEPT"), 7.5, MUTED, True)
    page_p = footer.add_paragraph()
    add_page_number(page_p)


def add_screenshot(doc, filename, caption, width=6.25):
    image = ASSETS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=4, after=3, keep=True)
    p.add_run().add_picture(str(image), width=Inches(width))
    add_text(doc, caption, 7.5, MUTED, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8)


def add_metric_strip(doc, items):
    table = doc.add_table(rows=1, cols=len(items))
    set_table_width(table, [6.5 / len(items)] * len(items))
    for cell, (value, label) in zip(table.rows[0].cells, items):
        set_cell_fill(cell, PURPLE_LIGHT)
        set_cell_border(cell, "DDD8F5", 8)
        set_cell_margins(cell, 130, 120, 130, 120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=1)
        set_font(p.add_run(value), 14, PURPLE, True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p2, after=0)
        set_font(p2.add_run(label), 7.5, MUTED, True)
    return table


def add_card_grid(doc, cards, columns=2):
    rows = (len(cards) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    widths = [6.5 / columns] * columns
    set_table_width(table, widths)
    index = 0
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, 150, 150, 150, 150)
            set_cell_border(cell, LINE, 8)
            if index >= len(cards):
                set_cell_fill(cell, WHITE)
                index += 1
                continue
            title, body, accent = cards[index]
            set_cell_fill(cell, accent[1])
            p = cell.paragraphs[0]
            style_paragraph(p, after=4, keep=True)
            set_font(p.add_run(title), 10.5, accent[0], True)
            p2 = cell.add_paragraph()
            style_paragraph(p2, after=0, line=1.15)
            set_font(p2.add_run(body), 8.5, INK)
            index += 1
    return table


def add_workflow(doc):
    steps = [
        ("1", "Collect", "Residents submit preferences, leave, weekends, and master-schedule rankings."),
        ("2", "Configure", "Chiefs define institution patterns, services, coverage roles, and working hours."),
        ("3", "Generate", "The system pulls eligible residents from the master schedule and builds editable drafts."),
        ("4", "Validate", "Coverage gaps, clinic conflicts, hours, nights, and fairness are flagged before publishing."),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_width(table, [1.625] * 4)
    for cell, (number, title, body) in zip(table.rows[0].cells, steps):
        set_cell_fill(cell, TEAL_LIGHT if number in ("1", "4") else PURPLE_LIGHT)
        set_cell_border(cell, LINE, 8)
        set_cell_margins(cell, 120, 120, 120, 120)
        p = cell.paragraphs[0]
        style_paragraph(p, after=3)
        set_font(p.add_run(number), 15, TEAL if number in ("1", "4") else PURPLE, True)
        p2 = cell.add_paragraph()
        style_paragraph(p2, after=3)
        set_font(p2.add_run(title), 9.5, PURPLE_DARK, True)
        p3 = cell.add_paragraph()
        style_paragraph(p3, after=0, line=1.1)
        set_font(p3.add_run(body), 7.3, MUTED)
    return table


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.68)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.32)
section.footer_distance = Inches(0.32)
add_footer(section)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for level, size, color in ((1, 20, PURPLE_DARK), (2, 14, PURPLE), (3, 11, PURPLE_DARK)):
    style = styles[f"Heading {level}"]
    style.font.name = "Aptos Display"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)

# Page 1: cover
add_kicker(doc, "A smarter operating system for residency scheduling")
add_text(doc, "Clarity Schedule", 31, PURPLE_DARK, True, after=3, keep=True)
add_text(doc, "From resident requests to a fair, editable, publish-ready schedule",
         15, PURPLE, after=12)
add_text(
    doc,
    "A program-wide platform that replaces disconnected forms, spreadsheets, PDFs, and manual cross-checking with one coordinated workflow for residents, chiefs, and departments.",
    11.5, INK, after=14, line=1.25,
)
add_metric_strip(doc, [
    ("2", "connected portals"),
    ("13 or 26", "configurable blocks"),
    ("1", "source of scheduling truth"),
])
add_screenshot(
    doc,
    "resident-overview.png",
    "Resident experience: personal assignments, calls, protected time, requests, and schedule statistics in one place.",
    6.25,
)
add_text(
    doc,
    "Designed for the complexity of graduate medical education - flexible enough for Pediatrics, Internal Medicine, Emergency Medicine, and other training programs.",
    9.5, MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0,
)

# Page 2: problem and solution
page_break(doc)
add_kicker(doc, "The operational problem")
add_heading(doc, "Scheduling is not one spreadsheet", 1)
add_text(
    doc,
    "Every block combines annual rotations, monthly service schedules, protected educational time, individual preferences, leave, coverage requirements, work-hour limits, and last-minute changes. Today, these inputs often live in separate places and must be reconciled manually.",
    10.5, after=10, line=1.2,
)
add_card_grid(doc, [
    ("Fragmented inputs", "Google Forms, emails, PDFs, institutional clinic rules, and annual master schedules are reviewed separately.", (RED, RED_LIGHT)),
    ("Repeated manual work", "Chiefs rebuild the same logic each block and repeatedly check who is eligible, available, or over-assigned.", (AMBER, AMBER_LIGHT)),
    ("Hidden conflicts", "Clinic, didactics, post-call needs, adjacent calls, leave, and next-block assignments are easy to miss.", (RED, RED_LIGHT)),
    ("Limited visibility", "Residents struggle to understand assignments and fairness; leadership lacks a real-time readiness view.", (AMBER, AMBER_LIGHT)),
])
add_heading(doc, "One connected workflow", 2)
add_workflow(doc)
add_screenshot(
    doc,
    "chief-overview.png",
    "Chief overview: each block shows configured services, readiness, coverage status, and items requiring attention.",
    6.15,
)

# Page 3: chief capabilities
page_break(doc)
add_kicker(doc, "Built for chiefs")
add_heading(doc, "Generate faster. Edit freely. Publish with confidence.", 1)
add_text(
    doc,
    "Clarity does not lock chiefs into a rigid algorithm. It prepares a rules-aware draft, explains conflicts, and preserves human judgment through an Excel-like editing experience.",
    10.5, after=9, line=1.2,
)
add_screenshot(
    doc,
    "service-builder.png",
    "Configurable service builder: add any service, define teams or units, staffing roles, eligible PGY levels, and working-hour templates.",
    6.2,
)
add_card_grid(doc, [
    ("Program configuration", "Choose the academic year, number of blocks, block model, request deadlines, holidays, and protected events.", (PURPLE, PURPLE_LIGHT)),
    ("Service-level rules", "Create NICU, PICU, floor teams, consult services, jeopardy, night coverage, or any custom service.", (TEAL, TEAL_LIGHT)),
    ("Master-linked rosters", "Automatically pull residents assigned to a service in that block, while allowing call-pool or manual additions.", (PURPLE, PURPLE_LIGHT)),
    ("Editable schedule workspace", "Drag, move, copy, and revise assignments while coverage, hours, nights, and fairness recalculate.", (TEAL, TEAL_LIGHT)),
    ("Conflict intelligence", "Flag clinic overlap, didactics, leave, back-to-back calls, staffing gaps, and end-of-block night conflicts.", (AMBER, AMBER_LIGHT)),
    ("Readiness and analytics", "Track schedules as pending, draft, ready, or published and compare calls, weekends, nights, and work hours.", (AMBER, AMBER_LIGHT)),
])

# Page 4: resident capabilities
page_break(doc)
add_kicker(doc, "Built for residents")
add_heading(doc, "A clear schedule and a simpler way to request changes", 1)
add_text(
    doc,
    "Residents see only their private annual master schedule, while published service rosters remain visible program-wide. Requests become structured data instead of messages that must be interpreted later.",
    10.5, after=8, line=1.2,
)
add_screenshot(
    doc,
    "resident-overview.png",
    "One resident dashboard combines the monthly schedule, calls, golden weekends, clinics, didactics, leave, and expected hours.",
    5.8,
)
add_card_grid(doc, [
    ("Submit requests", "Vacation, PTO, sick leave, weekends, days off, special needs, fellowship timing, and master-schedule preferences.", (PURPLE, PURPLE_LIGHT)),
    ("Understand assignments", "See daily work hours, service coverage, call type, post-call time, clinic, didactics, conferences, and exams.", (TEAL, TEAL_LIGHT)),
    ("Call-switch marketplace", "Offer an assigned call, list acceptable alternatives, and volunteer for calls offered by peers.", (PURPLE, PURPLE_LIGHT)),
    ("Instant eligibility checks", "The system checks same-day clinic, next-day clinic, duplicate assignments, and back-to-back calls before submission.", (TEAL, TEAL_LIGHT)),
])
add_heading(doc, "Three people, one decision trail", 2)
add_text(
    doc,
    "Offerer -> receiving resident -> chief approval. The chief receives the proposed switch together with the eligibility evidence, so all three parties no longer need to search separate schedules.",
    10, color=PURPLE_DARK, bold=True, after=5,
)
add_text(
    doc,
    "The original assignment remains unchanged until final approval, protecting schedule integrity and creating a documented review process.",
    9.5, color=MUTED, after=0,
)

# Page 5: value and pilot
page_break(doc)
add_kicker(doc, "Why programs choose this approach")
add_heading(doc, "Benefits beyond making the calendar", 1)
benefit_table = doc.add_table(rows=4, cols=3)
set_table_width(benefit_table, [1.3, 2.6, 2.6])
headers = ["Stakeholder", "Immediate benefit", "Program-level impact"]
for cell, text in zip(benefit_table.rows[0].cells, headers):
    set_cell_fill(cell, PURPLE)
    set_cell_border(cell, PURPLE, 8)
    set_cell_margins(cell, 120, 120, 120, 120)
    p = cell.paragraphs[0]
    style_paragraph(p, after=0)
    set_font(p.add_run(text), 8.5, WHITE, True)
rows = [
    ("Chiefs", "Less manual reconciliation; faster drafting; clear conflicts and fairness summaries.", "More time for education, resident support, and higher-value leadership work."),
    ("Residents", "Transparent assignments, easier requests, private master schedules, and safer call switches.", "Improved trust, predictability, communication, and perceived fairness."),
    ("Department", "Consistent rules, documented approvals, readiness tracking, and exportable schedules.", "Better operational resilience, auditability, coverage awareness, and continuity."),
]
for row_cells, values in zip(benefit_table.rows[1:], rows):
    for index, (cell, text) in enumerate(zip(row_cells.cells, values)):
        set_cell_fill(cell, "FAFAFD" if index else PURPLE_LIGHT)
        set_cell_border(cell, LINE, 8)
        set_cell_margins(cell, 130, 130, 130, 130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.15)
        set_font(p.add_run(text), 8.4, PURPLE_DARK if index == 0 else INK, index == 0)

add_heading(doc, "What makes Clarity different", 2)
for item in [
    "Flexible by institution, specialty, service, unit, PGY level, shift type, and block structure.",
    "Human-in-the-loop: the system recommends and validates; chiefs retain final control.",
    "Rules are entered once and reused, while block-specific exceptions remain editable.",
    "Master schedule, monthly service schedules, resident requests, analytics, and approvals stay linked.",
    "Designed around real residency workflows rather than generic employee shift scheduling.",
]:
    add_bullet(doc, item)

add_heading(doc, "A practical pilot", 2)
add_card_grid(doc, [
    ("Phase 1 | Configure", "Load one academic year, resident profiles, institution patterns, and two representative services.", (PURPLE, PURPLE_LIGHT)),
    ("Phase 2 | Parallel test", "Generate one block beside the current process and compare coverage, conflicts, fairness, and editing time.", (TEAL, TEAL_LIGHT)),
    ("Phase 3 | Evaluate", "Collect chief and resident feedback, refine rules, then decide whether to expand program-wide.", (AMBER, AMBER_LIGHT)),
], columns=3)
add_text(
    doc,
    "The goal is simple: transform residency scheduling from a recurring manual project into a transparent, reusable, and safer program workflow.",
    13, PURPLE_DARK, True, align=WD_ALIGN_PARAGRAPH.CENTER, before=15, after=4, line=1.2,
)
add_text(
    doc,
    "Concept developed by Abdullah Abu Aqoulah  |  Clarity Schedule",
    8.5, MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0,
)

doc.core_properties.title = "Clarity Schedule - Program Director Brochure"
doc.core_properties.subject = "Residency scheduling platform concept"
doc.core_properties.author = "Abdullah Abu Aqoulah"
doc.save(OUTPUT)
print(OUTPUT)
