from pathlib import Path
from datetime import date
from textwrap import shorten

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)
DOCX_PATH = OUT / "Clarity_Schedule_Business_and_System_Requirements.docx"
PDF_PATH = OUT / "Clarity_Schedule_Executive_Implementation_Summary.pdf"
ASSETS = OUT / "brochure-v2-assets-anon"

PURPLE = "4F4698"
PURPLE_DARK = "292442"
TEAL = "2F9B8F"
AMBER = "D89C32"
RED = "C55454"
INK = "20263A"
MUTED = "697087"
PALE = "F4F3FB"
PALE_TEAL = "EAF7F5"
PALE_AMBER = "FFF6E7"
LINE = "D9DCE7"
WHITE = "FFFFFF"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Update the table of contents in Word."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])


def configure_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Title", 27, PURPLE_DARK, 0, 8),
        ("Subtitle", 13, MUTED, 0, 14),
        ("Heading 1", 18, PURPLE, 14, 7),
        ("Heading 2", 13.5, PURPLE_DARK, 11, 5),
        ("Heading 3", 11, PURPLE, 8, 3),
    ):
        st = styles[name]
        st.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.font.bold = name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "CLARITY SCHEDULE  |  BUSINESS AND SYSTEM REQUIREMENTS"
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = rgb(MUTED)
    footer = section.footer.paragraphs[0]
    footer.add_run("Implementation baseline  |  Confidential working document     ")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = rgb(MUTED)
    add_page_number(footer)
    return doc


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = rgb(TEAL)
    return p


def add_callout(doc, title, body, tone="purple"):
    fills = {"purple": PALE, "teal": PALE_TEAL, "amber": PALE_AMBER}
    accents = {"purple": PURPLE, "teal": TEAL, "amber": AMBER}
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.72)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fills[tone])
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = rgb(accents[tone])
    r.font.size = Pt(10.5)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            r = p.add_run(item[0] + ": ")
            r.bold = True
            p.add_run(item[1])
        else:
            p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers, rows, widths=None, font_size=8.3, header_fill=PURPLE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(headers) > 3 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(header))
        r.bold = True
        r.font.color.rgb = rgb(WHITE)
        r.font.size = Pt(font_size)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        no_split(table.rows[-1])
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx % 2:
                set_cell_shading(cell, "F8F8FC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(font_size)
    if widths:
        set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, filename, caption, width=6.65):
    path = ASSETS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("title", caption)
    doc_pr.set("descr", caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(6)
    r = cp.add_run(caption)
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = rgb(MUTED)


def add_requirement_section(doc, title, intro, requirements):
    doc.add_heading(title, level=2)
    if intro:
        doc.add_paragraph(intro)
    rows = []
    for req in requirements:
        rows.append((req[0], req[1], req[2], req[3]))
    add_table(doc, ["ID", "Requirement", "Priority", "Acceptance summary"], rows, [0.68, 3.15, 0.7, 2.15], 7.8)


def add_page_break(doc):
    doc.add_page_break()


def build_docx():
    doc = configure_doc()

    add_kicker(doc, "Implementation baseline")
    title = doc.add_paragraph(style="Title")
    title.add_run("Clarity Schedule")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Business Requirements Document and Software Requirements Specification")
    add_callout(
        doc,
        "Purpose",
        "Define the business processes, functional behavior, data, rules, controls, and delivery plan required to move the scheduling concept from an interactive prototype into a production system.",
        "teal",
    )
    doc.add_paragraph()
    add_table(
        doc,
        ["Document attribute", "Value"],
        [
            ("Version", "1.0 implementation baseline"),
            ("Date", date.today().strftime("%B %d, %Y")),
            ("Prepared for", "Residency program leadership, chief residents, implementation team, and institutional stakeholders"),
            ("Product", "Clarity Schedule"),
            ("Status", "Draft for validation and implementation planning"),
            ("Source basis", "Stakeholder requirements, uploaded schedules/instructions, and the current interactive prototype"),
        ],
        [1.65, 5.03],
        8.8,
    )
    doc.add_paragraph()
    add_callout(
        doc,
        "Important governance note",
        "The system may recommend schedules and identify conflicts, but medical education leadership must approve all duty-hour, supervision, leave, staffing, and eligibility policies. Configurable rules shall not be treated as regulatory advice.",
        "amber",
    )
    add_figure(doc, "chief-overview-main.png", "Prototype reference: chief overview and block readiness.")
    add_page_break(doc)

    doc.add_heading("Document Control", level=1)
    add_table(
        doc,
        ["Role", "Responsibility in review"],
        [
            ("Program director", "Approves business objectives, governance, and program policy."),
            ("Chief residents", "Validate operational workflows, service rules, exceptions, and publishing process."),
            ("Program coordinator", "Validates resident records, leave data, deadlines, imports, and communications."),
            ("Residents", "Validate usability, privacy, request submission, and call-switch workflows."),
            ("IT/security", "Approves identity, hosting, security, retention, monitoring, and integration design."),
            ("Development team", "Confirms feasibility, estimates scope, and maintains requirement traceability."),
        ],
        [1.65, 5.03],
    )
    doc.add_heading("Requirement Language", level=2)
    add_bullets(
        doc,
        [
            ("Shall", "mandatory production behavior."),
            ("Should", "important behavior that may be deferred with explicit approval."),
            ("May", "optional enhancement."),
            ("Hard constraint", "an invalid condition that blocks generation, assignment, approval, or publication."),
            ("Soft constraint", "a scored preference or fairness target that produces recommendations and warnings."),
        ],
    )
    doc.add_heading("Table of Contents", level=2)
    add_toc(doc.add_paragraph())
    add_page_break(doc)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Clarity Schedule is a configurable residency scheduling platform with separate chief and resident portals. "
        "It connects annual rotation planning, resident availability, recurring institutional obligations, service-specific staffing rules, "
        "monthly assignment generation, editable schedule workspaces, approvals, publication, workload analytics, and call exchanges."
    )
    add_callout(
        doc,
        "Core product promise",
        "Enter reliable facts once, collect requests in one place, generate explainable schedule drafts, let chiefs edit them, and keep every downstream view synchronized.",
        "purple",
    )
    doc.add_heading("1.1 Business Outcomes", level=2)
    add_bullets(
        doc,
        [
            "Reduce manual collection and reconciliation across forms, spreadsheets, email, PDFs, and messaging.",
            "Decrease time spent identifying eligible residents and checking clinics, didactics, leave, calls, and service coverage.",
            "Improve visibility into staffing gaps, excessive workload, unfair weekend distribution, and unsafe transitions.",
            "Preserve chief authority while giving chiefs decision support and a complete audit trail.",
            "Give residents timely, private, accurate access to their annual and monthly assignments.",
            "Support multiple specialties and institutions without hard-coded team names or shift structures.",
        ],
    )
    doc.add_heading("1.2 Recommended Product Boundaries", level=2)
    add_bullets(
        doc,
        [
            ("In scope", "scheduling workflow, approvals, schedule analytics, configurable rules, imports/exports, notifications, and audit history."),
            ("Not initially in scope", "payroll, credentialing, clinical documentation, patient assignments, billing, and autonomous approval of policy exceptions."),
            ("System posture", "decision-support and workflow automation, not an unsupervised clinical staffing authority."),
        ],
    )

    doc.add_heading("2. Business Problem and Opportunity", level=1)
    doc.add_paragraph(
        "Current scheduling requires chiefs to reconstruct a complete picture of each resident from several disconnected sources. "
        "A change to one request or assignment can require manual review of clinics, didactics, calls, leave, service staffing, weekends, and the next block."
    )
    add_table(
        doc,
        ["Current problem", "Operational effect", "Clarity response"],
        [
            ("Fragmented inputs", "Facts and requests are duplicated or missed.", "One validated resident and institution data model."),
            ("Spreadsheet-based assembly", "Edits are difficult to trace and propagate.", "Linked annual and block workspaces with version history."),
            ("Service rules vary", "Generic scheduling logic produces incorrect assumptions.", "Independent service templates, roles, locations, shifts, and constraints."),
            ("Manual fairness review", "Imbalances are discovered late.", "Real-time hours, calls, nights, weekends, clinics, and golden-weekend analytics."),
            ("Complex call switches", "Three parties repeat the same eligibility review.", "Marketplace, automatic checks, and chief approval workflow."),
            ("Limited readiness view", "Leadership cannot quickly see what is ready or blocked.", "Block-by-service lifecycle dashboard and attention center."),
        ],
        [1.55, 2.4, 2.73],
        8.0,
    )
    doc.add_heading("2.1 Success Measures", level=2)
    add_table(
        doc,
        ["Metric", "Baseline method", "Target for pilot"],
        [
            ("Chief preparation time", "Time study for one block", "At least 40% reduction"),
            ("Unresolved coverage defects at publish", "Count from prior schedules", "At least 50% reduction"),
            ("Resident data completeness", "Required fields complete before generation", "95% or greater"),
            ("Request decision turnaround", "Submission-to-decision duration", "Median under 7 days"),
            ("Manual call-switch review time", "Minutes per proposed exchange", "Under 5 minutes for standard cases"),
            ("Published schedule accuracy", "Corrections after publication", "Improvement over current baseline"),
        ],
        [1.6, 2.5, 2.58],
    )

    doc.add_heading("3. Stakeholders and User Roles", level=1)
    add_table(
        doc,
        ["Role", "Primary goals", "Access level"],
        [
            ("Program director", "Governance, compliance oversight, final escalation.", "Program-wide read; optional approval."),
            ("Chief administrator", "Configure, generate, edit, approve, publish, report.", "Program-wide administrative."),
            ("Chief scheduler", "Build master and block schedules, resolve conflicts.", "Assigned academic year/program."),
            ("Program coordinator", "Maintain profiles, imports, leave, deadlines, communications.", "Administrative without rule override unless granted."),
            ("Resident", "View private schedule, submit requests, offer/accept calls.", "Own private data plus published team schedules."),
            ("Outside rotator", "Provide institutional facts and view assigned schedule.", "Limited resident portal."),
            ("Read-only leader", "Review readiness and reports.", "Read-only program access."),
            ("System administrator", "Tenant setup, access, integrations, support.", "Technical administration; audited."),
        ],
        [1.28, 3.15, 2.25],
        7.9,
    )
    doc.add_heading("3.1 Privacy Principle", level=2)
    doc.add_paragraph(
        "Residents shall see their own private master schedule, requests, leave records, and personal analytics. "
        "They may see published service rosters needed for operational coverage, but shall not see another resident's private requests, preference rankings, protected circumstances, or unpublished annual plan."
    )

    doc.add_heading("4. Product Scope and Conceptual Architecture", level=1)
    add_numbered(
        doc,
        [
            "Program configuration defines the academic year, block model, services, institutions, roles, shifts, and rules.",
            "Resident profiles hold fixed facts, recurring patterns, eligibility, annual preferences, and block-specific events.",
            "Chiefs review and decide requests before approved requests become scheduling inputs.",
            "The annual master schedule assigns each resident to one rotation per block and checks capacity and transitions.",
            "Each service schedule pulls its eligible roster from the selected block of the master schedule.",
            "The scheduling engine creates a draft using hard constraints first and optimization goals second.",
            "Chiefs edit assignments in an Excel-like workspace with live coverage, fairness, and conflict feedback.",
            "Approved schedules are published to residents, analytics, reports, and call-switch eligibility.",
        ],
    )
    add_figure(doc, "chief-schedule.png", "Prototype reference: editable monthly service schedule with coverage and fairness insights.")

    doc.add_heading("5. Business Process Requirements", level=1)
    doc.add_heading("5.1 Annual Setup", level=2)
    add_numbered(
        doc,
        [
            "Create an academic year and select the number and dates of blocks.",
            "Define master rotations and annual capacity requirements.",
            "Define service schedules and map each service to one or more master rotations.",
            "Configure institutions, recurring clinic/didactic patterns, travel buffers, and role eligibility.",
            "Import or create residents and validate required fixed facts.",
            "Open annual preference submission and communicate deadlines.",
            "Review conflicts, approve inputs, and build the annual master schedule.",
        ],
    )
    doc.add_heading("5.2 Block Scheduling", level=2)
    add_numbered(
        doc,
        [
            "Select the block and review service readiness.",
            "Import or verify holidays, exams, conferences, leave, and one-time exceptions.",
            "Review and decide resident requests.",
            "Generate one draft for each configured service using its own roster and rules.",
            "Resolve staffing, protected-time, fairness, rest, and cross-block alerts.",
            "Edit, save versions, compare, and mark each service ready.",
            "Publish selected service schedules or the complete block.",
        ],
    )
    doc.add_heading("5.3 Call Switch", level=2)
    add_numbered(
        doc,
        [
            "A resident selects one of their eligible published calls and offers it for exchange.",
            "The resident optionally identifies acceptable replacement dates and adds a note.",
            "Other residents see the offer and receive an immediate eligibility result.",
            "A volunteer submits interest; the app rechecks both schedules and routes the proposal to chiefs.",
            "A chief approves, declines, or requests correction.",
            "Approval creates a controlled schedule revision, recalculates analytics, notifies both residents, and preserves the audit history.",
        ],
    )
    add_figure(doc, "call-switch-anon.png", "Prototype reference: chief review of a call switch with precomputed eligibility checks.")

    doc.add_heading("6. Functional Requirements", level=1)
    add_requirement_section(
        doc,
        "6.1 Identity, Access, and Program Administration",
        "The production system shall support multiple programs and enforce role- and record-level access.",
        [
            ("FR-IAM-001", "Authenticate users through approved institutional identity or secure local accounts.", "Must", "Authorized user reaches only assigned programs."),
            ("FR-IAM-002", "Enforce role-based permissions and private resident record boundaries.", "Must", "Resident cannot query another resident's private record."),
            ("FR-IAM-003", "Support delegated chief, coordinator, director, and read-only roles.", "Must", "Permission matrix passes automated tests."),
            ("FR-IAM-004", "Record login, export, approval, publication, rule change, and override events.", "Must", "Audit report identifies actor, time, and before/after state."),
            ("FR-IAM-005", "Support one tenant with multiple programs and academic years.", "Should", "Data remains segregated by tenant and program."),
        ],
    )
    add_requirement_section(
        doc,
        "6.2 Academic Year and Block Configuration",
        "",
        [
            ("FR-BLK-001", "Create 13-block, 26-block, calendar-month, or custom-date academic structures.", "Must", "Blocks cover approved dates without unintended overlap."),
            ("FR-BLK-002", "Define request deadlines as a configurable number of days before block start.", "Must", "Late submissions are labeled and routed correctly."),
            ("FR-BLK-003", "Store holidays, retreats, exams, conferences, and program-wide protected events.", "Must", "Events appear in applicable block builders."),
            ("FR-BLK-004", "Display lifecycle by block and service: setup, pending, draft, review, ready, published, archived.", "Must", "Overview updates after every lifecycle change."),
        ],
    )
    add_requirement_section(
        doc,
        "6.3 Resident and Institution Profiles",
        "",
        [
            ("FR-PRO-001", "Maintain resident identity, PGY level, institution, program, status, and contact fields.", "Must", "Required facts validate before activation."),
            ("FR-PRO-002", "Maintain block-by-block clinic patterns and recurring didactics.", "Must", "Protected events populate generated drafts."),
            ("FR-PRO-003", "Support institution profiles that automatically apply to linked residents.", "Must", "Institution changes affect future drafts, not published history."),
            ("FR-PRO-004", "Support individual overrides without changing the institution template.", "Must", "Override scope and source are visible."),
            ("FR-PRO-005", "Store service/role eligibility, call-pool eligibility, outside-rotator limitations, and supervision category.", "Must", "Ineligible assignments are blocked or explicitly overridden."),
            ("FR-PRO-006", "Import residents and recurring patterns from approved spreadsheet templates.", "Should", "Import preview identifies errors before commit."),
        ],
    )
    add_figure(doc, "institution-rules.png", "Prototype reference: reusable institution profile with recurring protected time and eligibility.")
    add_requirement_section(
        doc,
        "6.4 Service, Location, Role, and Shift Configuration",
        "",
        [
            ("FR-SVC-001", "Create, rename, reorder, deactivate, and delete configurable services.", "Must", "All linked views update consistently."),
            ("FR-SVC-002", "Map a service to one or more master rotations.", "Must", "Block roster is pulled from the configured source."),
            ("FR-SVC-003", "Define coverage locations such as unit, floor, team, consult, or responsibility.", "Must", "Assignments identify where coverage occurs."),
            ("FR-SVC-004", "Define required roles and counts by PGY, seniority, credential, institution, or custom tag.", "Must", "Coverage checker evaluates each role separately."),
            ("FR-SVC-005", "Define custom assignments such as day, night, long call, short call, procedure call, crossover, and backup.", "Must", "Chiefs can add new types without code changes."),
            ("FR-SVC-006", "Configure times, colors, display style, hour treatment, recurrence, and post-call behavior per assignment type.", "Must", "Hours and visual cards follow saved configuration."),
            ("FR-SVC-007", "Version service templates and permit block-specific overrides.", "Must", "Historical schedules retain the version used."),
        ],
    )
    add_figure(doc, "service-builder-main.png", "Prototype reference: independent service template, staffing roles, coverage locations, and shift library.")
    add_requirement_section(
        doc,
        "6.5 Resident Requests and Chief Decisions",
        "",
        [
            ("FR-REQ-001", "Accept weekend, golden-weekend, day-off, call restriction, vacation, elective, fellowship, PTO, sick leave, and special-circumstance requests.", "Must", "Each request has type, scope, dates, reason category, and status."),
            ("FR-REQ-002", "Separate fixed facts from preferences and requests.", "Must", "Facts do not require repeated monthly entry."),
            ("FR-REQ-003", "Show deadlines and clearly label late submissions.", "Must", "Late request requires explicit handling."),
            ("FR-REQ-004", "Allow chiefs to approve, decline, partially approve, request clarification, or reopen.", "Must", "Decision, reason, actor, and timestamp are stored."),
            ("FR-REQ-005", "Use only approved requests as scheduling inputs.", "Must", "Pending/declined requests do not constrain generation."),
            ("FR-REQ-006", "Identify competing requests and show advisory priority factors without making the final decision.", "Must", "Chief retains decision authority."),
            ("FR-REQ-007", "Restrict sensitive circumstance details to authorized roles.", "Must", "Residents and unrelated staff cannot access protected notes."),
        ],
    )
    add_requirement_section(
        doc,
        "6.6 Annual Master Schedule",
        "",
        [
            ("FR-MST-001", "Display residents by row and blocks by column with editable rotation assignments.", "Must", "Each active resident has one primary assignment per block."),
            ("FR-MST-002", "Configure rotation capacities, PGY mix, annual requirements, vacation eligibility, and outside-rotator eligibility.", "Must", "Validation uses saved rotation policy."),
            ("FR-MST-003", "Display resident rankings, approved preferences, fellowship timing, and leave constraints.", "Must", "Chief can compare assignment to preference."),
            ("FR-MST-004", "Flag over/under capacity, missing annual requirements, duplicate assignments, and unsafe transitions.", "Must", "Conflict inspector links to affected cells."),
            ("FR-MST-005", "Support locking assignments during optimization.", "Must", "Locked assignments remain unchanged."),
            ("FR-MST-006", "Prevent or warn about an inpatient block after a final night assignment based on configured policy.", "Must", "Transition check spans adjacent blocks."),
            ("FR-MST-007", "Save versions, compare changes, restore an approved version, and identify editor.", "Must", "Version history is complete."),
        ],
    )
    add_figure(doc, "chief-master.png", "Prototype reference: editable annual master schedule with live capacity and conflict inspection.")
    add_requirement_section(
        doc,
        "6.7 Schedule Generation and Editing",
        "",
        [
            ("FR-GEN-001", "Generate a separate draft for every active service in the selected block.", "Must", "Each draft uses the correct service template."),
            ("FR-GEN-002", "Pull the initial roster from the master schedule, then identify eligible supplemental/call-pool residents.", "Must", "Roster sources are labeled."),
            ("FR-GEN-003", "Apply approved protected time, leave, institutional patterns, and one-time events before optimization.", "Must", "Generated assignments do not silently overwrite protected time."),
            ("FR-GEN-004", "Satisfy hard constraints before optimizing preferences and fairness.", "Must", "Unsatisfied hard constraint blocks ready/publish unless authorized override exists."),
            ("FR-GEN-005", "Explain why each resident was selected and why a constraint is unmet.", "Must", "Chief can inspect source rule and affected records."),
            ("FR-GEN-006", "Provide a four-week/block editor with drag/drop, copy, move, clear, custom time, resident add/remove, and row reorder.", "Must", "Edits recalculate immediately."),
            ("FR-GEN-007", "Support undo/redo, autosave, manual save, named versions, and chief notes.", "Must", "Editing session can recover from interruption."),
            ("FR-GEN-008", "Recalculate daily coverage, hours, calls, nights, weekend work, golden weekends, clinic conflicts, and rest after every edit.", "Must", "Insights update within accepted response time."),
            ("FR-GEN-009", "Allow an authorized chief to override a warning or permitted constraint with reason.", "Must", "Override is visible and audited."),
        ],
    )
    add_requirement_section(
        doc,
        "6.8 Publishing and Resident Experience",
        "",
        [
            ("FR-PUB-001", "Support draft, ready, published, corrected, and archived schedule states.", "Must", "Only published schedules appear to residents."),
            ("FR-PUB-002", "Allow selective service publication and controlled block publication.", "Must", "Pending services remain clearly identified."),
            ("FR-PUB-003", "Show resident personal monthly assignments, work hours, calls, protected activities, rest, and days off.", "Must", "Resident view matches published source."),
            ("FR-PUB-004", "Show each resident only their private annual master schedule.", "Must", "Privacy test prevents cross-resident access."),
            ("FR-PUB-005", "Show read-only published service schedules with names, roles, location, responsibility, date, and hours.", "Must", "Operational roster is complete."),
            ("FR-PUB-006", "Notify affected users when a published assignment changes.", "Must", "Notification identifies what changed and effective time."),
            ("FR-PUB-007", "Export approved schedules to PDF and spreadsheet with publication/version metadata.", "Must", "Export matches on-screen schedule."),
        ],
    )
    add_figure(doc, "resident-schedule-main.png", "Prototype reference: private resident monthly schedule, calls, protected time, and days off.")
    add_requirement_section(
        doc,
        "6.9 Call Switches",
        "",
        [
            ("FR-SWX-001", "Allow residents to offer only their own eligible published call assignments.", "Must", "Non-call or unpublished assignments cannot be offered."),
            ("FR-SWX-002", "Display open offers to eligible residents without exposing private profile details.", "Must", "Marketplace reveals only operationally necessary data."),
            ("FR-SWX-003", "Check clinic on call date, clinic on following day, back-to-back calls, duplicate assignment, leave, service eligibility, rest, and configured duty rules.", "Must", "Result lists pass/fail reason for every check."),
            ("FR-SWX-004", "Recheck eligibility when a volunteer submits and when a chief decides.", "Must", "Stale schedule changes cannot bypass validation."),
            ("FR-SWX-005", "Require chief approval before changing the published schedule.", "Must", "Resident actions alone do not mutate assignments."),
            ("FR-SWX-006", "Apply approved switch atomically and retain original assignment history.", "Must", "Both sides update together or not at all."),
        ],
    )
    add_requirement_section(
        doc,
        "6.10 Analytics, Overview, and Reporting",
        "",
        [
            ("FR-ANA-001", "Show block-by-service readiness and pending attention items.", "Must", "New, renamed, and deleted services update overview."),
            ("FR-ANA-002", "Calculate resident hours, days, nights, calls by type, clinic sessions, weekend days, days off, and golden weekends.", "Must", "Values derive from assignment templates."),
            ("FR-ANA-003", "Compare fairness only among residents in equivalent roles or configured peer groups.", "Must", "Comparison basis is visible."),
            ("FR-ANA-004", "Show staffing count and role coverage for each service and date.", "Must", "Shortage and excess are distinguishable."),
            ("FR-ANA-005", "Provide resident-, service-, block-, and academic-year reports.", "Should", "Filters and export work consistently."),
            ("FR-ANA-006", "Distinguish warnings from policy violations and from informational suggestions.", "Must", "Severity and required action are clear."),
        ],
    )

    doc.add_heading("7. Business Rules and Constraint Model", level=1)
    add_callout(
        doc,
        "Rule hierarchy",
        "1) legal/program safety and approved leave; 2) staffing and eligibility; 3) protected educational/clinical time; 4) approved requests; 5) fairness and preferences. The exact hierarchy must be approved per program.",
        "amber",
    )
    add_table(
        doc,
        ["ID", "Rule", "Default classification", "Configuration scope"],
        [
            ("BR-001", "One primary master rotation per resident per block.", "Hard", "Program"),
            ("BR-002", "Only eligible PGY/institution/credential may fill a role.", "Hard", "Service role"),
            ("BR-003", "Meet minimum daily staffing by location and role.", "Hard or overrideable", "Service/block"),
            ("BR-004", "Approved leave blocks work assignments.", "Hard", "Resident/event"),
            ("BR-005", "Clinic and didactics are protected according to effect.", "Hard or warning", "Institution/resident"),
            ("BR-006", "Maximum consecutive nights.", "Hard or warning", "Service"),
            ("BR-007", "Post-call recovery after configured night/call.", "Hard", "Assignment type"),
            ("BR-008", "No back-to-back calls unless policy explicitly permits.", "Hard", "Program/call type"),
            ("BR-009", "No call on clinic day or before clinic when configured.", "Hard", "Program/institution"),
            ("BR-010", "Target weekend days and at least one golden weekend.", "Soft unless policy says otherwise", "Service/block"),
            ("BR-011", "Balance nights, calls, hours, and clinic burden among peers.", "Soft", "Service/peer group"),
            ("BR-012", "End-of-block night to next inpatient block transition.", "Hard or warning", "Program/rotation"),
            ("BR-013", "Only approved resident requests influence generation.", "Hard workflow", "Program"),
            ("BR-014", "Chief override requires authority, reason, and audit entry.", "Hard workflow", "Tenant"),
            ("BR-015", "Published changes create a new immutable schedule version.", "Hard workflow", "System"),
        ],
        [0.62, 3.18, 1.45, 1.43],
        7.5,
    )
    doc.add_heading("7.1 Scheduling Objective Function", level=2)
    doc.add_paragraph(
        "The engine should use constraint optimization rather than a single opaque score. Hard constraints must be proven satisfied or returned as explicit infeasibility. "
        "Soft objectives should be weighted, independently reportable, and configurable."
    )
    add_bullets(
        doc,
        [
            "Maximize required coverage and valid role mix.",
            "Minimize hard-rule violations and unauthorized overrides.",
            "Honor approved leave and protected time.",
            "Maximize approved high-priority requests within capacity.",
            "Minimize variance in nights, weekend days, calls, and hours among peer groups.",
            "Maximize golden-weekend attainment.",
            "Minimize disruptive transitions and changes to locked/published assignments.",
            "Produce an explanation and score breakdown for each generated draft.",
        ],
    )

    doc.add_heading("8. Conceptual Data Model", level=1)
    add_table(
        doc,
        ["Entity", "Key information", "Important relationships"],
        [
            ("Tenant / Program", "Name, specialty, timezone, policy set", "Owns users, years, institutions, services"),
            ("Academic Year / Block", "Dates, deadline, lifecycle", "Contains master and service schedules"),
            ("User / Role", "Identity, role, status", "Linked to resident/staff profile"),
            ("Resident", "PGY, institution, status, tags", "Has patterns, requests, assignments"),
            ("Institution Profile", "Protected time, buffers, eligibility", "Applied to linked residents"),
            ("Rotation", "Capacity, inpatient flag, requirements", "Assigned in master schedule"),
            ("Service Template", "Master links, locations, roles, rules", "Instantiates service schedule"),
            ("Assignment Type", "Time, hours, category, color, recovery", "Used in schedule cells"),
            ("Master Assignment", "Resident, block, rotation, lock", "Sources service roster"),
            ("Request", "Type, dates, priority facts, status", "Decision controls engine use"),
            ("Protected Event", "Scope, time, effect, source", "Blocks or warns assignment"),
            ("Service Schedule Version", "Status, template version, author", "Contains assignments"),
            ("Daily Assignment", "Resident, service, location, role, time", "Feeds coverage and analytics"),
            ("Call Switch", "Offer, volunteer, checks, decision", "Creates schedule revision"),
            ("Audit Event", "Actor, action, before/after, reason", "Attached to every controlled change"),
            ("Notification", "Recipient, event, delivery status", "Triggered by workflow changes"),
        ],
        [1.48, 2.62, 2.58],
        7.5,
    )
    doc.add_heading("8.1 Data Provenance", level=2)
    doc.add_paragraph(
        "Every scheduling input shall retain its source: resident submission, institution template, chief entry, spreadsheet import, integration, or generated recommendation. "
        "The UI shall distinguish inherited data from individual overrides and published history."
    )

    doc.add_heading("9. Permissions Matrix", level=1)
    add_table(
        doc,
        ["Capability", "Resident", "Chief", "Coordinator", "Director", "Read-only"],
        [
            ("View own private schedule", "Yes", "Yes", "Yes", "Yes", "No"),
            ("View published team schedule", "Yes", "Yes", "Yes", "Yes", "Yes"),
            ("Submit own request/leave", "Yes", "On behalf", "On behalf", "No", "No"),
            ("Review private request detail", "Own", "Yes", "Yes", "Escalated", "No"),
            ("Approve requests/switches", "No", "Yes", "Configurable", "Configurable", "No"),
            ("Edit master schedule", "No", "Yes", "Configurable", "Configurable", "No"),
            ("Configure services/rules", "No", "Yes", "Limited", "Approve", "No"),
            ("Generate/edit schedules", "No", "Yes", "Configurable", "No", "No"),
            ("Publish/correct schedules", "No", "Yes", "Configurable", "Configurable", "No"),
            ("Export reports", "Own", "Yes", "Yes", "Yes", "Limited"),
        ],
        [1.85, 0.72, 0.72, 0.92, 0.82, 0.85],
        7.1,
    )

    doc.add_heading("10. Interface and Navigation Requirements", level=1)
    doc.add_heading("10.1 Chief Portal", level=2)
    add_bullets(
        doc,
        [
            ("Overview", "block/service readiness, coverage score, workload summary, conflicts, request decisions, and attention center."),
            ("Schedule builder", "program setup, people, protected time, requests, service templates, validation, and generation."),
            ("Schedules", "block and service tabs, four-week editor, insights, versions, export, and publication."),
            ("Analytics", "resident and service workload/fairness with filters and exports."),
            ("Master schedule", "editable annual grid, capacity, preference score, conflict inspector, and locks."),
            ("Residents and approvals", "profiles, recurring patterns, events, requests, leave, eligibility, and call switches."),
            ("Institution rules", "teams/tabs, shift templates, institution profiles, service requirements, and global rules."),
        ],
    )
    doc.add_heading("10.2 Resident Portal", level=2)
    add_bullets(
        doc,
        [
            ("My overview", "next assignments, key workload indicators, calls, protected events, and pending actions."),
            ("My schedule", "private block calendar with work, calls, clinic, didactics, leave, and rest."),
            ("My master schedule", "private annual rotations and submitted annual preferences."),
            ("Published schedules", "read-only shared service schedules."),
            ("Requests and approvals", "submission status, decision, deadline, and chief comments."),
            ("Call switches", "own offers, open marketplace, eligibility result, and approval status."),
            ("PTO and sick leave", "planned and urgent leave workflow with privacy controls."),
        ],
    )
    add_figure(doc, "resident-nav.png", "Prototype reference: resident portal navigation.")

    doc.add_heading("11. Integration and Import/Export Requirements", level=1)
    add_table(
        doc,
        ["Interface", "Purpose", "Phase"],
        [
            ("Spreadsheet import/export", "Residents, master schedule, service schedules, rules, and analytics", "MVP"),
            ("PDF export", "Publication and archival copies", "MVP"),
            ("Email notifications", "Deadlines, decisions, publication, corrections", "MVP"),
            ("Institutional SSO", "Authentication and lifecycle management", "Pilot/production"),
            ("Calendar feed", "Personal published assignments and protected events", "Phase 2"),
            ("Form migration/API", "Replace or ingest existing Google/Microsoft forms", "Phase 2"),
            ("Residency management system", "Resident demographics and rotation synchronization", "Phase 2/3"),
            ("Messaging integration", "Optional Teams/Slack notifications", "Phase 3"),
        ],
        [1.65, 3.55, 1.48],
        8.0,
    )
    doc.add_heading("11.1 Import Controls", level=2)
    add_bullets(
        doc,
        [
            "Provide downloadable templates and field definitions.",
            "Preview additions, updates, duplicates, and validation errors.",
            "Require confirmation before commit.",
            "Record source file, importer, timestamp, and row-level result.",
            "Make imports idempotent using stable external identifiers.",
            "Never silently replace published schedule history.",
        ],
    )

    doc.add_heading("12. Nonfunctional Requirements", level=1)
    add_table(
        doc,
        ["Category", "Requirement"],
        [
            ("Security", "Encryption in transit and at rest; least privilege; secure session handling; secrets management; dependency scanning."),
            ("Privacy", "Record-level controls, minimum necessary display, sensitive-note restrictions, retention policy, and audited exports."),
            ("Availability", "Target service availability and maintenance window shall be agreed before production."),
            ("Performance", "Normal page response under 2 seconds; edit recalculation under 1 second where feasible; generation progress visible."),
            ("Scalability", "Support multiple programs, 100+ residents per program, configurable blocks, services, and multi-year history."),
            ("Reliability", "Atomic approvals/publications, backups, recovery plan, optimistic locking, and conflict-safe concurrent editing."),
            ("Auditability", "Immutable history for policy, approval, assignment, publication, export, and override events."),
            ("Accessibility", "WCAG 2.2 AA target, keyboard navigation, color-independent status, readable zoom, and labeled controls."),
            ("Usability", "Desktop-first scheduling workspace with responsive resident portal and clear error recovery."),
            ("Localization", "Program timezone, locale-aware dates/times, and daylight-saving handling."),
            ("Maintainability", "Modular rule engine, versioned API, automated tests, documented migrations, and observable jobs."),
            ("Compliance", "Institutional legal, privacy, records, labor, accreditation, and security review before production."),
        ],
        [1.35, 5.33],
        8.0,
    )
    add_callout(
        doc,
        "Sensitive information",
        "Leave and special-circumstance requests may contain sensitive employment or health-related information. The implementation must minimize collected detail and complete institutional privacy and legal review. HIPAA applicability must be determined by the institution; it should not be assumed.",
        "amber",
    )

    doc.add_heading("13. Technical Implementation Guidance", level=1)
    doc.add_heading("13.1 Suggested Logical Components", level=2)
    add_bullets(
        doc,
        [
            "Web client for chief and resident portals.",
            "Application API with tenant, program, and record-level authorization.",
            "Relational database for normalized scheduling and audit data.",
            "Rules service that evaluates constraints and produces explanations.",
            "Optimization worker for long-running schedule generation.",
            "Notification worker for email and future messaging/calendar channels.",
            "Object storage for imports, exports, and generated documents.",
            "Audit and observability pipeline for operational and security events.",
        ],
    )
    doc.add_heading("13.2 API Domains", level=2)
    add_table(
        doc,
        ["Domain", "Representative operations"],
        [
            ("Programs and years", "Create year, configure blocks, set deadline, archive"),
            ("People and institutions", "Import resident, edit profile, set recurring pattern, override"),
            ("Services and rules", "Create service, version template, define role/location/shift/rule"),
            ("Requests", "Submit, review, decide, reopen, list conflicts"),
            ("Master schedule", "Read grid, update cell, lock, validate, optimize, version"),
            ("Service schedules", "Generate, edit assignment, validate, mark ready, publish, correct"),
            ("Call switches", "Offer, volunteer, evaluate, approve, apply"),
            ("Analytics", "Resident totals, service coverage, fairness, readiness, export"),
            ("Audit and notifications", "History, delivery status, acknowledgement"),
        ],
        [1.75, 4.93],
    )
    doc.add_heading("13.3 Concurrency and Versioning", level=2)
    add_bullets(
        doc,
        [
            "Every editable master and service schedule shall have a version identifier.",
            "The API shall reject stale writes or offer an explicit merge workflow.",
            "Publication shall create an immutable snapshot.",
            "Corrections shall create a later published version with affected-user notifications.",
            "Template changes shall not retroactively alter published snapshots.",
        ],
    )

    doc.add_heading("14. Validation and Acceptance Strategy", level=1)
    add_table(
        doc,
        ["Test layer", "Required coverage"],
        [
            ("Unit", "Time calculations, overnight hours, recurrence, rule evaluation, permission checks"),
            ("Constraint", "Each hard/soft rule with positive, negative, boundary, and exception cases"),
            ("Integration", "Imports, approvals, master-to-service roster, publication, notifications"),
            ("End-to-end", "Annual setup through resident publication and call-switch correction"),
            ("Security", "Authorization bypass, private record access, export controls, audit integrity"),
            ("Performance", "Generation, large master grid, concurrent editors, analytics recalculation"),
            ("Accessibility", "Keyboard, screen reader, contrast, zoom, error identification"),
            ("User acceptance", "Chief, coordinator, director, resident, and outside rotator scenarios"),
        ],
        [1.32, 5.36],
    )
    doc.add_heading("14.1 Minimum MVP Acceptance Scenarios", level=2)
    add_numbered(
        doc,
        [
            "Configure a new program with custom block dates, services, roles, and shifts without code changes.",
            "Import synthetic residents and institution patterns with a validation preview.",
            "Collect resident requests and ensure only approved requests influence generation.",
            "Build and validate a master schedule with capacity and transition warnings.",
            "Generate at least three service schedules with different staffing and shift models.",
            "Edit assignments and observe immediate coverage, hours, night, weekend, and clinic recalculation.",
            "Publish schedules and verify private versus shared resident access.",
            "Complete a call switch from offer through chief approval and schedule correction.",
            "Export a published schedule and resident/service analytics.",
            "Retrieve a complete audit history for the above workflow.",
        ],
    )

    doc.add_heading("15. Delivery Roadmap", level=1)
    add_table(
        doc,
        ["Phase", "Primary deliverables", "Exit criteria"],
        [
            ("0. Discovery and policy validation", "Rule workshops, data dictionary, security review, workflow sign-off", "Approved BRD/SRS and prioritized backlog"),
            ("1. Foundation", "Identity, tenant/program model, residents, institutions, blocks, services, shifts, imports", "Validated configuration and synthetic dataset"),
            ("2. Master schedule", "Annual grid, capacities, preferences, validation, versions", "Chiefs can build and save a valid annual plan"),
            ("3. Block scheduling MVP", "Roster pull, generator, editor, coverage, analytics, export", "Pilot block can be generated and edited"),
            ("4. Resident portal and approvals", "Private schedules, requests, leave, published rosters, notifications", "Residents complete pilot workflows"),
            ("5. Call switches and controlled corrections", "Marketplace, eligibility, approval, schedule revision", "End-to-end switch passes UAT"),
            ("6. Production hardening", "SSO, monitoring, backups, performance, accessibility, migration", "Institutional release approval"),
        ],
        [1.55, 3.35, 1.78],
        7.6,
    )
    doc.add_heading("15.1 Recommended MVP Boundary", level=2)
    add_bullets(
        doc,
        [
            "One institution/tenant and one residency program.",
            "One academic year with configurable blocks.",
            "Configurable services, roles, locations, and shift types.",
            "Resident/institution profiles and spreadsheet import.",
            "Request approvals, master schedule, block generation, editor, analytics, publication, and exports.",
            "Call-switch workflow may launch immediately after the core publishing model is stable.",
            "Email notifications before calendar or residency-system integration.",
        ],
    )

    doc.add_heading("16. Risks and Mitigations", level=1)
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ("Unwritten scheduling rules", "Incorrect or rejected schedules", "Rule workshops, examples, owner, effective dates, tests"),
            ("Conflicting policies", "Infeasible generation", "Priority hierarchy and explicit infeasibility report"),
            ("Poor source data", "Unsafe recommendations", "Completeness gates, validation, provenance, import preview"),
            ("Over-automation", "Loss of trust and unsafe assumptions", "Chief approval, explanations, overrides, staged rollout"),
            ("Sensitive leave details", "Privacy exposure", "Data minimization, restricted fields, retention and access review"),
            ("Concurrent spreadsheet use", "Divergent source of truth", "Migration plan, freeze windows, controlled exports"),
            ("Rule/template changes", "Historical inconsistency", "Version templates and immutable published snapshots"),
            ("Algorithm perceived as unfair", "Low adoption", "Peer-group definitions, visible metrics, configurable weights, decision logs"),
            ("Scope expansion across specialties", "Delayed MVP", "Configurable foundation with one-program pilot"),
        ],
        [1.65, 2.0, 3.03],
        7.7,
    )

    doc.add_heading("17. Open Decisions Requiring Stakeholder Approval", level=1)
    add_table(
        doc,
        ["ID", "Decision", "Recommended owner"],
        [
            ("OD-001", "Which rules are absolute, overrideable, or advisory for each service?", "Program director + chiefs"),
            ("OD-002", "Exact peer groups and fairness weights for hours, calls, nights, weekends, and clinics?", "Chiefs"),
            ("OD-003", "Duty-hour policy source and calculation window?", "Program leadership/compliance"),
            ("OD-004", "Who may approve PTO, sick leave, special circumstances, and call switches?", "Program director"),
            ("OD-005", "What resident information is visible in published rosters?", "Leadership/privacy"),
            ("OD-006", "What data may be included in special-circumstance requests?", "HR/legal/privacy"),
            ("OD-007", "SSO provider, hosting environment, retention, backup, and recovery targets?", "Institutional IT"),
            ("OD-008", "Which current forms and spreadsheets are authoritative during migration?", "Coordinator/chiefs"),
            ("OD-009", "When is a block considered complete and who may publish/correct it?", "Chiefs/director"),
            ("OD-010", "Which integrations are required for pilot versus later phases?", "Product sponsor/IT"),
        ],
        [0.68, 4.45, 1.55],
        7.6,
    )

    doc.add_heading("18. Definition of Ready for Development", level=1)
    add_bullets(
        doc,
        [
            "Program sponsor and operational product owner are named.",
            "MVP scope and pilot program are approved.",
            "Service templates and rule classifications are validated with real examples.",
            "Data dictionary and import templates are approved.",
            "Privacy, security, and records requirements are documented.",
            "Acceptance scenarios are converted into testable backlog items.",
            "Unresolved decisions are assigned owners and due dates.",
            "Synthetic test data is available before real resident data is loaded.",
        ],
    )
    add_callout(
        doc,
        "Implementation recommendation",
        "Begin with discovery workshops and a production data model, not by extending the current front-end prototype as the database. The prototype is a valuable interaction reference; production requires authenticated APIs, persistent versioned data, a validated rule engine, and audited workflows.",
        "teal",
    )

    doc.add_heading("Appendix A. Glossary", level=1)
    add_table(
        doc,
        ["Term", "Definition"],
        [
            ("Block", "A defined scheduling period in the academic year."),
            ("Master schedule", "Annual assignment of each resident to a rotation for every block."),
            ("Service schedule", "Daily assignment schedule for one team/service within one block."),
            ("Golden weekend", "A Saturday and Sunday off together, subject to program definition."),
            ("Call pool / Jeopardy", "Residents eligible to provide backup or call coverage based on rotation and policy."),
            ("Protected time", "Clinic, didactics, exam, conference, leave, or another event that limits assignment."),
            ("Hard constraint", "Condition the system must satisfy unless a specifically authorized override exists."),
            ("Soft constraint", "Preference or fairness goal optimized when feasible."),
            ("Coverage location", "Unit, floor, team, or responsibility that must be staffed."),
            ("Assignment type", "Reusable work/activity template with time, category, hours, and scheduling effects."),
            ("Published version", "Immutable resident-visible schedule snapshot."),
        ],
        [1.65, 5.03],
        8.1,
    )

    doc.add_heading("Appendix B. Requirement Traceability Summary", level=1)
    add_table(
        doc,
        ["Business objective", "Primary requirement groups"],
        [
            ("Reduce manual work", "FR-PRO, FR-REQ, FR-MST, FR-GEN, FR-SWX"),
            ("Protect coverage and eligibility", "FR-SVC, FR-MST, FR-GEN, BR-001 through BR-009"),
            ("Improve fairness", "FR-GEN-008, FR-ANA, BR-010 through BR-011"),
            ("Support institutional variation", "FR-BLK, FR-PRO, FR-SVC"),
            ("Preserve chief control", "FR-REQ, FR-GEN-009, FR-PUB, audit requirements"),
            ("Improve resident experience", "FR-PUB, FR-SWX, privacy requirements"),
            ("Enable safe implementation", "NFRs, validation strategy, roadmap, open decisions"),
        ],
        [2.1, 4.58],
    )

    doc.core_properties.title = "Clarity Schedule Business and System Requirements"
    doc.core_properties.subject = "Implementation baseline for residency scheduling platform"
    doc.core_properties.author = "Business Analysis"
    doc.core_properties.keywords = "residency scheduling, requirements, business analysis, SRS, BRD"
    doc.save(DOCX_PATH)
    return doc


def pdf_header_footer(canvas, doc):
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#D9DCE7"))
    canvas.setLineWidth(0.5)
    canvas.line(0.72 * inch, 10.35 * inch, 7.78 * inch, 10.35 * inch)
    canvas.setFont("Helvetica-Bold", 7.5)
    canvas.setFillColor(colors.HexColor("#697087"))
    canvas.drawString(0.72 * inch, 10.48 * inch, "CLARITY SCHEDULE | BUSINESS AND SYSTEM REQUIREMENTS")
    canvas.setFont("Helvetica", 7.5)
    canvas.drawString(0.72 * inch, 0.38 * inch, "Implementation baseline | Confidential working document")
    canvas.drawRightString(7.78 * inch, 0.38 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf():
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "BodyX", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.4,
        leading=11, textColor=colors.HexColor("#20263A"), spaceAfter=5,
    )
    h1 = ParagraphStyle(
        "H1X", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=17,
        leading=20, textColor=colors.HexColor("#4F4698"), spaceBefore=10, spaceAfter=7,
    )
    h2 = ParagraphStyle(
        "H2X", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12,
        leading=14, textColor=colors.HexColor("#292442"), spaceBefore=8, spaceAfter=5,
    )
    h3 = ParagraphStyle(
        "H3X", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=9.5,
        leading=11, textColor=colors.HexColor("#4F4698"), spaceBefore=6, spaceAfter=3,
    )
    small = ParagraphStyle("SmallX", parent=body, fontSize=7.2, leading=9)
    title_style = ParagraphStyle(
        "TitleX", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=27,
        leading=31, textColor=colors.HexColor("#292442"), alignment=TA_LEFT, spaceAfter=8,
    )
    subtitle_style = ParagraphStyle(
        "SubtitleX", parent=body, fontName="Helvetica", fontSize=12, leading=15,
        textColor=colors.HexColor("#697087"), spaceAfter=14,
    )
    bullet = ParagraphStyle(
        "BulletX", parent=body, leftIndent=14, firstLineIndent=-7, bulletIndent=3, spaceAfter=3,
    )

    class NumberedDocTemplate(BaseDocTemplate):
        pass

    pdf = NumberedDocTemplate(
        str(PDF_PATH), pagesize=letter,
        leftMargin=0.72 * inch, rightMargin=0.72 * inch,
        topMargin=0.72 * inch, bottomMargin=0.62 * inch,
        title="Clarity Schedule Executive Implementation Summary",
        author="Business Analysis",
    )
    frame = Frame(pdf.leftMargin, pdf.bottomMargin, pdf.width, pdf.height, id="normal")
    pdf.addPageTemplates(PageTemplate(id="all", frames=frame, onPage=pdf_header_footer))
    story = []

    def P(text, style=body):
        story.append(Paragraph(text, style))

    def H(text, level=1):
        P(text, {1: h1, 2: h2, 3: h3}[level])

    def bullets(items):
        for item in items:
            if isinstance(item, tuple):
                text = f"<b>{item[0]}:</b> {item[1]}"
            else:
                text = item
            story.append(Paragraph("• " + text, bullet))

    def table(headers, rows, widths=None, font=7.0):
        data = [[Paragraph(f"<b>{x}</b>", small) for x in headers]]
        for row in rows:
            data.append([Paragraph(str(x), ParagraphStyle("cell", parent=small, fontSize=font, leading=font + 2)) for x in row])
        if widths:
            col_widths = [x * inch for x in widths]
        else:
            col_widths = None
        t = Table(data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F4698")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D9DCE7")),
            ("BACKGROUND", (0, 2), (-1, -1), colors.HexColor("#FAFAFD")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.extend([t, Spacer(1, 7)])

    P("IMPLEMENTATION BASELINE", ParagraphStyle("kick", parent=small, textColor=colors.HexColor("#2F9B8F"), fontName="Helvetica-Bold", spaceAfter=4))
    P("Clarity Schedule", title_style)
    P("Business Requirements Document and Software Requirements Specification", subtitle_style)
    table(
        ["Document attribute", "Value"],
        [
            ("Version", "1.0 implementation baseline"),
            ("Date", date.today().strftime("%B %d, %Y")),
            ("Prepared for", "Program leadership, chief residents, implementation team, and institutional stakeholders"),
            ("Status", "Draft for validation and implementation planning"),
        ],
        [1.55, 5.35],
        8,
    )
    P("<b>Purpose.</b> Define the processes, behavior, data, rules, controls, and delivery plan required to move the concept from prototype to production.", body)
    P("<b>Governance.</b> The system recommends and validates; authorized program leadership approves policies, exceptions, and published schedules.", body)
    if ASSETS.joinpath("chief-overview-main.png").exists():
        story.append(Image(str(ASSETS / "chief-overview-main.png"), width=6.8 * inch, height=3.85 * inch))
    story.append(PageBreak())

    # The PDF is a concise but complete reading copy. The DOCX remains the authoritative detailed matrix.
    sections = [
        ("1. Executive Summary", [
            "Clarity Schedule is a configurable residency scheduling platform with separate chief and resident portals. It connects annual rotation planning, resident availability, institution patterns, service rules, monthly generation, editable schedules, approvals, publication, analytics, and call exchanges.",
            "The product promise is to enter reliable facts once, collect requests in one place, generate explainable drafts, allow chief editing, and synchronize every downstream view.",
        ]),
        ("2. Business Problems Solved", [
            "Fragmented forms, spreadsheets, PDFs, messages, and institutional calendars create duplicate work and missed constraints.",
            "Each service has different staffing, role, unit, shift, night, weekend, and call requirements.",
            "Manual fairness and call-switch review is slow and difficult to reproduce.",
            "Annual rotation decisions and monthly service schedules are not reliably linked.",
        ]),
        ("3. Users and Access", [
            "Chief administrators configure, generate, edit, approve, publish, and report.",
            "Coordinators maintain profiles, imports, events, deadlines, and communications under delegated permissions.",
            "Residents see their own private schedule and requests plus read-only published service rosters.",
            "Program directors and read-only leaders receive oversight access without unnecessary private details.",
        ]),
        ("4. End-to-End Workflow", [
            "Configure program, blocks, rotations, services, roles, shifts, institutions, and rules.",
            "Import residents and collect annual and block requests.",
            "Approve requests, build the annual master schedule, and validate capacity/transitions.",
            "Pull each service roster from the selected master-schedule block and generate independent drafts.",
            "Edit schedules with live coverage, workload, fairness, and protected-time checks.",
            "Publish controlled versions and support resident call-switch proposals with chief approval.",
        ]),
    ]
    for title, paras in sections:
        H(title, 1)
        for p in paras:
            P(p)
    H("5. Functional Capability Map", 1)
    table(
        ["Domain", "Required capability"],
        [
            ("Program configuration", "Custom years, block dates, request deadlines, services, rotations, roles, shifts, locations, and colors."),
            ("Profiles", "Resident facts, institution, PGY, clinic, didactics, eligibility, leave, and exceptions."),
            ("Requests", "Weekend, day off, vacation, elective, fellowship, PTO, sick leave, special circumstances, and decision workflow."),
            ("Master schedule", "Editable annual grid, capacities, requirements, preference scoring, locks, conflicts, and versions."),
            ("Service scheduling", "Master-linked roster, per-service generator, four-week editor, call pool, live checks, undo/redo, and versions."),
            ("Resident portal", "Private annual/monthly schedule, calls, protected time, requests, PTO, and published rosters."),
            ("Call switches", "Offer marketplace, eligibility checks, volunteer, chief decision, and atomic correction."),
            ("Analytics", "Hours, calls, nights, weekends, golden weekends, clinics, role coverage, fairness, and readiness."),
        ],
        [1.55, 5.35],
        7.5,
    )
    story.append(PageBreak())
    H("6. Requirement Baseline", 1)
    req_groups = [
        ("Identity and access", "Institutional authentication, role-based permissions, private resident boundaries, tenant/program segregation, and full audit history."),
        ("Blocks", "13, 26, monthly, or custom block models; deadlines; holidays; protected program events; lifecycle status."),
        ("Institutions", "Reusable clinic/didactic patterns, buffers, eligibility, automatic inheritance, and individual/block overrides."),
        ("Services", "Custom services mapped to master rotations with locations, roles, counts, shifts, hours, colors, and versioned rules."),
        ("Requests", "Only approved requests constrain generation; competing requests receive advisory priority information and chief decision."),
        ("Master schedule", "One rotation per resident/block, capacity and annual requirements, preference context, transition checks, locks, and versions."),
        ("Generation", "Hard constraints first, soft optimization second, explicit infeasibility, explainable selection, and supplemental eligibility."),
        ("Editing", "Drag, move, copy, clear, custom assignments, add/remove resident, undo/redo, autosave, versions, notes, and live recalculation."),
        ("Publishing", "Ready/published/corrected states, immutable snapshots, selective publication, resident notification, PDF/XLSX export."),
        ("Call switches", "Clinic-date, next-day clinic, back-to-back call, duplicate, leave, eligibility, rest, and policy checks before approval."),
    ]
    table(["Area", "Mandatory requirement"], req_groups, [1.55, 5.35], 7.4)
    H("7. Rule and Optimization Model", 1)
    P("Rules shall be classified as hard, overrideable hard, warning, or optimization objective. Each rule has an owner, scope, effective dates, version, severity, explanation, and test cases.")
    bullets([
        "Legal/program safety, approved leave, staffing, and eligibility precede preferences.",
        "Protected clinics and didactics apply from institution or resident profiles.",
        "Fairness compares configurable peer groups rather than unrelated roles.",
        "Chief overrides require authority, reason, and immutable audit history.",
        "Published schedule corrections create a new version and notifications.",
    ])
    H("8. Conceptual Data Model", 1)
    table(
        ["Entity", "Purpose"],
        [
            ("Program / academic year / block", "Scheduling scope, dates, deadlines, lifecycle"),
            ("Resident / institution", "Facts, patterns, eligibility, and overrides"),
            ("Rotation / master assignment", "Annual placement and capacity"),
            ("Service template / role / location / assignment type", "Reusable monthly scheduling model"),
            ("Request / protected event / decision", "Resident input and controlled approval"),
            ("Schedule version / daily assignment", "Editable and published operational records"),
            ("Call switch / eligibility result", "Controlled exchange workflow"),
            ("Audit event / notification", "Traceability and communication"),
        ],
        [2.7, 4.2],
        7.6,
    )
    H("9. Nonfunctional Requirements", 1)
    bullets([
        ("Security", "encryption, least privilege, secure sessions, dependency and vulnerability management."),
        ("Privacy", "minimum necessary access, restricted sensitive notes, audited exports, retention and deletion policy."),
        ("Reliability", "atomic approval/publication, backups, recovery, optimistic locking, and immutable history."),
        ("Performance", "responsive editing and visible progress for generation jobs."),
        ("Accessibility", "WCAG 2.2 AA target, keyboard use, contrast, screen-reader labels, and zoom."),
        ("Maintainability", "modular rule engine, versioned APIs, automated tests, documented migrations, and observability."),
    ])
    H("10. Implementation Architecture", 1)
    bullets([
        "Authenticated web client for chief and resident portals.",
        "Application API enforcing tenant, program, role, and record-level access.",
        "Relational database with normalized scheduling and audit records.",
        "Rule-evaluation service and asynchronous optimization worker.",
        "Notification worker and object storage for import/export artifacts.",
        "Monitoring, error tracking, security audit logging, backups, and deployment pipeline.",
    ])
    story.append(PageBreak())
    H("11. Acceptance and Testing", 1)
    bullets([
        "Configure a program and custom services without code changes.",
        "Import residents and institution patterns with validation preview.",
        "Approve requests and prove pending requests do not affect generation.",
        "Build a valid master schedule with capacity and transition feedback.",
        "Generate distinct service schedules with independent rules and rosters.",
        "Edit assignments and immediately recalculate coverage and fairness.",
        "Publish schedules with correct private/shared resident access.",
        "Complete an approved call switch and preserve original history.",
        "Export schedules and retrieve a complete audit trail.",
    ])
    H("12. Delivery Roadmap", 1)
    table(
        ["Phase", "Output"],
        [
            ("0. Discovery", "Policy workshops, data dictionary, security review, signed requirements"),
            ("1. Foundation", "Identity, program model, profiles, services, rules, imports"),
            ("2. Master schedule", "Annual grid, capacity, preferences, validation, versions"),
            ("3. Scheduling MVP", "Generator, editor, coverage, analytics, publishing, export"),
            ("4. Resident workflows", "Private views, requests, leave, notifications"),
            ("5. Call switches", "Marketplace, eligibility, approval, controlled correction"),
            ("6. Production hardening", "SSO, monitoring, backups, accessibility, migration"),
        ],
        [1.8, 5.1],
        7.6,
    )
    H("13. Open Decisions", 1)
    bullets([
        "Hard versus overrideable rules by service.",
        "Fairness peer groups and objective weights.",
        "Duty-hour policy source and calculation window.",
        "Approval authority for requests, leave, switches, publication, and correction.",
        "Privacy rules for special circumstances and published rosters.",
        "SSO, hosting, retention, backup, and recovery requirements.",
        "Authoritative migration source and pilot integrations.",
    ])
    H("14. Definition of Ready", 1)
    bullets([
        "Named product owner and sponsor.",
        "Approved MVP and pilot program.",
        "Validated service templates and rule classifications.",
        "Approved data dictionary and import templates.",
        "Security/privacy requirements completed.",
        "Acceptance scenarios converted to a prioritized backlog.",
        "Synthetic test dataset available.",
    ])
    pdf.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
